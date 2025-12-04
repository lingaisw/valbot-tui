from agent_plugins.common_agent_imports import *
from agent_plugins.agent_plugin import AgentPlugin
import asyncio
import os
import json
import re
from pathlib import Path
from typing import List, Dict, Union, Optional, Any
import PyPDF2
import os
import json
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
import hashlib
import pickle
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from docx import Document
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
from rich.prompt import Confirm
from rich import box
from rich.table import Table
from rich.live import Live
from rich.panel import Panel
import readchar
import sys
from prompt_toolkit import prompt
from prompt_toolkit.key_binding import KeyBindings
from agent_plugins.common_tools.terminal_tools import list_files, find_files, grep_in_files, show_file_tree, run_shell_command
from agent_plugins.common_tools.file_tools import read_file, FileContent

class MyCustomPlugin(AgentPlugin):
    """
    SPF to PDL Converter Agent (Two-Stage Approach)
    
    This agent converts Structural Pattern Framework (SPF) test sequences to 
    Tessent's Procedural Description Language (PDL) format using a two-stage approach:
    
    STAGE 1 - UNDERSTAND: 
        Analyze the SPF code to extract the test intent, objectives, register/signal 
        operations, and sequence logic. This creates a semantic understanding of what 
        the test is trying to accomplish rather than doing direct syntax translation.
    
    STAGE 2 - GENERATE: 
        Using the understanding from Stage 1 and relevant PDL syntax knowledge from 
        Tessent manuals (via RAG), generate proper PDL code that accomplishes the 
        same test objectives using correct PDL constructs and idioms.
    
    This approach avoids inconsistent direct translation and ensures the generated 
    PDL uses appropriate commands and syntax from the Tessent documentation.
    """
    
    REQUIRED_ARGS = {}
    
    def get_initializer_args(self):
        """Override to conditionally prompt for input_path only for mode 1."""
        # Import here to avoid issues if not available
        from rich.console import Console
        
        console = Console()
        
        # Menu choices
        choices = [
            ("1", "Convert SPF to PDL"),
            ("2", "PDL Expert Consultation (interactive Q&A)")
        ]
        
        selected_index = 0
        
        def render():
            """Render the menu with current selection highlighted."""
            lines = []
            for i, (key, description) in enumerate(choices):
                if i == selected_index:
                    # Highlight selected option with arrow
                    lines.append(f"  [bold green]→[/bold green]   [bold green]{description}[/bold green]")
                else:
                    # Non-selected option
                    lines.append(f"      [dim]{description}[/dim]")
            return "\n".join(lines)
        
        # Display header
        console.print("\n[bold cyan]Select an option (use ↑ ↓ and Enter)[/bold cyan]\n")
        
        # Interactive selection with arrow keys
        try:
            with Live(render(), refresh_per_second=10, console=console) as live:
                while True:
                    key = readchar.readkey()
                    if key == readchar.key.UP:
                        selected_index = (selected_index - 1) % len(choices)
                    elif key == readchar.key.DOWN:
                        selected_index = (selected_index + 1) % len(choices)
                    elif key == readchar.key.ENTER or key == '\r' or key == '\n':
                        break
                    live.update(render())
        except Exception as e:
            # Fallback to simple input if readchar doesn't work
            # console.print(f"[yellow]Arrow key navigation not available: {e}[/yellow]")
            console.print("\n[bold cyan]SPF to PDL Converter Agent[/bold cyan]")
            console.print("[dim]Select a mode to continue[/dim]\n")
            
            table = Table(show_header=False, box=box.ROUNDED, padding=(0, 2))
            table.add_column("Key", style="cyan", width=5)
            table.add_column("Mode", style="white")
            for key, description in choices:
                table.add_row(key, description)
            
            console.print(table)
            console.print()
            
            mode_value = None
            while mode_value not in ['1', '2']:
                mode_value = console.input("[bold yellow]Select an option:[/bold yellow] ").strip()
                if mode_value not in ['1', '2']:
                    console.print("[red]Invalid selection. Please enter 1 or 2.[/red]")
            
            selected_index = int(mode_value) - 1
        
        mode_value = choices[selected_index][0]
        console.print(f"\n[green]✓ Selected:[/green] {choices[selected_index][1]}\n")
        
        # Only ask for input_path if mode is 1
        if mode_value == '1':
            console.print()
            input_path = console.input("[bold yellow]Enter the path to an SPF file or directory:[/bold yellow] ").strip()
            return {
                'mode': mode_value,
                'input_path': input_path
            }
        else:
            return {'mode': mode_value}

    def __init__(self, model, **kwargs):
        super().__init__(model, **kwargs)
        # Initialize Rich console for beautiful formatting
        self.console = Console()
        
        # Initialize RAG Knowledge Base
        agent_dir = Path(__file__).parent
        self.agent_dir = agent_dir  # Store for later reference
        
        # Set up the resources directory for RAG files
        resources_dir = agent_dir / "pdl_agent_resources"
        resources_dir.mkdir(parents=True, exist_ok=True)
        
        try:
            # Initialize RAG database (will load existing or create new)
            # Note: Don't pass cache_dir - let it default to pdf_dir/.rag_cache
            # Note: Use same collection name as valbot_tui.py for compatibility
            self.rag_kb = RAGKnowledgeBase(
                pdf_dir=resources_dir,
                chunk_size=1000,
                chunk_overlap=200,
                collection_name="rag_knowledge_base"
            )
            
        except Exception as e:
            self.console.print(f"[yellow]⚠ RAG initialization failed: {e}[/yellow]")
            self.rag_kb = None
        
        # Create specialized agents for two-stage conversion
        self.understanding_agent = self.create_understanding_agent(model)
        self.pdl_generation_agent = self.create_pdl_generation_agent(model)
        # Legacy: keep for backward compatibility
        self.converter_agent = self.create_agent(model)
        # PDL Expert for interactive consultation
        self.pdl_expert_agent = self.create_pdl_expert_agent(model)
        # Provide tools for agent to use autonomously
        self.add_agentic_tools()

    def get_relevant_knowledge(self, query: str, top_k: int = 5) -> str:
        """
        Retrieve relevant knowledge from RAG system based on query.
        Falls back to full text if RAG is unavailable.
        
        Args:
            query: Search query for relevant context
            top_k: Number of relevant chunks to retrieve
            
        Returns:
            Relevant knowledge context
        """
        if self.rag_kb:
            # Use RAG for semantic search
            context = self.rag_kb.get_context_for_query(
                query=query,
                top_k=top_k,
                include_metadata=True
            )
            return context
        else:
            # Fallback: load full PDF text (legacy method)
            return self._load_legacy_knowledge_base()
    
    def _load_legacy_knowledge_base(self) -> str:
        """
        Legacy method: Load and extract full text from PDFs.
        Used as fallback if RAG system is unavailable.
        """
        knowledge = []
        agent_dir = Path(__file__).parent
        resources_dir = agent_dir / "pdl_agent_resources"
        
        pdf_files = [
            "tessent_shell_reference_manual.pdf",
            "tessent_cell_library_manual.pdf",
            "DTEG_ITPP_Reader_Commands.pdf"
        ]
        
        self.console.print("[bold yellow]Loading knowledge base (legacy mode)...[/bold yellow]")
        
        for pdf_file in pdf_files:
            pdf_path = resources_dir / pdf_file
            if pdf_path.exists():
                try:
                    with open(pdf_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        num_pages = len(pdf_reader.pages)
                        for page_num in range(num_pages):
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text()
                        knowledge.append(f"=== {pdf_file} ===\n{text[:50000]}")
                        self.console.print(f"  [green]✓[/green] Loaded {pdf_file} ({num_pages} pages)")
                except Exception as e:
                    self.console.print(f"  [red]✗[/red] Error loading {pdf_file}: {e}")
        
        return "\n\n".join(knowledge) if knowledge else "No knowledge base available."

    def create_agent(self, model):
        """Create the SPF to PDL conversion agent with specialized system prompt."""
        system_prompt = """
        You are an expert in pre-silicon validation and test pattern languages.
        Your task is to understand SPF test sequences and generate equivalent PDL code.
        
        IMPORTANT: Do NOT attempt direct translation. Instead:
        1. First understand what the SPF is trying to accomplish
        2. Then write PDL using proper PDL constructs for that purpose
        
        You will work in TWO stages with relevant Tessent manual knowledge provided.
        """
        
        return Agent[str](
            model=model,
            system_prompt=system_prompt,
            retries=3
        )
    
    def create_understanding_agent(self, model):
        """Create an agent specialized in understanding SPF intent."""
        system_prompt = """
        You are an expert in analyzing Structural Pattern Framework (SPF) test sequences.
        
        Your task is to analyze SPF code and extract ONLY the test operations and data.
        Do NOT translate to PDL. Do NOT mention any PDL commands or syntax.
        
        Focus ONLY on:
        - Which registers/signals are accessed (names and addresses)
        - What values are written (exact hex/binary values)
        - What values are read/checked (expected values)
        - What timing delays are needed (exact durations)
        - The sequence of operations (step by step)
        - Any loops or conditional logic
        - SPF label values (these will be used for documentation/comments)
        
        CRITICAL: Distinguish between register fields and bit numbers:
        - Register FIELD: Named field in register (e.g., register.FIELD_NAME or register.Enable)
        - Bit NUMBER: Numeric bit position (e.g., register[7] or register[15:8])
        - Do NOT confuse bit numbers with field names
        - Preserve the exact notation used in SPF (brackets for bits, dots for fields)
        
        CRITICAL: Extract exact value formats with proper prefixes:
        - Binary values: MUST have 0b prefix (e.g., 0b1010, 0b0, 0b111)
        - Hexadecimal values: MUST have 0x prefix (e.g., 0x1234, 0xABCD, 0x0)
        - Decimal values: Use 0d prefix explicitly (e.g., 0d10, 0d255) or no prefix for simple decimals
        - Preserve the exact prefix format from SPF
        
        OUTPUT FORMAT:
        === TEST OBJECTIVE ===
        [High-level purpose in one sentence]
        
        === OPERATIONS ===
        [List each operation as: WRITE register.field = value, READ register.field expect value, WAIT time, CALL procedure, etc.]
        [For bit access use: WRITE register[bit] or register[high:low], not register.bit]
        [Include LABEL: "text" for each operation that has an SPF label]
        
        === SEQUENCE ===
        1. [First operation with data] LABEL: "spf_label_text"
        2. [Second operation with data] LABEL: "spf_label_text"
        ...
        
        === TIMING ===
        [Any timing constraints or delays]
        
        CRITICAL RULES:
        1. Focus on WHAT data goes WHERE, not on how to implement it
        2. Extract and preserve any SPF label values - they describe what each step does
        3. PRESERVE EXACT signal paths and register names as they appear in SPF - do NOT modify, shorten, or rename them
        4. Keep the full hierarchical paths exactly as written (e.g., module.submodule.register.field)
        5. DO NOT convert bit numbers to field names - keep [bit] notation for bit access
        6. Use register.FIELD_NAME only when SPF uses named fields
        """
        
        return Agent[str](
            model=model,
            system_prompt=system_prompt,
            retries=2
        )
    
    def create_pdl_generation_agent(self, model):
        """Create an agent specialized in generating PDL from requirements."""
        system_prompt = """
        You are an expert in Tessent Procedural Description Language (PDL).
        
        You will receive:
        1. A description of test operations (register names, values, timing)
        2. Example PDL code and syntax from Tessent manuals
        
        Your task is to write PDL code using ONLY the commands and syntax shown in the examples.
        
        CRITICAL RULES FOR PDL CODE GENERATION:
        1. ONLY use PDL commands that appear in the provided manual excerpts
        2. DO NOT invent or hallucinate PDL functions
        3. Follow the EXACT syntax patterns from the examples
        4. Use the TCL-like syntax shown (iProc, iWrite, iRead, iCall, iApply, etc.)
        5. If a command is not in the examples, do NOT use it
        6. ALWAYS start PDL files with these TWO REQUIRED PDL header lines (in this exact order):
           iProcsForModule [get_single_name [get_current_design -icl]] 
           source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
           NOTE: These headers are PDL-specific, NOT used in SPF
        7. PRIORITIZE iSim commands for signal polling, peeking, and macro mapping
        8. COMPULSORY: Validate iSim command parameters match examples exactly:
           - iSim poll_signal: signal_path expected_value timeout step_size
           - iSim peek_signal: signal_path expected_value
           - iSim macro_map: alias rtl_path
        9. CRITICAL: Use iSim macro_map for repeating signal paths:
           - Identify long or repeated signal hierarchies in the test operations
           - Create macro mappings at the start of procedures using: iSim macro_map alias full_rtl_path
           - Use the backtick prefix (`) when referencing mapped macros in subsequent commands
           - Example pattern from example.pdl:
             iSim macro_map man_fuse_sip_rel_req pcd_tb.pcd.parfuse.parfuse_pwell_wrapper.fuse_top1.i_chassis_fuse_controller_top.i_fuse_array_cntrl.i_fuse_array_cntrl_tap.man_fuse_sip_release_req
             iSim poll_signal `man_fuse_sip_rel_req 0x1 10us 5ns
           - This improves code readability and maintainability
        10. Use iNote "description" in PDL code to document what each section or operation is doing
        10. When converting from SPF, use the SPF label values as descriptions in PDL iNote commands
            NOTE: iNote is a PDL command for documentation, NOT an SPF command
        11. PRESERVE EXACT signal paths and register names from the test operations - do NOT modify, rename, or shorten them
        12. Keep full hierarchical paths exactly as provided (e.g., pcd_tb.pcd.module.register.field)
        13. CRITICAL: Distinguish between register fields and bit numbers:
            - Use register.FIELD_NAME for named fields (e.g., register_name.Enable)
            - DO NOT treat bit numbers as field names
            - If test operations show [bit] notation, that's a bit number, NOT a field name
            - Do NOT convert register[7] to register.7 or register."7"
        14. CRITICAL: Use explicit prefixes for all PDL values:
            - Binary values: ALWAYS use 0b prefix (0b1010, 0b0, 0b111, 0b11111111)
            - Hexadecimal values: ALWAYS use 0x prefix (0x1234, 0xABCD, 0x0, 0xFFFF)
            - Decimal values: Use 0d prefix explicitly (0d10, 0d100) or no prefix for clarity
            - NEVER use values without proper prefixes (e.g., write 0b1 not just 1, write 0x10 not just 10)
            - Check PDL examples for proper value format patterns
        15. COMPULSORY: Check the provided manual examples for EXACT parameter requirements:
            - What parameters each command requires (order, count, type)
            - What format values must be in (hex 0x..., binary 0b..., decimal, string)
            - What data types are expected (integer, string, list, etc.)
            - Match parameter formats EXACTLY as shown in examples
        
        16. CRITICAL - DYNAMIC REGISTER INSTANCE HANDLING (HIGHEST PRIORITY):
            ALWAYS prioritize setting register instances dynamically using get_icl_instances and loops.
            This is STRONGLY PREFERRED over hardcoding individual register names.
            
            ✓ CORRECT APPROACH (Dynamic - USE THIS):
            ```
            set fuse_reg [get_icl_instances -of_modules [get_icl_module fuse_cltap_JTAGControl_DR]]
            foreach_in_collection reg $fuse_reg {
                set reg_name [get_single_name $reg]
                iWrite $reg_name.DIS_DCG 0x0
                iWrite $reg_name.DEFER_SB 0x0
                iWrite $reg_name.CLOCK_MUX 0x0
                iApply
            }
            ```
            
            ✗ AVOID (Hardcoded - Only use if absolutely necessary):
            ```
            iWrite SECUREDSOCTAPC_MISCDFXCFG.core_out_sel 2
            ```
            
            RULES FOR DYNAMIC APPROACH:
            - When accessing registers of the same module/type, use get_icl_instances to get all instances
            - Loop through instances using foreach_in_collection
            - Extract instance name with get_single_name
            - Apply operations to $reg_name variable (dynamically resolved)
            - Benefits: Scalable, maintainable, works with multiple instances automatically
            - ONLY use hardcoded register names when:
              * Operating on a single specific register that won't scale
              * The register is truly unique (not part of a repeated module)
            - If test operations show multiple registers of similar type, ALWAYS use dynamic approach
            - Look for patterns in register names that indicate module instances (common prefixes/suffixes)
        
        Common PDL commands you should look for in examples (PRIORITIZE iSim commands):
        - iSim commands (CHECK THESE FIRST)  # Simulation commands (DTEG dftPdl) - PRIORITY
          - iSim poll_signal signal_path expected_value timeout step_size
          - iSim peek_signal signal_path expected_value
          - iSim macro_map alias rtl_path  # Define short aliases for long signal paths
            * Place macro_map at start of procedures (often in initialization iProc like fuse_rtl_map)
            * Reference with backtick: `alias_name
            * Example: iSim macro_map my_signal pcd_tb.pcd.module.submodule.long_signal_name
            * Then use: iSim poll_signal `my_signal 0x1 10us 5ns
        - iProc <name> {args} { ... }  # Define procedure
        - iWrite register.field value   # Write to register field (prefer dynamic: $reg_name.field)
        - iRead register.field value    # Read and check register field (prefer dynamic: $reg_name.field)
        - iApply                        # Apply the register operations
        - iCall procedure args          # Call another procedure
        - iRunLoop count -tck           # Run clock cycles
        - iNote "message"               # Add comment/note in PDL (USE THIS FREQUENTLY IN PDL CODE)
        - get_icl_instances             # Get register/module instances (USE FOR DYNAMIC ACCESS)
        - get_icl_module                # Get module definition (USE WITH get_icl_instances)
        - get_single_name               # Extract name from instance object (USE IN LOOPS)
        - foreach_in_collection         # Loop through collections (USE FOR DYNAMIC REGISTER ACCESS)
        - set variable_name value       # Set TCL variable (USE TO STORE INSTANCE COLLECTIONS)
        
        BEFORE writing code, CHECK the examples for:
        - FIRST: Can registers be accessed dynamically? (get_icl_instances + foreach_in_collection)
        - SECOND: iSim command parameters (poll_signal, peek_signal, macro_map)
        - Opportunities to use macro_map for repeated/long signal paths
        - Opportunities to use dynamic register access instead of hardcoding
        - Parameter count and order for each command
        - Value formats (0x for hex, 0b for binary, decimal, quotes for strings)
        - Required vs optional parameters
        - Proper syntax for loops, conditionals, variable assignments
        
        OUTPUT FORMAT:
        === PDL CODE ===
        iProcsForModule [get_single_name [get_current_design -icl]] 
        source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
        
        [Rest of PDL code using ONLY commands from the manual examples]
        [PRIORITIZE iSim commands where applicable]
        [Include iNote "description" statements throughout to document PDL operations]
        [NOTE: iNote is PDL-specific syntax for documentation, not used in SPF]
        [Use EXACT signal paths and register names as provided in test operations]
        [Use EXACT parameter formats as shown in examples]
        [Validate iSim command parameters match examples]
        
        === IMPLEMENTATION NOTES ===
        [List which PDL commands you used and where you found them in the examples]
        [If iSim commands used, confirm their parameters match examples exactly]
        [Confirm parameter formats match the examples (hex/binary/decimal/string)]
        """
        
        return Agent[str](
            model=model,
            system_prompt=system_prompt,
            retries=3
        )
    
    def create_pdl_expert_agent(self, model):
        """Create an interactive PDL expert agent for consultation."""
        system_prompt = """
        You are an expert consultant on Tessent Procedural Description Language (PDL).
        
        Users will ask you questions about PDL syntax, commands, best practices, or how to 
        implement specific test scenarios. You will be provided with relevant excerpts from 
        Tessent manuals, DTEG ITPP Reader Commands documentation, and example PDL code to answer their questions.
        
        Your role:
        - Answer questions about PDL syntax and commands
        - Explain how PDL constructs work
        - Provide code examples using ONLY commands from the manual excerpts
        - Help debug PDL issues
        - Suggest best practices for test implementation
        - Guide users on proper PDL patterns
        - Explain DTEG iSim commands for simulation control
        
        IMPORTANT RULES:
        1. Base your answers on the provided manual excerpts and examples
        2. Use ONLY PDL commands that exist in the documentation
        3. Cite which manual section or example you're referencing
        4. When referencing files from the knowledge base (e.g., example.pdl), ALWAYS include the full file path for user reference
        5. If you don't have information in the provided context, say so
        6. Provide clear, actionable code examples when relevant
        7. ALWAYS include the required TWO PDL header lines in PDL code examples:
           iProcsForModule [get_single_name [get_current_design -icl]] 
           source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
           NOTE: These headers are required for PDL files only, not for SPF files
        8. Use iNote "description" liberally to document what PDL code is doing
           NOTE: iNote is a PDL command for documentation, not used in SPF
        9. CRITICAL BEST PRACTICE - PRIORITIZE DYNAMIC REGISTER ACCESS:
           When users ask about register access patterns, ALWAYS recommend the dynamic approach:
           - Use get_icl_instances to get register/module instances dynamically
           - Loop through instances with foreach_in_collection
           - Extract names with get_single_name
           - Apply operations to $reg_name variable
           - This approach is STRONGLY PREFERRED over hardcoding individual register names
           - Benefits: Scalable, maintainable, works with multiple instances automatically
           
           CORRECT (Dynamic - Always recommend this):
           ```pdl
           set fuse_reg [get_icl_instances -of_modules [get_icl_module fuse_cltap_JTAGControl_DR]]
           foreach_in_collection reg $fuse_reg {
               set reg_name [get_single_name $reg]
               iWrite $reg_name.DIS_DCG 0x0
               iWrite $reg_name.DEFER_SB 0x0
               iApply
           }
           ```
           
           AVOID (Hardcoded - Only for truly unique registers):
           ```pdl
           iWrite SECUREDSOCTAPC_MISCDFXCFG.core_out_sel 2
           ```
           
           Explain that hardcoding should only be used for truly unique, non-repeating registers.
        
        CRITICAL: Format ALL responses using proper Markdown:
        - Use ### for section headers
        - Use ```pdl for PDL code blocks
        - Use `command` for inline command names
        - Use **bold** for emphasis
        - Use bullet points (- or *) for lists
        - Use > for important notes or warnings
        
        Example response format:
        The `iWrite` command is used to write values to register fields.
        
        ```pdl
        iProcsForModule [get_single_name [get_current_design -icl]] 
        source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
        
        iNote "Write to register field"
        iWrite register_name.field_name value
        iApply
        ```
        
        **Key points:**
        - Multiple `iWrite` commands can be grouped
        - `iApply` executes the writes
        
        > **Note:** Always call `iApply` after `iWrite` operations
        
        **Reference:** example.pdl  
        **Full Path:** C:/path/to/agents/example.pdl (line 45)
        
        Common PDL commands to reference:
        - iProc, iWrite, iRead, iApply, iCall, iRunLoop, iNote
        - Dynamic access (PRIORITIZE THESE): get_icl_instances, get_icl_module, get_single_name, foreach_in_collection
        - iSim commands (DTEG dftPdl):
          - iSim poll_signal signal_path expected_value timeout step_size
          - iSim peek_signal signal_path expected_value
          - iSim macro_map alias rtl_path
            * Use to create short aliases for long/repeated signal paths
            * Define at start of procedures (e.g., in initialization iProc)
            * Reference with backtick prefix: `alias_name
        - TCL constructs: set, foreach_in_collection, if/else, for
        
        Be helpful, accurate, and always ground your answers in the documentation.
        Use proper Markdown formatting in ALL responses.
        """
        
        return Agent[str](
            model=model,
            system_prompt=system_prompt,
            retries=2
        )
    
    def add_agentic_tools(self):
        """
        Add autonomous tools for the agent to use during conversion.
        These tools help with file parsing and structure analysis.
        """
        # Tools can be added here if needed for autonomous operation
        self.pdl_expert_agent.tool(list_files)
        self.pdl_expert_agent.tool(find_files)
        self.pdl_expert_agent.tool(grep_in_files)
        self.pdl_expert_agent.tool(show_file_tree)
        self.pdl_expert_agent.tool(run_shell_command)
        self.pdl_expert_agent.tool(read_file)

    def find_spf_files(self, input_path: str) -> List[str]:
        """
        Find all SPF files in the given path.
        
        Args:
            input_path: File path or directory path
            
        Returns:
            List of SPF file paths
        """
        spf_files = []
        path = Path(input_path)
        
        if path.is_file():
            if path.suffix.lower() == '.spf':
                spf_files.append(str(path))
            else:
                self.console.print(f"[yellow]Warning: {input_path} is not an SPF file[/yellow]")
        elif path.is_dir():
            self.console.print(f"[cyan]Scanning directory:[/cyan] {input_path}")
            spf_files = [str(f) for f in path.rglob('*.spf')]
            self.console.print(f"[green]Found {len(spf_files)} SPF files[/green]")
        else:
            self.console.print(f"[red]Error: Path does not exist: {input_path}[/red]")
        
        return spf_files

    def read_spf_file(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Read and parse SPF file structure.
        
        Args:
            file_path: Path to SPF file
            
        Returns:
            Dictionary containing file metadata and content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Basic structure analysis
            spf_data = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "content": content,
                "line_count": len(content.split('\n')),
                "has_macros": "MACRO" in content.upper(),
                "has_vectors": "VECTOR" in content.upper(),
                "has_scan": "SCAN" in content.upper(),
                "size_bytes": len(content)
            }
            
            return spf_data
        except Exception as e:
            self.console.print(f"[red]Error reading {file_path}: {e}[/red]")
            return None

    def analyze_spf_structure(self, spf_data: Dict) -> str:
        """
        Analyze SPF file structure to provide context for conversion.
        
        Args:
            spf_data: SPF file data dictionary
            
        Returns:
            Analysis summary string
        """
        analysis = f"""
        SPF File Analysis:
        - File: {spf_data['file_name']}
        - Lines: {spf_data['line_count']}
        - Size: {spf_data['size_bytes']} bytes
        - Contains Macros: {spf_data['has_macros']}
        - Contains Vectors: {spf_data['has_vectors']}
        - Contains Scan Operations: {spf_data['has_scan']}
        """
        
        # Extract key patterns
        content = spf_data['content']
        patterns = {
            'CALL': len(re.findall(r'\bCALL\b', content, re.IGNORECASE)),
            'MACRO': len(re.findall(r'\bMACRO\b', content, re.IGNORECASE)),
            'IOPIN': len(re.findall(r'\bIOPIN\b', content, re.IGNORECASE)),
            'TIMING': len(re.findall(r'\bTIMING\b', content, re.IGNORECASE)),
        }
        
        if any(patterns.values()):
            analysis += "\n        Key Commands Found:\n"
            for cmd, count in patterns.items():
                if count > 0:
                    analysis += f"        - {cmd}: {count} occurrences\n"
        
        return analysis

    def convert_spf_to_pdl(self, spf_data: Dict) -> Dict[str, Any]:
        """
        Convert SPF content to PDL using two-stage AI approach:
        1. Understand what the SPF does
        2. Generate PDL based on that understanding
        
        Args:
            spf_data: SPF file data dictionary
            
        Returns:
            Dictionary containing conversion results
        """
        # Analyze structure first
        analysis = self.analyze_spf_structure(spf_data)
        
        try:
            # ====================================================================
            # STAGE 1: UNDERSTAND SPF INTENT
            # ====================================================================
            self.console.print(f"  [cyan]Stage 1:[/cyan] Understanding SPF intent...")
            
            understanding_prompt = f"""
            Analyze the following SPF test code and explain what it does.
            Focus on the test objectives, register/signal operations, and sequence.
            
            {analysis}
            
            SPF CONTENT:
            {spf_data['content']}
            
            Provide a clear description of the test's purpose and operations.
            """
            
            understanding_result = asyncio.run(self.understanding_agent.run(understanding_prompt))
            test_understanding = understanding_result.data
            
            self.console.print(f"  [green]✓[/green] SPF intent analyzed")
            
            # ====================================================================
            # STAGE 2: GENERATE PDL FROM UNDERSTANDING
            # ====================================================================
            self.console.print(f"  [cyan]Stage 2:[/cyan] Generating PDL from understanding...")
            
            # Extract key information to query RAG
            # Create a better search query based on the operations found
            search_terms = []
            content_lower = spf_data['content'].lower()
            
            # Identify key operations to search for in PDL manual
            if 'scan' in content_lower:
                search_terms.append('scan')
            if 'vector' in content_lower or 'pattern' in content_lower:
                search_terms.append('vector pattern')
            if 'macro' in content_lower or 'call' in content_lower:
                search_terms.append('iProc procedure call iCall')
            if 'iopin' in content_lower or 'signal' in content_lower:
                search_terms.append('iWrite iRead signal register')
            if 'timing' in content_lower or 'clock' in content_lower:
                search_terms.append('iRunLoop timing clock wait')
            if 'register' in content_lower or 'reg' in content_lower:
                search_terms.append('iWrite iRead iApply register')
            
            # Always search for core PDL syntax
            search_terms.append('iProc iWrite iRead iApply iCall iRunLoop iNote')
            
            # Build RAG query - PRIORITIZE iSim commands
            rag_query = f"iSim poll_signal peek_signal macro_map DTEG simulation commands {' '.join(search_terms)} PDL syntax examples"
            
            # Retrieve relevant PDL knowledge (example.pdl is now in RAG)
            self.console.print(f"  [dim]Retrieving PDL syntax examples from knowledge base...[/dim]")
            relevant_pdl_knowledge = self.get_relevant_knowledge(rag_query, top_k=12)
            
            # Generate PDL based on understanding and knowledge
            pdl_generation_prompt = f"""
            Write PDL code to implement the following test operations.
            
            TEST OPERATIONS TO IMPLEMENT:
            {test_understanding}
            
            PDL SYNTAX EXAMPLES FROM TESSENT MANUALS, DTEG ITPP READER COMMANDS, AND example.pdl:
            {relevant_pdl_knowledge}
            
            CRITICAL INSTRUCTIONS:
            1. Study the PDL examples above carefully
            2. Use ONLY the commands and syntax shown in those examples
            3. Follow the exact patterns you see (iProc, iWrite, iRead, iCall, iApply, etc.)
            4. DO NOT invent commands - if you don't see it in the examples, don't use it
            5. Match the TCL-like syntax structure from the examples
            6. ALWAYS start with these TWO header lines (in this exact order):
               iProcsForModule [get_single_name [get_current_design -icl]] 
               source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
            7. PRIORITIZE using iSim commands when applicable (poll_signal, peek_signal, macro_map)
            8. COMPULSORY: Validate iSim command parameters from examples:
               - iSim poll_signal: requires signal_path, expected_value, timeout, step_size
               - iSim peek_signal: requires signal_path, expected_value
               - iSim macro_map: requires alias, rtl_path
               - Check examples for exact parameter formats and values
            9. CRITICAL: Use iSim macro_map for repeating signal paths:
               - If signal hierarchies are long or used multiple times, create macro mappings
               - Place macro_map commands at the start of procedures (typically in an initialization iProc)
               - Pattern: iSim macro_map short_alias full.hierarchical.path.to.signal
               - Reference mapped signals with backtick prefix: `short_alias
               - Example from example.pdl:
                 iSim macro_map man_fuse_sip_rel_req pcd_tb.pcd.parfuse.parfuse_pwell_wrapper.fuse_top1.i_chassis_fuse_controller_top.i_fuse_array_cntrl.i_fuse_array_cntrl_tap.man_fuse_sip_release_req
                 iSim poll_signal `man_fuse_sip_rel_req 0x1 10us 5ns
               - This makes code cleaner and easier to maintain
            10. Use iNote "description" in your PDL code to document what each section is doing
            10. If converting from SPF, use the SPF label values in your PDL iNote descriptions
                NOTE: iNote is a PDL command - it's not used in SPF source files
            11. CRITICAL - PRIORITIZE DYNAMIC REGISTER ACCESS:
                - BEFORE writing any iWrite/iRead commands, check if registers can be accessed dynamically
                - If test operations show multiple registers from the same module, use get_icl_instances
                - Pattern: set reg_collection [get_icl_instances -of_modules [get_icl_module MODULE_NAME]]
                - Then: foreach_in_collection reg $reg_collection {{{{ set reg_name [get_single_name $reg]; ... }}}}
                - Only hardcode register names when dealing with truly unique, non-repeating registers
                - Dynamic approach is ALWAYS preferred for scalability and maintainability
            12. PRESERVE EXACT signal paths and register names - do NOT modify, rename, or shorten them
            13. Copy signal/register names character-for-character from the test operations description
            14. CRITICAL: Distinguish bit numbers from field names:
                - If test operations show register[7] or register[15:8], that's BIT ACCESS, not a field
                - Use register.FIELD_NAME only for actual named fields (e.g., register.Enable)
                - DO NOT convert bit numbers to field names (register[7] ≠ register.7)
                - Check PDL examples for proper bit/field access syntax
            15. CRITICAL: Use explicit prefixes for ALL PDL values:
                - Binary values: ALWAYS use 0b prefix (0b1, 0b0, 0b1010, 0b11111111)
                - Hexadecimal values: ALWAYS use 0x prefix (0x0, 0x1234, 0xABCD, 0xFFFFFFFF)
                - Decimal values: Use 0d prefix or no prefix (0d10, 0d255, or 100)
                - NEVER write ambiguous values - always make the base explicit
                - Examples: Write 0b1 not 1, write 0x10 not 10 (unless clearly decimal)
            16. COMPULSORY: Check examples for parameter requirements BEFORE writing each command:
                - How many parameters does each command need?
                - What format should values be in? (0x1234 for hex, 0b1010 for binary, decimal, "string")
                - What is the correct parameter order?
                - Match parameter formats EXACTLY as shown in examples
            
            Look at the examples to understand:
            - FIRST: How to access registers dynamically (get_icl_instances, foreach_in_collection)
            - SECOND: How to use iSim commands (poll_signal, peek_signal, macro_map)
            - How to define procedures (iProc)
            - How to write registers dynamically (iWrite $reg_name.field)
            - How to read registers dynamically (iRead $reg_name.field)
            - How to apply operations (iApply)
            - How to add timing (iRunLoop)
            - How to call procedures (iCall)
            - How to add descriptive notes (iNote "description")
            - What value formats each command expects (hex 0x, binary 0b, decimal, strings)
            
            Format your response EXACTLY as:
            === PDL CODE ===
            iProcsForModule [get_single_name [get_current_design -icl]] 
            source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl
            
            [Rest of PDL code using ONLY commands from the examples above]
            [PRIORITIZE dynamic register access using get_icl_instances + foreach_in_collection loops]
            [ONLY hardcode register names if they are truly unique and non-repeating]
            [PRIORITIZE iSim commands where applicable]
            [Use iSim macro_map for any repeated or long signal paths - define mappings early]
            [Reference mapped signals with backtick prefix: `alias_name]
            [Include iNote statements to describe what each PDL section does]
            [NOTE: The two header lines and iNote are PDL-specific, not used in SPF]
            [Use EXACT signal paths and register names from test operations - copy them exactly]
            [Use correct value formats for each parameter - check examples first]
            [Ensure iSim commands have all required parameters]
            [Use register.FIELD_NAME for named fields, NOT for bit numbers]
            [Use explicit prefixes: 0b for binary, 0x for hex, 0d for decimal]
            
            === IMPLEMENTATION NOTES ===
            [FIRST: State whether dynamic register access was used and justify the approach]
            [If dynamic: List the get_icl_instances and foreach_in_collection patterns used]
            [If hardcoded: Explain why hardcoding was necessary for this specific case]
            [List which PDL commands you used and confirm they came from the examples]
            [If iSim commands used, list them and confirm parameters are complete and correct]
            [If iSim macro_map used, list the mappings created and why they were beneficial]
            [Confirm that all signal paths and register names match the test operations exactly]
            [Confirm that all parameter formats match the examples (hex/binary/decimal/string)]
            [List what parameter formats you used for each command]
            [Confirm bit numbers vs field names are correctly distinguished]
            [Confirm all values use explicit prefixes (0b, 0x, 0d)]
            """
            
            pdl_result = asyncio.run(self.pdl_generation_agent.run(pdl_generation_prompt))
            
            self.console.print(f"  [green]✓[/green] PDL code generated")
            
            # Parse the result
            pdl_code = ""
            implementation_notes = ""
            
            response = pdl_result.data
            if "=== PDL CODE ===" in response:
                parts = response.split("=== PDL CODE ===")
                if len(parts) > 1:
                    rest = parts[1]
                    if "=== IMPLEMENTATION NOTES ===" in rest:
                        code_part, notes_part = rest.split("=== IMPLEMENTATION NOTES ===", 1)
                        pdl_code = code_part.strip()
                        implementation_notes = notes_part.strip()
                    else:
                        pdl_code = rest.strip()
            else:
                # If format not followed, use entire response as code
                pdl_code = response
            
            # Combine both token usage
            total_tokens = understanding_result.usage().total_tokens + pdl_result.usage().total_tokens
            
            return {
                "success": True,
                "spf_file": spf_data['file_path'],
                "test_understanding": test_understanding,
                "pdl_code": pdl_code,
                "implementation_notes": implementation_notes,
                "conversion_notes": f"Two-stage conversion:\n\n{test_understanding}\n\n{implementation_notes}",
                "tokens_used": total_tokens
            }
        except Exception as e:
            self.console.print(f"  [red]Conversion failed: {e}[/red]")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "spf_file": spf_data['file_path'],
                "error": str(e)
            }

    def save_pdl_file(self, conversion_result: Dict, output_dir: Optional[str] = None) -> Optional[str]:
        """
        Save converted PDL to file along with understanding and notes.
        
        Args:
            conversion_result: Conversion result dictionary
            output_dir: Optional output directory (defaults to current working directory)
            
        Returns:
            Path to saved PDL file
        """
        if not conversion_result.get('success'):
            return None
        
        spf_path = Path(conversion_result['spf_file'])
        
        # Determine output directory
        if output_dir:
            out_dir = Path(output_dir)
        else:
            out_dir = Path.cwd()  # Use current working directory
        
        out_dir.mkdir(parents=True, exist_ok=True)
        
        # Create PDL filename
        pdl_filename = spf_path.stem + '.pdl'
        pdl_path = out_dir / pdl_filename
        
        # Write PDL file
        try:
            pdl_code = conversion_result['pdl_code']
            
            # Ensure the required TWO header lines are at the top
            required_line1 = "iProcsForModule [get_single_name [get_current_design -icl]]"
            required_line2 = "source $::env(DUVE_M_HOME)/verif/pdl/common/tap_utils.pdl"
            
            # Remove any leading whitespace
            pdl_code = pdl_code.strip()
            
            # Check if headers are already present
            has_line1 = pdl_code.startswith(required_line1)
            has_line2 = required_line2 in pdl_code.split('\n')[:5]  # Check first 5 lines
            
            if not (has_line1 and has_line2):
                # Remove headers if partially present
                lines = pdl_code.split('\n')
                filtered_lines = []
                for line in lines:
                    if line.strip() != required_line1 and line.strip() != required_line2:
                        filtered_lines.append(line)
                
                # Add both headers at the top
                pdl_code = required_line1 + "\n" + required_line2 + "\n\n" + '\n'.join(filtered_lines)
            
            with open(pdl_path, 'w', encoding='utf-8') as f:
                f.write(pdl_code)
            
            # Save test understanding if available
            if conversion_result.get('test_understanding'):
                understanding_path = out_dir / (spf_path.stem + '_test_understanding.txt')
                with open(understanding_path, 'w', encoding='utf-8') as f:
                    f.write("=" * 80 + "\n")
                    f.write(f"SPF Test Understanding: {spf_path.name}\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(conversion_result['test_understanding'])
                    f.write("\n\n" + "=" * 80 + "\n")
                    f.write("PDL Implementation Notes\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(conversion_result.get('implementation_notes', ''))
            
            # Save conversion notes if available (backward compatibility)
            elif conversion_result.get('conversion_notes'):
                notes_path = out_dir / (spf_path.stem + '_conversion_notes.txt')
                with open(notes_path, 'w', encoding='utf-8') as f:
                    f.write(conversion_result['conversion_notes'])
            
            return str(pdl_path)
        except Exception as e:
            self.console.print(f"[red]Error saving PDL file: {e}[/red]")
            return None

    def run_pdl_expert_mode(self, context, **kwargs):
        """
        Interactive PDL expert consultation mode with conversation history.
        Users can ask questions about PDL and get answers based on the RAG knowledge base.
        The agent remembers previous questions and answers in the session.
        """
        from rich.markdown import Markdown
        
        # Display header
        self.console.rule("[bold blue]PDL Expert Consultation Mode[/bold blue]", style="blue")
        self.console.print("\n[bold cyan]Ask questions about Tessent PDL syntax, commands, and best practices.[/bold cyan]")
        self.console.print("[dim]Tips:[/dim]")
        self.console.print("[dim]  • Press alt+Enter to submit your question or Enter for a new line[/dim]")
        self.console.print("[dim]  • Type 'exit', 'quit', or 'done' to end the session[/dim]")
        self.console.print("[dim]  • The agent remembers your conversation history[/dim]\n")
        
        # Check RAG knowledge base
        if self.rag_kb:
            stats = self.rag_kb.get_stats()
            self.console.print(f"[green]✓ Knowledge base loaded: {stats['total_chunks']} chunks available[/green]")
            # self.console.print(f"[green]✓ example.pdl loaded in knowledge base[/green]\n")
        else:
            self.console.print("[yellow]⚠ RAG knowledge base not available. Using legacy mode.[/yellow]\n")
        
        # Interactive loop with conversation history
        session_history = []
        conversation_context = []  # For maintaining conversation memory
        question_count = 0
        
        while True:
            try:
                # Get user question with multiline support
                # Use prompt_toolkit for proper Shift+Enter support
                self.console.print("[bold yellow]Your question:[/bold yellow]")
                user_question = prompt("", multiline=True).strip()
                
                if not user_question:
                    continue
                
                question_count += 1
                
                # Build RAG query from the question
                rag_query = f"PDL {user_question}"
                
                # Retrieve relevant knowledge with spinner
                with self.console.status("[bold cyan]Just a sec...", spinner="dots") as status:
                    relevant_knowledge = self.get_relevant_knowledge(rag_query, top_k=10)
                
                # Build conversation history for context
                conversation_history_text = ""
                if conversation_context:
                    conversation_history_text = "\n\nPREVIOUS CONVERSATION IN THIS SESSION:\n"
                    for i, conv in enumerate(conversation_context[-3:], 1):  # Last 3 exchanges for context
                        conversation_history_text += f"\nQ{i}: {conv['question']}\nA{i}: {conv['answer'][:300]}...\n"
                
                # Build consultation prompt with conversation history
                consultation_prompt = f"""
                {conversation_history_text}
                
                CURRENT USER QUESTION: {user_question}
                
                RELEVANT PDL DOCUMENTATION AND EXAMPLES:
                {relevant_knowledge}
                
                Please answer the user's current question based on the documentation above.
                If the question refers to previous conversation (e.g., "what about...", "can you explain more...", "how does that work..."),
                use the conversation history context to provide a coherent response.
                
                Provide clear explanations and code examples when appropriate.
                Use ONLY commands and syntax shown in the documentation.
                
                IMPORTANT: 
                - When citing sources, use the FULL FILE PATH provided in the [Full Path: ...] headers above
                - Format your response using proper Markdown syntax
                - Include file references so users can locate the source files
                - Maintain context from previous questions in this session when relevant
                """
                
                # Get answer from PDL expert agent with spinner
                with self.console.status("[bold cyan]Just a sec...", spinner="dots") as status:
                    result = asyncio.run(self.pdl_expert_agent.run(consultation_prompt))
                
                # Display answer with markdown formatting
                self.console.rule("[bold green]Answer[/bold green]", style="green")
                
                # Render as markdown
                md = Markdown(result.data)
                self.console.print(md)
                self.console.print()
                
                # Save to conversation context (for maintaining conversation memory)
                conversation_context.append({
                    "question": user_question,
                    "answer": result.data
                })
                
                # Save to session history (for final summary)
                session_history.append({
                    "question": user_question,
                    "answer": result.data,
                    "tokens": result.usage().total_tokens
                })
                
            except KeyboardInterrupt:
                self.console.print("\n\n[cyan]Session interrupted. Goodbye![/cyan]")
                break
            except Exception as e:
                self.console.print(f"\n[red]Error: {e}[/red]\n")
                import traceback
                traceback.print_exc()
        
        # Display session summary
        if session_history:
            self.console.rule("[bold blue]Session Summary[/bold blue]", style="blue")
            total_tokens = sum(item['tokens'] for item in session_history)
            self.console.print(f"""
[bold green]Summary:[/bold green]
  • Questions asked: {question_count}
  • Total tokens used: {total_tokens}
            """)

    def run_agent_flow(self, context, **kwargs):
        """
        Main entry point - routes to appropriate mode based on user selection.
        """
        # Get mode from initializer_args or prompt if not set
        mode = self.initializer_args.get('mode')
        
        # If no mode in initializer_args, this means get_initializer_args wasn't called yet
        # This can happen in some agent framework flows
        if not mode:
            # Call get_initializer_args to get user input
            init_args = self.get_initializer_args()
            self.initializer_args.update(init_args)
            mode = self.initializer_args.get('mode', '1')
        
        if mode == '2':
            # PDL Expert consultation mode
            return self.run_pdl_expert_mode(context, **kwargs)
        else:
            # SPF to PDL conversion mode (default)
            return self.run_spf_conversion_mode(context, **kwargs)
    
    def run_spf_conversion_mode(self, context, **kwargs):
        """
        SPF to PDL conversion workflow.
        """
        # Get input_path from initializer_args
        input_path = self.initializer_args.get('input_path')
        
        # Also check kwargs in case it's passed differently
        if not input_path:
            input_path = kwargs.get('input_path')
        
        if not input_path:
            self.console.print("[red]Error: input_path is required for conversion mode[/red]")
            self.console.print(f"[dim]Debug: initializer_args = {self.initializer_args}[/dim]")
            self.console.print(f"[dim]Debug: kwargs = {kwargs}[/dim]")
            return
        
        # Display header
        self.console.rule("[bold blue]SPF to PDL Converter Agent[/bold blue]", style="blue")
        self.console.print(f"\n[bold cyan]Input Path:[/bold cyan] {input_path}\n")
        
        # Stage 1: Find SPF files
        self.console.print("[bold yellow]STAGE 1: Finding SPF Files[/bold yellow]", justify="center")
        spf_files = self.find_spf_files(input_path)
        
        if not spf_files:
            self.console.print("\n[bold red]No SPF files found![/bold red]")
            return
        
        self.console.print(f"\n[bold green]Found {len(spf_files)} SPF file(s) to convert[/bold green]\n")
        
        # Stage 2: Convert each file
        self.console.print("[bold yellow]STAGE 2: Converting SPF to PDL (Two-Stage Approach)[/bold yellow]", justify="center")
        self.console.print("[dim]Stage 2a: Understanding SPF intent | Stage 2b: Generating PDL[/dim]\n", justify="center")
        
        conversion_results = []
        total_tokens = 0
        
        for idx, spf_file in enumerate(spf_files, 1):
            self.console.print(f"\n[bold white]({idx}/{len(spf_files)})[/bold white] Processing: [cyan]{Path(spf_file).name}[/cyan]")
            
            # Read SPF file
            spf_data = self.read_spf_file(spf_file)
            if not spf_data:
                continue
            
            # Convert to PDL
            result = self.convert_spf_to_pdl(spf_data)
            conversion_results.append(result)
            
            if result.get('success'):
                total_tokens += result.get('tokens_used', 0)
                self.console.print(f"  [green]✓ Conversion successful[/green]")
            else:
                self.console.print(f"  [red]✗ Conversion failed: {result.get('error')}[/red]")
        
        # Stage 3: Save results
        self.console.print(f"\n[bold yellow]STAGE 3: Saving PDL Files[/bold yellow]", justify="center")
        
        saved_files = []
        failed_files = []
        
        for result in conversion_results:
            if result.get('success'):
                pdl_path = self.save_pdl_file(result)
                if pdl_path:
                    saved_files.append(pdl_path)
                    self.console.print(f"  [green]✓[/green] Saved: {Path(pdl_path).name}")
                else:
                    failed_files.append(Path(result['spf_file']).name)
            else:
                failed_files.append(Path(result['spf_file']).name)
        
        # Stage 4: Generate summary report
        self.console.print(f"\n[bold yellow]STAGE 4: Generating Summary[/bold yellow]", justify="center")
        
        summary = {
            "total_files_processed": len(spf_files),
            "successful_conversions": len(saved_files),
            "failed_conversions": len(failed_files),
            "total_ai_tokens_used": total_tokens,
            "converted_files": saved_files,
            "failed_files": failed_files,
            "conversion_results": conversion_results
        }
        
        # Save summary JSON
        summary_file = Path.cwd() / "spf_to_pdl_conversion_summary.json"
        
        try:
            with open(summary_file, 'w', encoding='utf-8') as f:
                json.dump(summary, f, indent=2)
            self.console.print(f"\n[green]Summary saved to:[/green] [bold]{summary_file}[/bold]")
        except Exception as e:
            self.console.print(f"\n[red]Error saving summary: {e}[/red]")
        
        # Final summary display
        self.console.rule("[bold blue]Conversion Complete[/bold blue]", style="blue")
        self.console.print(f"""
[bold green]Summary:[/bold green]
  • Total SPF files: {len(spf_files)}
  • Successfully converted: {len(saved_files)}
  • Failed: {len(failed_files)}
  • AI Tokens used: {total_tokens}
        """, justify="center")
        
        if saved_files:
            self.console.print("\n[bold green]Converted PDL Files:[/bold green]")
            for pdl_file in saved_files:
                self.console.print(f"  • {pdl_file}")
        
        if failed_files:
            self.console.print("\n[bold red]Failed Conversions:[/bold red]")
            for failed_file in failed_files:
                self.console.print(f"  • {failed_file}")

class RAGKnowledgeBase:
    """
    RAG-based knowledge base that uses vector embeddings and semantic search
    to retrieve relevant information from PDF documents.
    """
    
    def __init__(
        self,
        pdf_dir: Path,
        cache_dir: Optional[Path] = None,
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        collection_name: str = "tessent_knowledge"
    ):
        """
        Initialize RAG Knowledge Base.
        
        Args:
            pdf_dir: Directory containing PDF files
            cache_dir: Directory for caching embeddings (defaults to pdf_dir/.rag_cache)
            embedding_model: HuggingFace model name for embeddings
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks for context preservation
            collection_name: Name of the ChromaDB collection
        """
        self.console = Console()
        self.pdf_dir = Path(pdf_dir)
        self.cache_dir = cache_dir or (self.pdf_dir / ".rag_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.embedding_model_name = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.collection_name = collection_name
        
        # Initialize components
        self.embedding_model = None
        self.chroma_client = None
        self.collection = None
        
        # Check dependencies
        self._check_dependencies()
        
        # Initialize models and database
        self._initialize_components()
    
    def _check_dependencies(self):
        """Check if required dependencies are installed."""
        missing = []
        
        if PyPDF2 is None:
            missing.append("PyPDF2")
        if SentenceTransformer is None:
            missing.append("sentence-transformers")
        if chromadb is None:
            missing.append("chromadb")
        
        if missing:
            error_msg = f"Missing required packages: {', '.join(missing)}"
            self.console.print(f"[bold red]Error: {error_msg}[/bold red]")
            self.console.print("\n[yellow]Install with:[/yellow]")
            self.console.print(f"  pip install {' '.join(missing)}")
            raise ImportError(error_msg)
    
    def _initialize_components(self):
        """Initialize embedding model and vector database (load existing or create new)."""
        try:
            # Initialize ChromaDB first to check if we have existing data
            chroma_path = str(self.cache_dir / "chroma_db")
            self.chroma_client = chromadb.PersistentClient(path=chroma_path)
            
            # Get or create collection (same as valbot_tui.py)
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            # Check if we have existing embeddings
            existing_count = self.collection.count()
            is_new = existing_count == 0
            
            if is_new:
                self.console.print("[dim]Initializing new knowledge base...[/dim]")
            else:
                self.console.print(f"[dim]Found existing knowledge base with {existing_count} chunks[/dim]")
            
            # Load embedding model
            if is_new:
                self.console.print("[dim]Loading embedding model (this may take a moment)...[/dim]")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            if is_new:
                self.console.print("[green]✓ Embedding model loaded[/green]")
            
        except Exception as e:
            self.console.print(f"[red]Error initializing components: {e}[/red]")
            raise
    
    def _get_pdf_hash(self, pdf_path: Path) -> str:
        """Generate hash of PDF file for caching."""
        hash_obj = hashlib.md5()
        with open(pdf_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    
    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """
        Extract all text from a PDF file.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {pdf_path.name}: {e}[/red]")
            return ""
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """
        Extract all text from a .docx file.
        
        Args:
            docx_path: Path to .docx file
            
        Returns:
            Extracted text content
        """
        if Document is None:
            self.console.print(f"[yellow]⚠ python-docx not installed. Install with: pip install python-docx[/yellow]")
            return ""
        
        try:
            doc = Document(str(docx_path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {docx_path.name}: {e}[/red]")
            return ""
    
    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks for better context preservation.
        
        Args:
            text: Full text to chunk
            metadata: Metadata to attach to each chunk
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            # Extract chunk
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                last_period = chunk_text.rfind('. ')
                last_newline = chunk_text.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > self.chunk_size * 0.5:  # Don't break too early
                    chunk_text = chunk_text[:break_point + 1]
                    end = start + break_point + 1
            
            # Create chunk with metadata
            chunk = {
                'text': chunk_text.strip(),
                'metadata': {
                    **metadata,
                    'chunk_id': chunk_id,
                    'start_char': start,
                    'end_char': end
                }
            }
            
            if chunk['text']:  # Only add non-empty chunks
                chunks.append(chunk)
                chunk_id += 1
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
        
        return chunks
    
    def load_text_file(self, file_path: Path, source_name: str = None, force_reload: bool = False):
        """
        Removed - database loading is disabled.
        """
        pass
    
    def load_pdfs(self, pdf_files: Optional[List[str]] = None, force_reload: bool = False):
        """
        Removed - database loading is disabled.
        """
        pass
    
    def load_docx_files(self, docx_files: Optional[List[str]] = None, force_reload: bool = False):
        """
        Removed - database loading is disabled.
        """
        pass
    
    def retrieve_relevant_context(
        self, 
        query: str, 
        top_k: int = 5,
        min_similarity: float = 0.0
    ) -> List[Dict[str, Any]]:
        """
        Retrieve the most relevant text chunks for a given query.
        
        Args:
            query: Search query
            top_k: Number of top results to return
            min_similarity: Minimum similarity score (0-1)
            
        Returns:
            List of relevant chunks with text, metadata, and similarity scores
        """
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=top_k
            )
            
            # Format results
            relevant_chunks = []
            if results['documents'] and results['documents'][0]:
                for i, (doc, metadata, distance) in enumerate(zip(
                    results['documents'][0],
                    results['metadatas'][0],
                    results['distances'][0]
                )):
                    # Convert distance to similarity (cosine distance -> similarity)
                    similarity = 1 - distance
                    
                    if similarity >= min_similarity:
                        relevant_chunks.append({
                            'text': doc,
                            'metadata': metadata,
                            'similarity': similarity,
                            'rank': i + 1
                        })
            
            return relevant_chunks
        
        except Exception as e:
            self.console.print(f"[red]Error retrieving context: {e}[/red]")
            return []
    
    def get_context_for_query(
        self, 
        query: str, 
        top_k: int = 5,
        include_metadata: bool = False
    ) -> str:
        """
        Get formatted context string for a query.
        
        Args:
            query: Search query
            top_k: Number of chunks to retrieve
            include_metadata: Include source metadata in output
            
        Returns:
            Formatted context string
        """
        chunks = self.retrieve_relevant_context(query, top_k=top_k)
        
        if not chunks:
            return "No relevant context found in knowledge base."
        
        context_parts = []
        for chunk in chunks:
            source_name = chunk['metadata']['source']
            # Construct full file path for reference
            full_path = self.pdf_dir / source_name
            
            if include_metadata:
                header = f"[Source: {source_name} | Full Path: {full_path} | Similarity: {chunk['similarity']:.2f}]"
                context_parts.append(f"{header}\n{chunk['text']}\n")
            else:
                # Include file path reference even when include_metadata is False
                # This helps the agent provide proper file references to users
                header = f"[Source: {source_name} | Full Path: {full_path}]"
                context_parts.append(f"{header}\n{chunk['text']}\n")
        
        return "\n\n---\n\n".join(context_parts)
    
    def search_knowledge_base(self, query: str, top_k: int = 5) -> None:
        """
        Interactive search and display of knowledge base results.
        
        Args:
            query: Search query
            top_k: Number of results to display
        """
        self.console.print(f"\n[bold cyan]Searching for:[/bold cyan] '{query}'")
        
        chunks = self.retrieve_relevant_context(query, top_k=top_k)
        
        if not chunks:
            self.console.print("[yellow]No relevant results found[/yellow]")
            return
        
        self.console.print(f"\n[green]Found {len(chunks)} relevant chunks:[/green]\n")
        
        for i, chunk in enumerate(chunks, 1):
            self.console.rule(f"[bold]Result {i}[/bold]", style="blue")
            self.console.print(f"[dim]Source:[/dim] {chunk['metadata']['source']}")
            self.console.print(f"[dim]Similarity:[/dim] {chunk['similarity']:.3f}")
            self.console.print(f"\n{chunk['text'][:500]}...")
            self.console.print()
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base."""
        count = 0
        if self.collection:
            try:
                count = self.collection.count()
            except:
                count = 0
        
        return {
            "total_chunks": count,
            "embedding_model": self.embedding_model_name,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "cache_dir": str(self.cache_dir)
        }
