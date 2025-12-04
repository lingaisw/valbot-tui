# -*- coding: utf-8 -*-
"""
ValBot TUI - A Modern Material Design Terminal Interface for ValBot
Inspired by Charm's Crush interface with beautiful gradients and sleek styling

FEATURE PARITY WITH CLI:
=======================
This TUI implementation includes all major features from the CLI version:

✅ System Prompt Integration:
   - Automatically loads system_prompt from config on startup
   - Ensures consistent AI behavior according to user preferences

✅ GPT-5 Reasoning Support:
   - Displays thinking process when display_reasoning is enabled
   - Configurable reasoning effort levels (low/medium/high)
   - Streams reasoning text in real-time via ResponseAudioDeltaEvent
   - Shows progress indicators during reasoning phase

✅ Full Command Support:
   - /new - Start new conversation
   - /quit - Exit application
   - /help - Show comprehensive help
   - /model - Interactive model picker with arrow keys
   - /agent - Interactive agent selection with descriptions
   - /file - Display file contents with syntax highlighting
   - /terminal - Execute shell commands
   - /multi - Multi-line input via system editor
   - /prompts - Show custom prompts (via CommandManager)
   - /commands - Show all available commands (via CommandManager)
   - /settings - Settings information
   - /reload - Reinitialize chatbot
   - /update - Check for and install updates from GitHub
   - /add_agent - Add agent information
   - /add_tool - Add tool information

✅ CommandManager Integration:
   - Uses actual CommandManager from bot_commands.py
   - Supports custom prompts with argument parsing
   - Supports custom commands from agent plugins
   - Automatic delegation to plugin manager

✅ Agent System:
   - Interactive agent picker with arrow navigation
   - Shows agent names and descriptions
   - Executes agent workflows with full context
   - Error handling and status reporting

✅ Context Management:
   - Load single files or glob patterns
   - Integrates with ContextManager
   - Visual feedback on loaded files
   - File content preview with syntax highlighting

✅ Response Streaming:
   - Real-time text streaming
   - Reasoning display for GPT-5 models
   - Proper event handling (ResponseAudioDeltaEvent, ResponseTextDeltaEvent)
   - Visual progress indicators

✅ Markdown Rendering:
   - Full markdown support with syntax highlighting
   - Code blocks with copy buttons
   - Tables, lists, quotes, emphasis
   - Collapsible code sections

✅ Material Design UI:
   - Modern dark theme
   - Gradient accents
   - Smooth animations
   - Responsive layout
   - Keyboard shortcuts

✅ Theme Persistence:
   - Theme changes are automatically saved to ~/.valbot_tui_config.json
   - Saved theme is loaded on startup
   - Use Ctrl+P (command palette) to change themes
   - Supports all Textual built-in themes plus custom ValBot theme
"""

import asyncio
import sys
import re
import glob
import gzip
import json
import shlex
import traceback
from datetime import datetime
from typing import Optional, List, Dict, Any, Iterable, Tuple
from pathlib import Path

from textual import on, work, events
from textual.app import App, ComposeResult
from textual.binding import Binding
from textual.message import Message
from textual.containers import Container, Horizontal, Vertical, VerticalScroll, ScrollableContainer
from textual.reactive import reactive
from textual.screen import Screen
from textual.theme import Theme
from textual.widgets import (
    Header, Input, Static, Button, 
    DirectoryTree, Label, Markdown, RichLog, TabbedContent, TabPane, TextArea, OptionList, ProgressBar
)
from textual.widgets.option_list import Option
from textual.widgets._directory_tree import DirEntry
from textual.widgets._tree import TreeNode, TOGGLE_STYLE
from textual.worker import Worker, WorkerState
from rich.markdown import Markdown as RichMarkdown
from rich.text import Text
from rich.style import Style
from rich.panel import Panel
from rich.console import Group
from rich.syntax import Syntax

from openai.types.responses import (
    ResponseCreatedEvent,
    ResponseInProgressEvent,
    ResponseAudioDeltaEvent,
    ResponseTextDeltaEvent
)


def open_file_smart(file_path, mode='r', **kwargs):
    """
    Open regular or .gz compressed files transparently.
    
    Args:
        file_path: Path to file (can be .gz compressed or Path object)
        mode: File open mode ('r', 'rb', etc.)
        **kwargs: Additional arguments passed to open/gzip.open
        
    Returns:
        File handle for reading
    """
    file_str = str(file_path)  # Convert Path objects to string
    if file_str.endswith('.gz'):
        # For text mode, ensure we use text mode with gzip
        if 'b' not in mode:
            return gzip.open(file_str, mode + 't', **kwargs)
        return gzip.open(file_str, mode, **kwargs)
    return open(file_str, mode, **kwargs)

from chatbot import ChatBot
from config import ConfigManager
from client_setup import initialize_agent_model, initialize_chat_client
from agent_plugins.plugin_manager import PluginManager
from agent_plugins import utilities
from terminal_manager import TerminalManager
from bot_commands import CommandManager
from rich.console import Console
from console_extensions import HistoryConsole
from valbot_updater import ValbotUpdater
import sys
import os
import tempfile
import subprocess
import hashlib
import traceback

# RAG database imports - optional dependencies
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    chromadb = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import pandas as pd
    import openpyxl
except ImportError:
    pd = None
    openpyxl = None

try:
    from chonkie import TokenChunker
    CHONKIE_AVAILABLE = True
except ImportError:
    CHONKIE_AVAILABLE = False




# Platform-specific emoji/symbol mapping
# print(sys.platform)
IS_LINUX = sys.platform.startswith('linux')

# Define emoji mappings based on platform
if IS_LINUX:
    # ASCII symbols for Linux
    EMOJI = {
        'checkmark': '✓',
        'clipboard': '\\[>_]',
        'info': 'ℹ️',
        'lightbulb': '◐',
        'robot': '⚞⚟',
        'user': '❰❱',
        'cross': '✕',
        'gear': '⚙️',
        'open_folder': '◧',
        'folder': '◨',
        'file': '◫',
        'keyboard': '⌨️ ',
        'lightning': '⚡',
        'agent': '◈',
        'chat' : '◷',
        'link': '⎘',
        'send': '⌲',
        'stop': '◼',
        'database': '⛁',
        'hourglass': '⌛',
        'refresh': '⟳',
    }
else:
    # Unicode emojis for Windows/Mac
    EMOJI = {
        'checkmark': '✅',
        'clipboard': '📋',
        'info': 'ℹ️',
        'lightbulb': '💡',
        'robot': '🤖',
        'user': '👤',
        'cross': '❌',
        'gear': '⚙️',
        'open_folder': '📂',
        'folder': '📁',
        'file': '📄',
        'keyboard': '⌨️',
        'lightning': '⚡',
        'agent': '🤖',
        'chat': '💬',
        'link': '🔗',
        'send': '⌲',
        'stop': '◼',
        'database': '📊',
        'hourglass': '⌛',
        'refresh': '🔄',
    }


# Register custom ValBot dark theme
VALBOT_DARK_THEME = Theme(
    name="valbot-dark (default)",
    primary="#989af5",           # Indigo - primary brand color
    secondary="#4d4eba",         # Indigo darker - secondary/user color
    accent="#0696d4",            # Light blue - accent highlights
    foreground="#fafafa",        # Alabaster white - foreground & text 
    warning="#f59e0b",           # Amber - warnings
    error="#ef4444",             # Red - errors
    success="#10b981",           # Emerald - success states
    background="#0f172a",        # Slate-950 - main background
    surface="#1e293b",           # Slate-800 - surface/panel background
    panel="#353e4f",             # Slate-700 - elevated panels
    dark=True,
    # Additional variables for text colors
    variables={
        "text": "#fafafa",       # Slate-200 - primary text
        "text-muted": "#94a3b8",  # Slate-400 - muted text
        "text-bright": "#ffffff", # White - bright text
        "text-dim": "#a5b4fc",    # Indigo-300 - dimmed text
        "gray-light": "#374151",  # Gray-700 - light gray surfaces
        "blue-dark": "#1d4ed8",   # Blue-700 - dark blue
        "blue": "#2563eb",        # Blue-600 - standard blue
        "indigo-dark": "#4f46e5", # Indigo-600 - dark indigo
    }
)

# Theme persistence configuration file path
THEME_CONFIG_PATH = Path.home() / ".valbot_tui_config.json"


def save_theme_config(theme_name: str) -> None:
    """Save the current theme to the config file."""
    try:
        # Load existing config to preserve other fields
        config = {}
        if THEME_CONFIG_PATH.exists():
            with open(THEME_CONFIG_PATH, 'r') as f:
                config = json.load(f)
        
        # Update theme
        config["theme"] = theme_name
        
        # Write back
        with open(THEME_CONFIG_PATH, 'w') as f:
            json.dump(config, f, indent=2)
    except Exception as e:
        # Silently fail if we can't save the theme config
        pass


def load_theme_config() -> Optional[str]:
    """Load the saved theme from the config file."""
    try:
        if THEME_CONFIG_PATH.exists():
            with open(THEME_CONFIG_PATH, 'r') as f:
                config = json.load(f)
                return config.get("theme")
    except Exception as e:
        # Silently fail if we can't load the theme config
        pass
    return None


def save_history_config(location: str, max_size_gb: int) -> None:
    """Save history configuration to the config file."""
    try:
        # Load existing config to preserve other fields
        config = {}
        if THEME_CONFIG_PATH.exists():
            with open(THEME_CONFIG_PATH, 'r') as f:
                config = json.load(f)
        
        # Update history config
        config["history_config"] = {
            "location": location,
            "max_size": max_size_gb
        }
        
        # Write back
        with open(THEME_CONFIG_PATH, 'w') as f:
            json.dump(config, f, indent=2)
    except Exception as e:
        # Silently fail if we can't save the history config
        pass


def load_history_config() -> Optional[Dict[str, Any]]:
    """Load history configuration from the config file."""
    try:
        if THEME_CONFIG_PATH.exists():
            with open(THEME_CONFIG_PATH, 'r') as f:
                config = json.load(f)
                return config.get("history_config")
    except Exception as e:
        # Silently fail if we can't load the history config
        pass
    return None


def ensure_history_directory() -> Path:
    """
    Ensure the chat history directory exists and return its path.
    Creates valbot_tui_chat_history folder at the configured location.
    """
    history_config = load_history_config()
    
    if history_config:
        base_location = Path(history_config.get("location", str(Path.cwd())))
    else:
        base_location = Path.cwd()
    
    # Create the main directory and the chat_history subdirectory
    history_dir = base_location / "valbot_tui_chat_history"
    history_dir.mkdir(parents=True, exist_ok=True)
    
    return history_dir


def get_directory_size(directory: Path) -> int:
    """
    Calculate the total size of all files in a directory in bytes.
    
    Args:
        directory: Path to the directory
        
    Returns:
        Total size in bytes
    """
    total_size = 0
    try:
        for file_path in directory.rglob('*'):
            if file_path.is_file():
                total_size += file_path.stat().st_size
    except Exception:
        pass
    return total_size


def cleanup_old_history(history_dir: Path, max_size_gb: int) -> None:
    """
    Check if chat history folder exceeds max_size and delete oldest files until it fits.
    
    Args:
        history_dir: Path to the chat history directory
        max_size_gb: Maximum size in gigabytes
    """
    try:
        # Convert GB to bytes
        max_size_bytes = max_size_gb * 1024 * 1024 * 1024
        
        # Get current directory size
        current_size = get_directory_size(history_dir)
        
        # If under the limit, nothing to do
        if current_size <= max_size_bytes:
            return
        
        # Get all .log files sorted by modification time (oldest first)
        log_files = []
        for file_path in history_dir.glob('*.log'):
            if file_path.is_file():
                try:
                    mtime = file_path.stat().st_mtime
                    size = file_path.stat().st_size
                    log_files.append((file_path, mtime, size))
                except Exception:
                    continue
        
        # Sort by modification time (oldest first)
        log_files.sort(key=lambda x: x[1])
        
        # Delete oldest files until we're under the limit
        for file_path, mtime, size in log_files:
            if current_size <= max_size_bytes:
                break
            
            try:
                file_path.unlink()
                current_size -= size
            except Exception:
                # If we can't delete a file, continue with the next one
                continue
                
    except Exception:
        # Silently fail if we encounter any errors during cleanup
        pass


# ============================================================================
# RAG Knowledge Base Class
# ============================================================================

class RAGKnowledgeBase:
    """
    RAG-based knowledge base that uses vector embeddings and semantic search
    to retrieve relevant information from documents.
    """
    
    def __init__(
        self,
        pdf_dir: Path,
        cache_dir: Optional[Path] = None,
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        collection_name: str = "rag_knowledge_base"
    ):
        """
        Initialize RAG Knowledge Base.
        
        Args:
            pdf_dir: Directory containing documents
            cache_dir: Directory for caching embeddings
            embedding_model: HuggingFace model name for embeddings
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks for context preservation
            collection_name: Name of the ChromaDB collection
        """
        self.console = Console()
        self.pdf_dir = Path(pdf_dir)
        self.cache_dir = cache_dir or (self.pdf_dir / ".rag_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.embedding_model_name = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.collection_name = collection_name
        
        # Initialize components
        self.embedding_model = None
        self.chroma_client = None
        self.collection = None
        
        # Check dependencies
        self._check_dependencies()
        
        # Initialize models and database
        self._initialize_components()
    
    def _check_dependencies(self):
        """Check if required dependencies are installed."""
        missing = []
        
        if PyPDF2 is None:
            missing.append("PyPDF2")
        if SentenceTransformer is None:
            missing.append("sentence-transformers")
        if chromadb is None:
            missing.append("chromadb")
        
        if missing:
            error_msg = f"Missing required packages: {', '.join(missing)}"
            self.console.print(f"[bold red]Error: {error_msg}[/bold red]")
            self.console.print("\n[yellow]Install with:[/yellow]")
            self.console.print(f"  pip install {' '.join(missing)}")
            raise ImportError(error_msg)
    
    def _initialize_components(self):
        """Initialize embedding model and vector database."""
        try:
            # Initialize ChromaDB first to check if we have existing data
            chroma_path = str(self.cache_dir / "chroma_db")
            self.chroma_client = chromadb.PersistentClient(path=chroma_path)
            
            # Get or create collection
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            # Check if we have existing embeddings
            existing_count = self.collection.count()
            is_new = existing_count == 0
            
            if is_new:
                self.console.print("[dim]Initializing new knowledge base...[/dim]")
            else:
                self.console.print(f"[dim]Found existing knowledge base with {existing_count} chunks[/dim]")
            
            # Load embedding model
            if is_new:
                self.console.print("[dim]Loading embedding model (this may take a moment)...[/dim]")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            if is_new:
                self.console.print("[green]✓ Embedding model loaded[/green]")
            
        except Exception as e:
            self.console.print(f"[red]Error initializing components: {e}[/red]")
            raise
    
    def _get_pdf_hash(self, file_path: Path) -> str:
        """Generate hash of file for caching."""
        hash_obj = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    
    def _validate_file_type(self, file_path: Path, allowed_extensions: List[str], method_name: str = "this method") -> bool:
        """Validate file type and provide helpful error messages for unsupported files.
        
        Args:
            file_path: Path to the file to validate
            allowed_extensions: List of allowed file extensions (e.g., ['.pdf', '.txt'])
            method_name: Name of the calling method for error messages
            
        Returns:
            True if file is valid, False otherwise
        """
        file_ext = file_path.suffix.lower()
        
        # If file has no extension and path is valid, treat as normal text file
        if not file_ext and file_path.is_file():
            return True
        
        # Check if file extension is in allowed list
        if file_ext not in allowed_extensions:
            self.console.print(f"[bold red]Error:[/bold red] File '{file_path.name}' is not supported by {method_name}.")
            
            # Provide specific hints for common unsupported types
            if file_ext in ['.xlsx', '.xls', '.xlsm', '.xlsb']:
                self.console.print(f"[yellow]Excel files are not supported for RAG database loading.[/yellow]")
            elif file_ext in ['.exe', '.dll', '.bin', '.so', '.dylib']:
                self.console.print(f"[yellow]Binary executable files cannot be loaded.[/yellow]")
            elif file_ext in ['.zip', '.rar', '.7z', '.tar']:
                self.console.print(f"[yellow]Archive files (except .gz) are not supported.[/yellow]")
            elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg']:
                self.console.print(f"[yellow]Image files cannot be loaded as text.[/yellow]")
            elif file_ext in ['.mp3', '.mp4', '.avi', '.mov', '.wav']:
                self.console.print(f"[yellow]Media files are not supported.[/yellow]")
            
            # Show allowed formats
            allowed_str = ", ".join(allowed_extensions)
            self.console.print(f"[yellow]Allowed formats:[/yellow] {allowed_str}")
            return False
        
        return True
    
    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract all text from a PDF file (supports .pdf and .pdf.gz)."""
        try:
            # Check if it's a .gz compressed PDF
            if str(pdf_path).endswith('.pdf.gz'):
                import tempfile
                # Decompress to temporary file
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                    tmp_path = tmp_file.name
                    with gzip.open(str(pdf_path), 'rb') as gz_file:
                        tmp_file.write(gz_file.read())
                try:
                    with open(tmp_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        return text
                finally:
                    import os
                    os.unlink(tmp_path)
            else:
                # Regular PDF file
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {pdf_path.name}: {e}[/red]")
            return ""
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract all text from a .docx file (supports .docx and .docx.gz)."""
        if DocxDocument is None:
            self.console.print(f"[yellow]⚠ python-docx not installed. Install with: pip install python-docx[/yellow]")
            return ""
        
        try:
            # Check if it's a .gz compressed DOCX
            if str(docx_path).endswith('.docx.gz'):
                import tempfile
                # Decompress to temporary file
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_path = tmp_file.name
                    with gzip.open(str(docx_path), 'rb') as gz_file:
                        tmp_file.write(gz_file.read())
                try:
                    doc = DocxDocument(tmp_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                        text += "\n"
                    return text
                finally:
                    import os
                    os.unlink(tmp_path)
            else:
                # Regular DOCX file
                doc = DocxDocument(str(docx_path))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
                return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {docx_path.name}: {e}[/red]")
            return ""
    
    def chunk_text(self, text: str, metadata: dict) -> List[dict]:
        """Split text into overlapping chunks for better context preservation."""
        chunks = []
        
        # Debug: Check text length
        if not text or not text.strip():
            return chunks
        
        text_length = len(text)
        
        # Use Chonkie if available for better semantic chunking
        if CHONKIE_AVAILABLE:
            try:
                # Use TokenChunker from Chonkie for semantic-aware chunking
                chunker = TokenChunker(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap
                )
                chonkie_chunks = chunker.chunk(text)
                
                for chunk_id, chunk_obj in enumerate(chonkie_chunks):
                    chunk = {
                        'text': chunk_obj.text.strip(),
                        'metadata': {
                            **metadata,
                            'chunk_id': chunk_id,
                            'start_char': chunk_obj.start_index if hasattr(chunk_obj, 'start_index') else 0,
                            'end_char': chunk_obj.end_index if hasattr(chunk_obj, 'end_index') else len(chunk_obj.text)
                        }
                    }
                    if chunk['text']:
                        chunks.append(chunk)
                
                return chunks
            except Exception as e:
                # Fall back to basic chunking if Chonkie fails
                self.console.print(f"[dim]Chonkie chunking failed, using basic chunking: {e}[/dim]")
        
        # Basic chunking (fallback)
        start = 0
        chunk_id = 0
        
        while start < len(text):
            # Extract chunk
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                last_period = chunk_text.rfind('. ')
                last_newline = chunk_text.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > self.chunk_size * 0.5:  # Don't break too early
                    chunk_text = chunk_text[:break_point + 1]
                    end = start + break_point + 1
            
            # Create chunk with metadata
            chunk = {
                'text': chunk_text.strip(),
                'metadata': {
                    **metadata,
                    'chunk_id': chunk_id,
                    'start_char': start,
                    'end_char': end
                }
            }
            
            if chunk['text']:  # Only add non-empty chunks
                chunks.append(chunk)
                chunk_id += 1
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
        
        return chunks
    
    def load_text_file(self, file_path: Path, source_name: str = None, force_reload: bool = False):
        """Load a text file into the vector database."""
        if not file_path.exists():
            self.console.print(f"[yellow]⚠[/yellow] File not found: {file_path}")
            return
        
        # Validate file type - allow common text file extensions
        allowed_extensions = ['.txt', '.md', '.py', '.js', '.java', '.c', '.cpp', '.h', '.hpp', '.cs', '.go', 
                             '.rs', '.rb', '.php', '.json', '.xml', '.yaml', '.yml', '.csv', '.log', 
                             '.sh', '.bash', '.sql', '.html', '.css', '.scss', '.ts', '.jsx', '.tsx']
        if not self._validate_file_type(file_path, allowed_extensions, "text file loading"):
            return
        
        source_name = source_name or file_path.name
        
        # Check if already processed
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
            
            # Check if file is already loaded
            file_hash = self._get_pdf_hash(file_path)
            if source_name in processed_docs and processed_docs[source_name] == file_hash:
                self.console.print(f"[dim]Skipping {source_name} (already processed)[/dim]")
                return
        
        # Read text file
        try:
            with open_file_smart(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except UnicodeDecodeError:
            self.console.print(f"[bold red]Error:[/bold red] Unable to read '{file_path.name}' as text. This appears to be a binary file.")
            self.console.print(f"[yellow]Supported formats:[/yellow] Text-based files only (.txt, .md, .py, .json, etc., also .gz compressed)")
            return
        except Exception as e:
            self.console.print(f"[red]Error reading {file_path.name}: {e}[/red]")
            return
        
        if not text:
            return
        
        # Create chunks
        file_hash = self._get_pdf_hash(file_path)
        metadata = {
            'source': source_name,
            'doc_hash': file_hash,
            'file_type': 'text'
        }
        chunks = self.chunk_text(text, metadata)
        
        if chunks:
            # Generate embeddings and store
            texts = [chunk['text'] for chunk in chunks]
            embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
            
            # Prepare for ChromaDB
            ids = [f"{source_name}_{chunk['metadata']['chunk_id']}" for chunk in chunks]
            metadatas = [chunk['metadata'] for chunk in chunks]
            
            # Add to collection
            self.collection.add(
                ids=ids,
                embeddings=embeddings.tolist(),
                documents=texts,
                metadatas=metadatas
            )
            
            # Update processed docs cache
            processed_docs[source_name] = file_hash
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
    
    def load_pdfs(self, pdf_files: Optional[List[str]] = None, force_reload: bool = False):
        """Load and process PDF files into the vector database."""
        from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
        
        # Check if knowledge base is already fully loaded
        if not force_reload:
            existing_count = self.collection.count()
            cache_file = self.cache_dir / "processed_docs.json"
            
            if existing_count > 0 and cache_file.exists():
                self.console.print(f"[dim]Knowledge base already loaded ({existing_count} chunks)[/dim]")
        
        # Find PDF files and validate they are PDFs
        if pdf_files is None:
            pdf_paths = list(self.pdf_dir.glob("*.pdf"))
        else:
            # Handle both absolute paths and relative filenames
            pdf_paths = []
            for pdf in pdf_files:
                pdf_path = Path(pdf)
                # If it's already an absolute path, use it directly (don't combine with pdf_dir)
                if pdf_path.is_absolute():
                    pdf_paths.append(pdf_path)
                else:
                    # Otherwise treat it as relative to pdf_dir
                    pdf_paths.append(self.pdf_dir / pdf)
            
            # Validate that all files are PDFs
            for pdf_path in pdf_paths:
                if not self._validate_file_type(pdf_path, ['.pdf'], "PDF loading"):
                    return
        
        if not pdf_paths:
            self.console.print("[yellow]No PDF files found[/yellow]")
            return
        
        self.console.print(f"\n[bold cyan]Loading {len(pdf_paths)} PDF(s) into RAG Knowledge Base[/bold cyan]")
        
        # Track processed documents
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
        
        all_chunks = []
        total_chunks = 0
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=self.console
        ) as progress:
            
            task = progress.add_task("[cyan]Processing PDFs...", total=len(pdf_paths))
            
            for pdf_path in pdf_paths:
                if not pdf_path.exists():
                    self.console.print(f"  [yellow]⚠[/yellow] File not found: {pdf_path.name}")
                    progress.advance(task)
                    continue
                
                # Check if already processed
                pdf_hash = self._get_pdf_hash(pdf_path)
                if pdf_path.name in processed_docs and processed_docs[pdf_path.name] == pdf_hash and not force_reload:
                    self.console.print(f"  [dim]Skipping (cached): {pdf_path.name}[/dim]")
                    progress.advance(task)
                    continue
                
                # Extract text
                progress.update(task, description=f"[cyan]Extracting: {pdf_path.name}")
                text = self.extract_text_from_pdf(pdf_path)
                
                if not text:
                    self.console.print(f"  [yellow]⚠[/yellow] No text extracted from {pdf_path.name}")
                    progress.advance(task)
                    continue
                
                # Create chunks
                progress.update(task, description=f"[cyan]Chunking: {pdf_path.name}")
                metadata = {
                    'source': pdf_path.name,
                    'doc_hash': pdf_hash,
                    'file_type': 'pdf'
                }
                chunks = self.chunk_text(text, metadata)
                
                self.console.print(f"  [green]✓[/green] {pdf_path.name}: {len(chunks)} chunks ({len(text):,} chars)")
                
                all_chunks.extend(chunks)
                total_chunks += len(chunks)
                processed_docs[pdf_path.name] = pdf_hash
                
                progress.advance(task)
        
        # Generate embeddings and store in ChromaDB
        if all_chunks:
            self.console.print(f"\n[yellow]Generating embeddings for {total_chunks} chunks...[/yellow]")
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=self.console
            ) as progress:
                
                task = progress.add_task("[cyan]Creating embeddings...", total=total_chunks)
                
                # Process in batches
                batch_size = 32
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i + batch_size]
                    texts = [chunk['text'] for chunk in batch]
                    embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                    
                    # Prepare for ChromaDB
                    ids = [f"{chunk['metadata']['source']}_{chunk['metadata']['chunk_id']}" for chunk in batch]
                    metadatas = [chunk['metadata'] for chunk in batch]
                    
                    # Add to collection
                    self.collection.add(
                        ids=ids,
                        embeddings=embeddings.tolist(),
                        documents=texts,
                        metadatas=metadatas
                    )
                    
                    progress.update(task, advance=len(batch))
            
            # Save processed documents cache
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
            
            self.console.print(f"[green]✓ Successfully loaded {total_chunks} chunks into vector database[/green]")
        else:
            self.console.print("[yellow]No chunks to process[/yellow]")
    
    def load_docx_files(self, docx_files: Optional[List[str]] = None, force_reload: bool = False):
        """Load and process .docx files into the vector database."""
        from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
        
        if DocxDocument is None:
            self.console.print(f"[yellow]⚠ python-docx not installed. Skipping .docx files. Install with: pip install python-docx[/yellow]")
            return
        
        # Find .docx files and validate they are DOCX
        if docx_files is None:
            docx_paths = list(self.pdf_dir.glob("*.docx"))
        else:
            # Handle both absolute paths and relative filenames
            docx_paths = []
            for docx in docx_files:
                docx_path = Path(docx)
                # If it's already an absolute path that exists, use it directly
                if docx_path.is_absolute() and docx_path.exists():
                    docx_paths.append(docx_path)
                else:
                    # Otherwise treat it as relative to pdf_dir
                    docx_paths.append(self.pdf_dir / docx)
            
            # Validate that all files are DOCX
            for docx_path in docx_paths:
                if not self._validate_file_type(docx_path, ['.docx'], "DOCX loading"):
                    return
        
        if not docx_paths:
            self.console.print("[yellow]No .docx files found[/yellow]")
            return
        
        self.console.print(f"\n[bold cyan]Loading {len(docx_paths)} .docx file(s) into RAG Knowledge Base[/bold cyan]")
        
        # Track processed documents
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
        
        all_chunks = []
        total_chunks = 0
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=self.console
        ) as progress:
            
            task = progress.add_task("[cyan]Processing .docx files...", total=len(docx_paths))
            
            for docx_path in docx_paths:
                if not docx_path.exists():
                    progress.update(task, advance=1)
                    continue
                
                # Check if already processed
                docx_hash = self._get_pdf_hash(docx_path)
                if not force_reload and docx_path.name in processed_docs:
                    if processed_docs[docx_path.name] == docx_hash:
                        progress.update(task, advance=1, description=f"[dim]Skipping {docx_path.name} (cached)[/dim]")
                        continue
                
                progress.update(task, description=f"[cyan]Processing {docx_path.name}...[/cyan]")
                
                # Extract text
                text = self.extract_text_from_docx(docx_path)
                
                if text:
                    # Create chunks
                    metadata = {
                        'source': docx_path.name,
                        'doc_hash': docx_hash,
                        'file_type': 'docx'
                    }
                    chunks = self.chunk_text(text, metadata)
                    all_chunks.extend(chunks)
                    total_chunks += len(chunks)
                    processed_docs[docx_path.name] = docx_hash
                
                progress.update(task, advance=1)
        
        # Generate embeddings and store in ChromaDB
        if all_chunks:
            self.console.print(f"\n[yellow]Generating embeddings for {total_chunks} chunks...[/yellow]")
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=self.console
            ) as progress:
                
                task = progress.add_task("[cyan]Creating embeddings...", total=total_chunks)
                
                # Process in batches
                batch_size = 32
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i + batch_size]
                    texts = [chunk['text'] for chunk in batch]
                    embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                    
                    # Prepare for ChromaDB
                    ids = [f"{chunk['metadata']['source']}_{chunk['metadata']['chunk_id']}" for chunk in batch]
                    metadatas = [chunk['metadata'] for chunk in batch]
                    
                    # Add to collection
                    self.collection.add(
                        ids=ids,
                        embeddings=embeddings.tolist(),
                        documents=texts,
                        metadatas=metadatas
                    )
                    
                    progress.update(task, advance=len(batch))
            
            # Save processed documents cache
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
            
            self.console.print(f"[green]✓ Successfully loaded {total_chunks} chunks into vector database[/green]")
        else:
            self.console.print("[yellow]No chunks to process[/yellow]")
    
    def retrieve_relevant_context(
        self, 
        query: str, 
        top_k: int = 5,
        min_similarity: float = 0.0
    ) -> List[dict]:
        """Retrieve the most relevant text chunks for a given query."""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=top_k
            )
            
            # Format results
            chunks = []
            if results['documents'] and results['documents'][0]:
                for i, doc in enumerate(results['documents'][0]):
                    metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                    distance = results['distances'][0][i] if results['distances'] else 0
                    similarity = 1 - distance  # Convert distance to similarity
                    
                    if similarity >= min_similarity:
                        chunks.append({
                            'text': doc,
                            'metadata': metadata,
                            'similarity': similarity
                        })
            
            return chunks
        
        except Exception as e:
            self.console.print(f"[red]Error retrieving context: {e}[/red]")
            return []
    
    def get_context_for_query(
        self, 
        query: str, 
        top_k: int = 5,
        include_metadata: bool = False
    ) -> str:
        """Get formatted context string for a query."""
        chunks = self.retrieve_relevant_context(query, top_k=top_k)
        
        if not chunks:
            return "No relevant information found."
        
        context_parts = []
        for chunk in chunks:
            if include_metadata:
                source = chunk['metadata'].get('source', 'Unknown')
                similarity = chunk['metadata'].get('similarity', 0)
                context_parts.append(f"[Source: {source} | Relevance: {similarity:.2%}]\n{chunk['text']}")
            else:
                context_parts.append(chunk['text'])
        
        return "\n\n---\n\n".join(context_parts)
    
    def get_stats(self) -> dict:
        """Get statistics about the knowledge base."""
        count = self.collection.count()
        
        return {
            "total_chunks": count,
            "embedding_model": self.embedding_model_name,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "cache_dir": str(self.cache_dir)
        }


# ============================================================================
# TUI Components
# ============================================================================

class TUIConsoleWrapper:
    """Wrapper to make TUI's ChatPanel work like a Rich Console for CommandManager."""
    
    def __init__(self, chat_panel, app):
        self.chat_panel = chat_panel
        self.app = app
        self.history_console = HistoryConsole()
        self._live_display = None
        self._live_content = []
        self._current_response = []
        
    def print(self, *args, **kwargs):
        """Print to the chat panel."""
        import re
        from rich.text import Text
        
        # Handle Rich Markdown objects specially to preserve markdown formatting
        if args and hasattr(args[0], '__class__') and args[0].__class__.__name__ == 'Markdown':
            # Extract the markdown text from Rich Markdown object
            markdown_obj = args[0]
            if hasattr(markdown_obj, 'markup'):
                message = markdown_obj.markup
            elif hasattr(markdown_obj, '_text'):
                message = markdown_obj._text
            else:
                message = str(markdown_obj)
        elif args and isinstance(args[0], Text):
            # Handle Rich Text objects - convert to markup string
            rich_text = args[0]
            message = str(rich_text)
        else:
            # Convert to string
            message = " ".join(str(arg) for arg in args)
            
            # Handle styled text by preserving Rich markup
            # BUT: Don't wrap content with code blocks, as they need markdown rendering
            style = kwargs.get('style', '')
            has_code_blocks = '```' in message
            
            if style and not has_code_blocks:
                # Wrap content with Rich markup tags to preserve styling
                # (only if no code blocks present)
                message = f"[{style}]{message}[/]"
            # If message has code blocks, don't add markup - let Markdown handle it
        
        # Skip empty messages
        if not message.strip():
            return
        
        # Determine message role based on content
        # System notifications include: file detection, tool usage info, etc.
        role = "assistant"
        if any(indicator in message for indicator in [
            "Found local file:",
            "Found local files:",
            "Using tools",
            "→ Using tools",
            "Running agent with tools",
            "Loading context", "Context loaded", "File not found", "Error loading",
            "Estimated token counts", "Total estimated tokens", "Warning:",
            "Database", "Processing", "Skipping", "Embedding", "Overwriting"
        ]):
            role = "system"
        
        # Always add messages during agent execution
        # Use call_from_thread for thread-safe UI updates
        try:
            self.app.call_from_thread(self.chat_panel.add_message, role, message)
        except Exception as e:
            # Fallback: try direct call if we're already in main thread
            try:
                self.chat_panel.add_message(role, message)
            except Exception:
                pass  # Silently fail if neither works
        
    def add_command_completer(self, command_manager):
        """Stub for console command completion (not needed in TUI)."""
        pass
    
    def set_live(self, live):
        """Set the live display context (compatibility with Rich Console)."""
        self._live_display = live
    
    def clear_live(self):
        """Clear the live display context."""
        self._live_display = None
        self._live_content = []
    
    def input(self, prompt="", **kwargs):
        """Stub for input (not used in TUI mode - handled by TUIConsole instead)."""
        return ""
    
    def status(self, message, **kwargs):
        """Status message display."""
        class StatusContext:
            def __enter__(self):
                return self
            def __exit__(self, *args):
                pass
        return StatusContext()
    
    def rule(self, title="", **kwargs):
        """Display a horizontal rule (just show the title)."""
        if title:
            self.print(f"\n{'─' * 40}\n{title}\n{'─' * 40}\n")
    
    def is_terminal(self):
        """Indicate this is a terminal-like interface."""
        return True
    
    @property
    def width(self):
        """Return a reasonable default width."""
        return 80
    
    @property
    def height(self):
        """Return a reasonable default height."""
        return 24
    
    def ask(self, prompt, **kwargs):
        """Ask for user input - not supported in TUI streaming mode."""
        self.print(f"[Input required: {prompt}]")
        return ""
    
    def start_agent_stream(self):
        """Start streaming agent output."""
        self._streaming_agent = True
        self._current_response = []
        # Create a message for agent output
        # Check if header should be shown
        show_header = not self.chat_panel._assistant_header_shown
        self.chat_panel._assistant_header_shown = True
        self._current_msg = ChatMessage("assistant", f"{EMOJI['robot']} Agent is working...", show_header=show_header)
        self.chat_panel.mount(self._current_msg)
        self.chat_panel.scroll_end(animate=False)
        
    def end_agent_stream(self):
        """End streaming agent output."""
        self._streaming_agent = False
        if hasattr(self, '_current_msg') and self._current_msg:
            # Final update with complete response
            if self._current_response:
                final_response = "\n".join(self._current_response)
                self._current_msg.update_content(final_response)
            
            # Schedule copy button detection for agent output (same as direct AI responses)
            # Must use call_from_thread since this is called from agent thread
            if self._current_msg and self._current_msg.parent:
                msg_widget = self._current_msg
                self.app.call_from_thread(
                    self.chat_panel._add_copy_buttons_to_message,
                    msg_widget
                )
            
            self._current_msg = None
        self._current_response = []



class CodeBlockHeader(Container):
    """A simple header bar for code blocks with a copy button."""
    
    DEFAULT_CSS = """
    CodeBlockHeader {
        height: 1;
        background: $panel;
        padding: 0 1;
        margin: 0;
        layout: horizontal;
    }
    
    CodeBlockHeader .spacer {
        width: 1fr;
    }
    
    CodeBlockHeader #copy-btn {
        width: 8;
        height: 1;
        min-width: 8;
        text-style: bold;
        border: none;
        padding: 0;
        dock: right;
    }

    CodeBlockHeader #copy-btn:focus {
        background: $success;
    }
    """
    
    def __init__(self, code: str, **kwargs):
        super().__init__(**kwargs)
        self.code = code
    
    def compose(self) -> ComposeResult:
        """Compose the header with copy button on the right."""
        from textual.widgets import Button
        
        yield Static("", classes="spacer")
        yield Button("Copy", id="copy-btn", variant="primary")
    
    @on(Button.Pressed, "#copy-btn")
    async def copy_code(self):
        """Copy code to clipboard."""
        try:
            # Use pyperclip if available
            import pyperclip
            pyperclip.copy(self.code)
            btn = self.query_one("#copy-btn", Button)
            btn.label = EMOJI['checkmark']
            await asyncio.sleep(2)
            btn.label = "Copy"
        except ImportError:
            # Fallback: just show feedback
            btn = self.query_one("#copy-btn", Button)
            btn.label = "Copied"
            await asyncio.sleep(2)
            btn.label = "Copy"


class CopyButton(Static):
    """Overlay copy button that floats over code blocks."""
    
    DEFAULT_CSS = """
    CopyButton {
        width: 8;
        height: 1;
        offset: 0 0;
        layer: overlay;
        background: $surface;
        content-align: center middle;
        border: none;
    }
    
    CopyButton:hover {
        background: $surface-lighten-1;
    }
    
    CopyButton.success {
        background: $success;
    }
    """
    
    def __init__(self, code_content: str, code_index: int):
        super().__init__()
        self.code_content = code_content
        self.code_index = code_index
        self.label = "Copy"
        self.can_focus = True  # Make it clickable
        
    def render(self) -> str:
        return self.label
    
    def on_click(self) -> None:
        """Handle copy button click."""
        try:
            import pyperclip
            pyperclip.copy(self.code_content)
            # Show success state
            self.label = EMOJI['checkmark']
            self.add_class("success")
            self.refresh()
            # Reset after 1.5 seconds
            self.set_timer(1.5, self._reset_button)
        except Exception:
            pass
    
    def _reset_button(self):
        """Reset button to default state."""
        self.label = "Copy"
        self.remove_class("success")
        self.refresh()


class ChatMessage(Container):
    """A beautifully styled chat message widget with Material Design principles."""
    
    DEFAULT_CSS = """
    ChatMessage {
        height: auto;
        margin: 0 0 1 0;
        padding: 0;
    }
    
    ChatMessage.user-message {
        border: solid $accent-muted;
        border-left: thick $accent;
        padding: 1 2;
    }
    
    ChatMessage.assistant-message {
        background: transparent;
        border: none;
        padding: 0;
    }
    
    ChatMessage.system-message {
        border-left: thick $foreground-muted;
        padding: 1 2;
    }
    
    ChatMessage.error-message {
        border-left: thick $error;
        padding: 1 2;
    }
    
    ChatMessage.welcome-message {
        background: transparent;
        border: none;
        padding: 0;
    }
    
    .message-header {
        color: $text-muted;
        text-style: bold;
        margin-bottom: 1;
    }
    
    .message-content {
        background: transparent;
    }
    
    /* RichLog styling to match Markdown */
    ChatMessage RichLog.message-content {
        background: transparent;
        border: none;
        padding: 0;
        scrollbar-size: 0 0;
        height: auto;
    }
    """
    
    can_focus = False  # Prevent focus/selection
    
    @staticmethod
    def _has_rich_color_markup(content: str) -> bool:
        """
        Check if content has Rich markup tags (colors, styles, etc.).
        Returns True if content has Rich formatting tags that need to be converted.
        """
        # Pattern to detect Rich markup tags:
        # 1. Color tags: [green], [bold green], [#ff0000], etc.
        # 2. Style tags: [bold], [italic], [dim], [underline], etc.
        # 3. Closing tags: [/bold], [/], etc.
        
        # Check for color tags (with or without style prefix)
        color_pattern = r'\[(?:(?:bold|italic|dim|underline|strike|blink|reverse)\s+)?(?:red|green|blue|yellow|cyan|magenta|white|black|orange|purple|pink|deep_pink|grey\d+|#[0-9a-fA-F]{6}|rgb\([^\)]+\))'
        if re.search(color_pattern, content, re.IGNORECASE):
            return True
        
        # Check for standalone style tags: [bold], [italic], [dim], [underline], [strike]
        style_pattern = r'\[(?:bold|italic|dim|underline|strike|blink|reverse)\]'
        if re.search(style_pattern, content, re.IGNORECASE):
            return True
        
        # Check for closing tags: [/bold], [/italic], [/dim], [/]
        closing_pattern = r'\[/(?:bold|italic|dim|underline|strike|blink|reverse)?\]'
        if re.search(closing_pattern, content, re.IGNORECASE):
            return True
        
        return False
    
    @staticmethod
    def _has_markdown_syntax(content: str) -> bool:
        """
        Check if content has markdown syntax (code blocks, headers, lists, etc.).
        Returns True if content needs markdown rendering.
        """
        # Check for common markdown patterns
        markdown_patterns = [
            r'```',           # Code blocks
            r'^#{1,6}\s',     # Headers
            r'^\s*[-*+]\s',   # Unordered lists
            r'^\s*\d+\.\s',   # Ordered lists
            r'\[.+?\]\(.+?\)', # Links
            r'^\|.+\|$',      # Tables
        ]
        for pattern in markdown_patterns:
            if re.search(pattern, content, re.MULTILINE):
                return True
        return False
    
    @staticmethod
    def _convert_rich_markup_to_markdown(content: str) -> str:
        """
        Convert Rich markup tags to markdown equivalents.
        Strips all Rich formatting and converts styles to markdown.
        
        Strategy:
        1. Protect escaped brackets (\\[) by temporarily replacing them
        2. Convert style tags (bold, italic) to markdown
        3. Remove all color/style tags completely
        4. Clean up any remaining Rich markup
        5. Restore escaped brackets as literal text
        """
        # Step 1: Protect escaped brackets by replacing them with placeholders
        # This prevents \\[bold] from being processed as Rich markup
        ESCAPED_OPEN_PLACEHOLDER = "\x00ESCAPED_OPEN_BRACKET\x00"
        ESCAPED_CLOSE_PLACEHOLDER = "\x00ESCAPED_CLOSE_BRACKET\x00"
        
        content = content.replace(r'\[', ESCAPED_OPEN_PLACEHOLDER)
        content = content.replace(r'\]', ESCAPED_CLOSE_PLACEHOLDER)
        
        # Step 2: Handle [bold ...] tags with any modifiers (colors, etc.)
        # Use a more aggressive pattern that captures the style word and ignores everything after
        # Match: [bold ...] content [/bold ...] OR [bold ...] content [/]
        
        # Bold: [bold X] -> **, [/bold X] or [/] -> **
        def replace_bold(match):
            return f"**{match.group(1)}**"
        content = re.sub(r'\[bold[^\]]*\](.*?)\[/bold[^\]]*\]', replace_bold, content, flags=re.IGNORECASE | re.DOTALL)
        content = re.sub(r'\[bold[^\]]*\](.*?)\[/\]', replace_bold, content, flags=re.IGNORECASE | re.DOTALL)
        
        # Italic: [italic X] -> *, [/italic X] or [/] -> *
        def replace_italic(match):
            return f"*{match.group(1)}*"
        content = re.sub(r'\[italic[^\]]*\](.*?)\[/italic[^\]]*\]', replace_italic, content, flags=re.IGNORECASE | re.DOTALL)
        content = re.sub(r'\[italic[^\]]*\](.*?)\[/\]', replace_italic, content, flags=re.IGNORECASE | re.DOTALL)
        
        # Dim: [dim X] -> _, [/dim X] or [/] -> _
        def replace_dim(match):
            return f"_{match.group(1)}_"
        content = re.sub(r'\[dim[^\]]*\](.*?)\[/dim[^\]]*\]', replace_dim, content, flags=re.IGNORECASE | re.DOTALL)
        content = re.sub(r'\[dim[^\]]*\](.*?)\[/\]', replace_dim, content, flags=re.IGNORECASE | re.DOTALL)
        
        # Step 3: Remove ALL remaining Rich markup tags
        # This includes colors, standalone styles, and any other Rich tags
        # Match any [tag] or [tag something] pattern
        content = re.sub(r'\[[^\]]+\]', '', content)
        
        # Step 4: Restore escaped brackets as literal brackets (without backslash)
        content = content.replace(ESCAPED_OPEN_PLACEHOLDER, '[')
        content = content.replace(ESCAPED_CLOSE_PLACEHOLDER, ']')
        
        return content
    
    @staticmethod
    def _process_backticks_only(content: str) -> str:
        """
        Process only backticks in user messages for inline code formatting.
        Escapes other markdown characters to prevent them from being interpreted,
        but preserves backticks so they render as inline code.
        """
        # Store backtick content temporarily with indices
        backtick_pattern = r'`([^`]+)`'
        backtick_matches = []
        
        # Find all backtick sections
        for match in re.finditer(backtick_pattern, content):
            backtick_matches.append({
                'start': match.start(),
                'end': match.end(),
                'full': match.group(0),
                'inner': match.group(1)
            })
        
        # Build the result by processing character by character
        result = []
        i = 0
        
        while i < len(content):
            # Check if we're at the start of a backtick section
            in_backtick = False
            for bt_match in backtick_matches:
                if i == bt_match['start']:
                    # Add the entire backtick section as-is
                    result.append(bt_match['full'])
                    i = bt_match['end']
                    in_backtick = True
                    break
            
            if not in_backtick:
                # Escape markdown special characters
                char = content[i]
                if char in '*_#[]()~>-+=|\\':
                    result.append('\\' + char)
                else:
                    result.append(char)
                i += 1
        
        return ''.join(result)
    
    def __init__(self, role: str, content: str, timestamp: Optional[datetime] = None, show_header: bool = True, agent_name: Optional[str] = None):
        super().__init__()
        self.role = role
        self.timestamp = timestamp or datetime.now()
        self.show_header = show_header
        self.agent_name = agent_name  # Store the agent name if provided
        self._header_widget = None  # Store header widget reference
        self._content_widget = None  # Store content widget reference
        self._widget_type = None  # Track which widget type we're using ('markdown' or 'richlog')
        
        # Store original content before processing
        self.original_content = content
        
        # Determine which widget to use based on content
        # Priority: RichLog for pure Rich markup, Markdown for markdown syntax or mixed content
        if role != "user":
            has_rich_colors = self._has_rich_color_markup(content)
            has_markdown = self._has_markdown_syntax(content)
            
            # Decision logic:
            # - If has markdown syntax (code blocks, headers, etc.) -> Use Markdown (convert Rich to markdown if present)
            # - If has ONLY Rich colors with no markdown -> Use RichLog (preserve colors)
            # - Otherwise -> Use Markdown (safe default)
            if has_markdown or not has_rich_colors:
                # Use Markdown widget - only convert Rich markup if it's actually present
                self._widget_type = 'markdown'
                if has_rich_colors:
                    content = self._convert_rich_markup_to_markdown(content)
            else:
                # Use RichLog widget - preserve Rich markup for color rendering
                self._widget_type = 'richlog'
                # Keep content as-is with Rich markup
        else:
            # User messages always use Markdown with backticks only
            self._widget_type = 'markdown'
            content = self._process_backticks_only(content)
        
        # Store the processed content
        self.content = content
        
        # Apply role-specific styling
        if role == "user":
            self.add_class("user-message")
        elif role == "assistant":
            self.add_class("assistant-message")
        elif role == "system":
            self.add_class("system-message")
        elif role == "error":
            self.add_class("error-message")
    
    def update_content(self, new_content: str):
        """Update the message content and refresh the display properly."""
        # Store original content
        self.original_content = new_content
        
        # Re-determine widget type based on new content
        if self.role != "user":
            has_rich_colors = self._has_rich_color_markup(new_content)
            has_markdown = self._has_markdown_syntax(new_content)
            
            if has_markdown or not has_rich_colors:
                new_widget_type = 'markdown'
                if has_rich_colors:
                    new_content = self._convert_rich_markup_to_markdown(new_content)
            else:
                new_widget_type = 'richlog'
                # Keep Rich markup for RichLog
        else:
            new_widget_type = 'markdown'
            new_content = self._process_backticks_only(new_content)
        
        self.content = new_content
        
        # If widget type changed, need to recompose
        if new_widget_type != self._widget_type:
            self._widget_type = new_widget_type
            self._content_widget = None
            self.refresh(recompose=True)
            return
        
        # Update the appropriate widget type
        try:
            if self._widget_type == 'markdown':
                content_widget = self.query_one(".message-content", Markdown)
                content_widget.update(new_content)
            elif self._widget_type == 'richlog':
                content_widget = self.query_one(".message-content", RichLog)
                content_widget.clear()
                content_widget.write(new_content, markup=True)
        except Exception:
            # If widget doesn't exist yet, do a full refresh
            self.refresh(recompose=True)
        
    def compose(self) -> ComposeResult:
        """Compose the message display with markdown."""
        role_icons = {
            "user": EMOJI['user'],
            "assistant": EMOJI['robot'],
            "system": EMOJI['info'],
            "error": EMOJI['cross']
        }
        
        role_labels = {
            "user": "You",
            "assistant": f"ValBot (Agent: {self.agent_name})" if self.agent_name else "ValBot",
            "system": "System",
            "error": "Error"
        }
        
        # Dynamically load theme colors from the app's theme
        theme = self.app.theme_variables
        role_colors = {
            "user": theme.get("accent", "#0696d4"),
            "assistant": theme.get("primary", "#989af5"),
            "system": theme.get("foreground-muted", "#fafafa"),
            "error": theme.get("error", "#ef4444")
        }
        
        icon = role_icons.get(self.role, "•")
        role_label = role_labels.get(self.role, self.role.title())
        role_color = role_colors.get(self.role, "#94a3b8")
        time_str = self.timestamp.strftime("%H:%M:%S")
        
        # Create header with colored role label (only if show_header is True)
        if self.show_header:
            if self._header_widget is None:
                header = f"{icon} [{role_color}]{role_label}[/] · {time_str}"
                self._header_widget = Static(header, classes="message-header", markup=True)
            yield self._header_widget
        
        # Create content widget based on widget type determined in __init__
        # - 'markdown': Use Markdown widget for markdown syntax (code blocks, headers, etc.)
        # - 'richlog': Use RichLog widget to preserve Rich color markup
        if self._content_widget is None:
            if self._widget_type == 'richlog':
                # Use RichLog to preserve Rich markup with colors
                richlog = RichLog(highlight=True, markup=True, classes="message-content")
                richlog.write(self.content)
                self._content_widget = richlog
            else:
                # Use Markdown for markdown syntax or converted content
                self._content_widget = Markdown(self.content, classes="message-content")
        yield self._content_widget
    
    def add_copy_buttons_for_code_blocks(self):
        """Add overlay copy buttons positioned over code blocks (legacy method)."""
        # This method is now primarily used by the new post-streaming detection system
        # Find code blocks in the content
        code_block_pattern = r'```(\w*)\n(.*?)\n```'
        matches = list(re.finditer(code_block_pattern, self.content, re.DOTALL))
        
        if not matches:
            return
        
        # Remove any existing copy buttons first
        for existing_btn in list(self.query(CopyButton)):
            existing_btn.remove()
        
        # Get the markdown widget to calculate positions
        try:
            markdown_widget = self.query_one(".message-content", Markdown)
        except Exception:
            return
        
        # Add overlay copy button for each code block
        for i, match in enumerate(matches):
            code_content = match.group(2)
            # Create overlay button
            btn = CopyButton(code_content, i)
            # Position it - we'll update positions after mounting
            self.mount(btn)
            
        # Use the enhanced positioning method with a delay
        self.call_after_refresh(lambda: self.set_timer(0.1, self._position_copy_buttons_accurately))
    
    def _position_copy_buttons(self):
        """Position copy buttons at top-right of each code block."""
        try:
            buttons = list(self.query(CopyButton))
            if not buttons:
                return
            
            # Get markdown widget to reference
            markdown_widget = self.query_one(".message-content", Markdown)
            
            # Find positions of code blocks in the content
            code_block_pattern = r'```(\w*)\n(.*?)\n```'
            matches = list(re.finditer(code_block_pattern, self.content, re.DOTALL))
            
            if not matches:
                return
            
            # Get the width of the parent container to position from right
            parent_width = self.size.width if self.size.width > 0 else 80
            x_position = parent_width - 10  # 10 chars from right (button is 8 wide + margin)
            
            # Calculate cumulative line positions for each code block
            lines_before_content = 1  # Reduced to account for markdown padding
            
            for i, (btn, match) in enumerate(zip(buttons, matches)):
                # Get text from start to this code block start
                match_start_content = self.content[:match.start()]
                lines_to_match = match_start_content.count('\n')
                
                # Position button at the start of this code block
                # Lower by 2 rows for better visual alignment
                y_position = lines_before_content + lines_to_match + 2
                
                # Set the offset (x, y) from top-left of the message
                btn.styles.offset = (x_position, y_position)
                
        except Exception as e:
            # Fallback: position buttons in a simple vertical stack on the right
            try:
                parent_width = self.size.width if self.size.width > 0 else 80
                x_position = parent_width - 10
                for i, btn in enumerate(buttons):
                    y_offset = 1 + (i * 10)  # Rough spacing, starting closer to top
                    btn.styles.offset = (x_position, y_offset)
            except:
                pass
    
    def on_resize(self) -> None:
        """Reposition copy buttons when the widget is resized."""
        # Reposition buttons after a short delay to ensure layout is updated
        self.set_timer(0.2, self._position_copy_buttons_accurately)
    
    def _position_copy_buttons_accurately(self):
        """Enhanced copy button positioning with better accuracy after layout settling."""
        try:
            buttons = list(self.query(CopyButton))
            if not buttons:
                return
            
            # Get the markdown widget and ensure it's fully rendered
            try:
                markdown_widget = self.query_one(".message-content", Markdown)
                # Give markdown time to fully render if needed
                if not hasattr(markdown_widget, 'size') or markdown_widget.size.width == 0:
                    # Schedule retry if layout not ready
                    self.set_timer(0.1, self._position_copy_buttons_accurately)
                    return
            except Exception:
                return
            
            # Find code block positions in content
            code_block_pattern = r'```(\w*)\n(.*?)\n```'
            matches = list(re.finditer(code_block_pattern, self.content, re.DOTALL))
            
            if not matches or len(matches) != len(buttons):
                return
            
            # Calculate positioning based on actual widget dimensions
            container_width = self.size.width if self.size.width > 0 else 80
            x_position = max(container_width - 12, 2)  # Position from right edge with margin
            
            # More accurate line counting for positioning
            # Account for the markdown widget's internal structure
            lines_before_content = 1  # Reduced to account for markdown padding
            
            for i, (btn, match) in enumerate(zip(buttons, matches)):
                # Find the line where this code block starts
                match_start_content = self.content[:match.start()]
                lines_to_match = match_start_content.count('\n')
                
                # Position button at the start of the code block
                # Adjust for markdown rendering (code blocks typically have 1 line padding)
                # Lower by 2 rows for better visual alignment
                y_position = lines_before_content + lines_to_match + 2
                
                # Fine-tune positioning to align with the top border of the code block
                # This accounts for the ``` delimiter line
                y_position = max(y_position, 0)
                
                # Apply positioning with better alignment
                btn.styles.offset = (x_position, y_position)
                
        except Exception as e:
            # Fallback to original positioning method
            self._position_copy_buttons()
    
    def add_copy_buttons_if_needed(self):
        """Legacy method - now deprecated in favor of post-streaming detection."""
        # This method is kept for compatibility but should not be used during streaming
        # Copy button detection is now handled after streaming completes for better accuracy
        pass


class DatabaseProgressWidget(Static):
    """A custom widget to display database creation progress in the chat."""
    
    DEFAULT_CSS = """
    DatabaseProgressWidget {
        height: auto;
        margin: 1 0;
        padding: 1 2;
        background: $surface;
        border: solid $primary;
    }
    
    DatabaseProgressWidget #progress-status {
        width: 100%;
        margin-bottom: 1;
    }
    
    DatabaseProgressWidget #progress-bar {
        width: 100%;
    }
    """
    
    def __init__(self, total: int, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.total = total
        self.border_title = f"{EMOJI['database']} Database Creation Progress"
        
    def compose(self) -> ComposeResult:
        """Compose the progress display."""
        yield Label(f"Processing files... (0/{int(self.total/100)})", id="progress-status")
        yield ProgressBar(total=self.total, show_eta=False, id="progress-bar")
    
    def update_progress(self, current: int, status: str):
        """Update progress bar and status label."""
        try:
            progress_bar = self.query_one("#progress-bar", ProgressBar)
            status_label = self.query_one("#progress-status", Label)
            
            # Update progress by setting the progress attribute directly
            progress_bar.progress = current
            progress_bar.refresh()
            
            # Update status text
            status_label.update(status)
        except Exception as e:
            pass  # Widget might not be mounted yet


class ChatTerminalOutput(Static):
    """A terminal output widget that displays inline in the chat with border."""
    
    DEFAULT_CSS = """
    ChatTerminalOutput {
        height: auto;
        margin: 0 0 1 0;
        padding: 1 2;
        background: $surface;
        border: solid $primary;
    }
    
    #terminal-output-log {
        background: $surface;
        height: auto;
        scrollbar-size: 0 0;
    }
    """
    
    def __init__(self, command: str, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.command = command
        self.border_title = f"{EMOJI['keyboard']} Terminal"
        self.terminal_manager = TerminalManager()
        
    def compose(self) -> ComposeResult:
        """Compose the terminal output display."""
        yield RichLog(highlight=True, markup=True, id="terminal-output-log")
        
    async def on_mount(self) -> None:
        """Execute command and display output when mounted."""
        log = self.query_one("#terminal-output-log", RichLog)
        log.write(Text(f"$ {self.command}", style="bold cyan"))
        
        async for output_type, line in self.terminal_manager.run_command(self.command):
            if output_type == "stdout":
                log.write(Text(line, style="white"))
            elif output_type == "stderr":
                log.write(Text(line, style="red"))
            elif output_type == "error":
                log.write(Text(f"Error: {line}", style="bold red"))
        
        log.write("")  # Empty line after command


class ChatPanel(VerticalScroll):
    """Modern Material Design panel for displaying chat messages."""
    
    DEFAULT_CSS = """
    ChatPanel {
        height: 1fr;
        background: $background;
        padding: 1 2;
        scrollbar-size: 1 1;
    }
    
    ChatPanel:focus {
        border: none;
    }
    
    ChatPanel > Static {
        width: 1fr;
    }
    
    ChatPanel > ChatMessage > Markdown {
        background: transparent;
    }
    """
    
    can_focus = False  # Prevent focus
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.messages: List[ChatMessage] = []
        self._assistant_header_shown = False  # Track if assistant header shown after last user message
        self.current_agent: Optional[str] = None  # Track the currently active agent
    
    def add_message(self, role: str, content: str, agent_name: Optional[str] = None, raw_content: Optional[str] = None, skip_autosave: bool = False):
        """Add a new message to the chat.
        
        Args:
            role: Message role (user, assistant, system, error)
            content: Formatted content for display
            agent_name: Optional agent name
            raw_content: Optional raw content for saving (without formatting). If not provided, content is used.
            skip_autosave: If True, skip autosaving this message (used when loading history)
        """
        # Reset header flag when user sends a message
        if role == "user":
            self._assistant_header_shown = False
        
        # For assistant messages, only show header if not already shown
        show_header = True
        if role == "assistant":
            show_header = not self._assistant_header_shown
            self._assistant_header_shown = True
        
        # Use the provided agent_name, or fall back to the current_agent if not specified
        effective_agent_name = agent_name if agent_name is not None else self.current_agent
        
        msg = ChatMessage(role, content, show_header=show_header, agent_name=effective_agent_name)
        self.messages.append(msg)
        self.mount(msg)
        self.scroll_end(animate=True)
        
        # Schedule copy button detection for messages with code blocks
        # (especially important for agent output)
        if '```' in content:
            # Use delayed detection to ensure proper layout and positioning
            self.app.call_later(self._add_copy_buttons_to_message, msg)
        
        # Autosave assistant and user messages (skip system messages)
        # Skip autosave if explicitly requested (e.g., when loading history)
        if not skip_autosave:
            # Get MainScreen to call autosave
            try:
                main_screen = self.app.screen
                if hasattr(main_screen, '_autosave_message'):
                    # Use raw_content for saving if provided, otherwise use content
                    content_to_save = raw_content if raw_content is not None else content
                    asyncio.create_task(main_screen._autosave_message(role, content_to_save, effective_agent_name))
            except Exception:
                pass  # Silently fail if autosave not available
    
    async def add_terminal_output(self, command: str):
        """Add a terminal output widget to the chat."""
        terminal_widget = ChatTerminalOutput(command)
        await self.mount(terminal_widget)
        self.scroll_end(animate=True)
    
    def _add_copy_buttons_to_message(self, message_widget):
        """
        Add copy buttons to a message with code blocks.
        Single-phase approach with adequate delay for layout settling.
        """
        if not message_widget or not message_widget.parent:
            return
        
        try:
            # Clear any existing copy buttons first
            for existing_btn in list(message_widget.query(CopyButton)):
                existing_btn.remove()
            
            # Detect code blocks in the final content
            code_block_pattern = r'```(\w*)\n(.*?)\n```'
            matches = list(re.finditer(code_block_pattern, message_widget.content, re.DOTALL))
            
            if matches:
                # Add copy buttons for each code block
                for i, match in enumerate(matches):
                    code_content = match.group(2)
                    btn = CopyButton(code_content, i)
                    message_widget.mount(btn)
                
                # Wait for layout to settle, then position accurately
                # Single delay is sufficient (0.5s total) instead of two phases (0.2s + 0.3s)
                self.set_timer(0.5, lambda: self._position_buttons_for_message(message_widget))
            
        except Exception:
            # Ignore errors in copy button detection
            pass
    
    def _position_buttons_for_message(self, message_widget):
        """Position copy buttons accurately after layout has settled."""
        if message_widget and message_widget.parent:
            try:
                # Use the improved positioning method
                message_widget._position_copy_buttons_accurately()
            except Exception:
                # Fallback to basic positioning
                try:
                    message_widget._position_copy_buttons()
                except Exception:
                    pass
        
    def clear_messages(self):
        """Clear all messages."""
        for msg in self.messages:
            msg.remove()
        self.messages.clear()
        # Reset the header flag and current agent
        self._assistant_header_shown = False
        self.current_agent = None
        # Also remove any ChatMessage widgets that were mounted directly
        # (e.g., streaming messages that weren't added to the messages list)
        for widget in self.query("ChatMessage"):
            widget.remove()
        # Remove any terminal output widgets
        for widget in self.query("ChatTerminalOutput"):
            widget.remove()

class CustomDirectoryTree(DirectoryTree):
    """Custom DirectoryTree with custom file and folder icons."""
    
    ICON_FILE = f"{EMOJI['file']} "
    ICON_FOLDER = f"{EMOJI['folder']} "
    ICON_FOLDER_OPEN = f"{EMOJI['open_folder']} "
    
    def render_label(self, node: TreeNode[DirEntry], base_style: Style, style: Style) -> Text:
        """Render tree label with custom folder and file icons."""
        node_label = node._label.copy()
        node_label.stylize(style)
        
        if node._allow_expand:
            # Use custom folder icons based on expand state
            icon = self.ICON_FOLDER_OPEN if node.is_expanded else self.ICON_FOLDER
            prefix = (icon, base_style)
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--folder", partial=True),
            )
        else:
            # Use custom file icon
            prefix = (self.ICON_FILE, base_style)
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--file", partial=True),
            )
            node_label.highlight_regex(
                r"\\..*$",
                self.get_component_rich_style("directory-tree--extension", partial=True),
            )
        
        if node._label.plain.startswith("."):
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--hidden", partial=True)
            )
        
        text = Text.assemble(prefix, node_label)
        return text


class FileExplorerPanel(Container):
    """Material Design file system navigation panel."""
    
    DEFAULT_CSS = """
    FileExplorerPanel {
        width: 35%;
        height: 100vh;
        background: $surface;
        border: solid $primary;
        padding: 1;
        display: none;
        layer: overlay;
        offset: 0 0;
    }
    
    FileExplorerPanel.visible {
        display: block;
    }
    
    #nav-buttons {
        width: 100%;
        height: 1;
        layout: horizontal;
        margin-bottom: 1;
    }
    
    #nav-buttons Button {
        width: 1fr;
        height: 1;
        min-width: 0;
        padding: 0;
        margin: 0 1 0 0;
        border: none;
        background: $panel;
        content-align: center middle;
    }
    
    #nav-buttons Button:hover {
        background: $surface-lighten-1;
    }
    
    #nav-buttons Button:disabled {
        opacity: 0.4;
    }
    
    #nav-buttons Button:last-child {
        margin: 0;
    }
    
    #address-bar {
        width: 100%;
        height: 1;
        margin-bottom: 1;
        padding: 0 1;
        background: $panel;
        border: none;
    }
    
    #file-tree {
        height: 1fr;
        scrollbar-size: 1 1;
    }
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.border_title = f"{EMOJI['folder']} File Explorer"
        self.current_path = Path.cwd()  # Start with current working directory
        self.initial_path = Path.cwd()  # Store initial directory for reset
        self.history_back = []  # History for back navigation
        self.history_forward = []  # History for forward navigation
        self._navigating = False  # Lock to prevent concurrent navigation
        
    def compose(self) -> ComposeResult:
        """Compose the file explorer."""
        # Navigation buttons
        with Horizontal(id="nav-buttons"):
            yield Button("↻", id="nav-refresh", variant="primary")
            yield Button("↑", id="nav-up", variant="primary")
            yield Button("←", id="nav-back", variant="primary")
            yield Button("→", id="nav-forward", variant="primary")
            yield Button("✕", id="nav-close", variant="primary")
        
        # Address bar input
        address_input = Input(
            placeholder="Enter path...",
            value=str(self.current_path),
            id="address-bar"
        )
        yield address_input
        # Directory tree starting from current working directory
        yield CustomDirectoryTree(str(self.current_path), id="file-tree")
    
    async def on_mount(self) -> None:
        """Update button states on mount."""
        self.update_nav_buttons()
    
    def update_nav_buttons(self) -> None:
        """Update navigation button states based on history."""
        try:
            back_btn = self.query_one("#nav-back", Button)
            back_btn.disabled = len(self.history_back) == 0
        except Exception:
            pass
        
        try:
            forward_btn = self.query_one("#nav-forward", Button)
            forward_btn.disabled = len(self.history_forward) == 0
        except Exception:
            pass
    
    def on_input_changed(self, event: Input.Changed) -> None:
        """Handle address bar text changes - auto-navigate if valid path."""
        if event.input.id != "address-bar":
            return
        
        new_path = event.value.strip()
        if not new_path:
            return
        
        # Try to expand and resolve the path
        try:
            expanded_path = Path(new_path).expanduser().resolve()
            
            # Only auto-navigate if path exists and is a directory
            if expanded_path.exists() and expanded_path.is_dir():
                # Navigate without changing focus (keep focus on address bar)
                self.run_worker(self.navigate_to_path(expanded_path, add_to_history=False, keep_focus_on_input=True))
        except Exception:
            # Invalid path - ignore
            pass
    
    async def on_input_submitted(self, event: Input.Submitted) -> None:
        """Handle address bar submission - navigate and move focus to tree."""
        if event.input.id != "address-bar":
            return
        
        new_path = event.value.strip()
        if not new_path:
            return
        
        # Expand user home directory and resolve the path
        try:
            expanded_path = Path(new_path).expanduser().resolve()
            
            if not expanded_path.exists():
                self.app.bell()
                event.input.value = str(self.current_path)
                return
            
            if not expanded_path.is_dir():
                self.app.bell()
                event.input.value = str(self.current_path)
                return
            
            # Navigate to the new path (with history) and move focus to tree
            await self.navigate_to_path(expanded_path, add_to_history=True, keep_focus_on_input=False)
            
        except Exception as e:
            # Invalid path - reset to current path
            self.app.bell()
            event.input.value = str(self.current_path)
    
    async def navigate_to_path(self, new_path: Path, add_to_history: bool = True, keep_focus_on_input: bool = False) -> None:
        """Navigate to a new directory path."""
        # Prevent concurrent navigation
        if self._navigating:
            return
        
        self._navigating = True
        
        try:
            # Add current path to back history if navigating to a different path
            if add_to_history and new_path != self.current_path:
                self.history_back.append(self.current_path)
                self.history_forward.clear()  # Clear forward history on new navigation
            
            # Update the current path
            self.current_path = new_path
            
            # Update address bar only if not actively typing in it
            if not keep_focus_on_input:
                try:
                    address_bar = self.query_one("#address-bar", Input)
                    address_bar.value = str(self.current_path)
                except Exception:
                    pass
            
            # Remove old directory tree and create new one
            try:
                old_tree = self.query_one("#file-tree", CustomDirectoryTree)
                await old_tree.remove()
            except Exception:
                pass
            
            # Mount new directory tree with the new path
            new_tree = CustomDirectoryTree(str(self.current_path), id="file-tree")
            await self.mount(new_tree)
            
            # Reload the tree to show contents
            new_tree.reload()
            
            # Update navigation button states
            self.update_nav_buttons()
            
            # Focus management
            if not keep_focus_on_input:
                # Focus the tree when explicitly navigating (Enter key, clicking folders/buttons)
                # Use set_timer to ensure tree is fully loaded before focusing
                def focus_tree_delayed():
                    try:
                        tree = self.query_one("#file-tree", CustomDirectoryTree)
                        tree.focus()
                        # Move cursor to the first item in the tree
                        tree.cursor_line = 0
                        tree.scroll_to_line(0)
                    except Exception as e:
                        pass
                
                self.set_timer(0.1, focus_tree_delayed)
            # If keep_focus_on_input is True, don't change focus at all (typing in address bar)
        finally:
            self._navigating = False
    
    async def on_directory_tree_directory_selected(self, event: DirectoryTree.DirectorySelected) -> None:
        """Handle directory selection in the tree."""
        selected_path = Path(event.path)
        # Navigate to the selected directory
        await self.navigate_to_path(selected_path, add_to_history=True, keep_focus_on_input=False)
        # Prevent default expansion behavior
        event.stop()
    
    @on(Button.Pressed, "#nav-refresh")
    async def action_refresh(self):
        """Refresh to initial working directory."""
        await self.navigate_to_path(self.initial_path, add_to_history=True, keep_focus_on_input=False)
    
    @on(Button.Pressed, "#nav-up")
    async def action_go_up(self):
        """Go up one directory level."""
        parent = self.current_path.parent
        if parent != self.current_path:  # Check we're not at root
            await self.navigate_to_path(parent, add_to_history=True, keep_focus_on_input=False)
    
    @on(Button.Pressed, "#nav-back")
    async def action_go_back(self):
        """Go back to previous directory."""
        if self.history_back:
            # Move current to forward history
            self.history_forward.append(self.current_path)
            # Pop from back history
            previous_path = self.history_back.pop()
            await self.navigate_to_path(previous_path, add_to_history=False, keep_focus_on_input=False)
    
    @on(Button.Pressed, "#nav-forward")
    async def action_go_forward(self):
        """Go forward to next directory."""
        if self.history_forward:
            # Move current to back history
            self.history_back.append(self.current_path)
            # Pop from forward history
            next_path = self.history_forward.pop()
            await self.navigate_to_path(next_path, add_to_history=False, keep_focus_on_input=False)
    
    @on(Button.Pressed, "#nav-close")
    def action_close_explorer(self):
        """Close the file explorer panel."""
        self.remove_class("visible")


class ChatHistoryTree(DirectoryTree):
    """Custom DirectoryTree that only shows .log files without folders, sorted by date."""
    
    ICON_FILE = ""  # Empty icon for history files
    
    def filter_paths(self, paths: Iterable[Path]) -> Iterable[Path]:
        """Filter to only show .log files, no directories."""
        for path in paths:
            # Only yield files that end with .log
            if path.is_file() and path.suffix == '.log':
                yield path
    
    def _populate_node(self, node, content):
        """Override to sort files by modification time before populating."""
        # Get the directory entries
        from textual.widgets._directory_tree import DirEntry
        
        if hasattr(content, '__iter__'):
            # Convert to list and sort by modification time (newest first)
            entries = list(content)
            # Filter and sort only .log files
            # Check if entries are Path objects or DirEntry objects
            if entries and isinstance(entries[0], Path):
                log_entries = [e for e in entries if e.is_file() and e.suffix == '.log']
                log_entries.sort(key=lambda e: e.stat().st_mtime, reverse=True)
            else:
                # DirEntry objects
                log_entries = [e for e in entries if hasattr(e, 'path') and e.path.is_file() and e.path.suffix == '.log']
                log_entries.sort(key=lambda e: e.path.stat().st_mtime, reverse=True)
            
            # Call parent with sorted entries
            content = iter(log_entries)
        
        return super()._populate_node(node, content)
    
    def render_label(self, node: TreeNode[DirEntry], base_style: Style, style: Style) -> Text:
        """Render tree label without .log extension and timestamp."""
        node_label = node._label.copy()
        # Remove .log extension from the label
        label_text = str(node_label.plain)
        if label_text.endswith('.log'):
            label_text = label_text[:-4]
            
            # Remove last two underscore-separated parts (date and time)
            # Pattern: prefix_20241204_123045 -> prefix
            parts = label_text.split('_')
            if len(parts) > 2:
                # Remove last 2 items
                parts.pop()
                parts.pop()
                label_text = '_'.join(parts)
            
            node_label = Text(label_text, style=node_label.style)
        
        node_label.stylize(style)
        
        if node._allow_expand:
            prefix = ("▶ " if node.is_expanded else "▸ ", base_style + TOGGLE_STYLE)
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--folder", partial=True),
            )
        else:
            prefix = (self.ICON_FILE, base_style)
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--file", partial=True),
            )
            node_label.highlight_regex(
                r"\\..*$",
                self.get_component_rich_style("directory-tree--extension", partial=True),
            )
        
        if node._label.plain.startswith("."):
            node_label.stylize_before(
                self.get_component_rich_style("directory-tree--hidden", partial=True)
            )
        
        text = Text.assemble(prefix, node_label)
        return text


class ChatHistoryPanel(Container):
    """Material Design chat history panel for browsing saved chat logs."""
    
    DEFAULT_CSS = """
    ChatHistoryPanel {
        width: 35%;
        height: 100vh;
        background: $surface;
        border: solid $primary;
        padding: 1;
        display: none;
        layer: overlay;
        offset: 0 0;
    }
    
    ChatHistoryPanel.visible {
        display: block;
    }
    
    #history-header {
        width: 100%;
        height: 1;
        layout: horizontal;
        align: right middle;
        margin-bottom: 1;
    }
    
    #history-header Button {
        width: auto;
        height: 1;
        min-width: 3;
        padding: 0 1;
        margin: 0 1 0 0;
        border: none;
        background: $panel;
        content-align: center middle;
    }
    
    #history-header Button:hover {
        background: $surface-lighten-1;
    }
    
    #history-header Button:last-child {
        margin: 0;
    }
    
    #history-tree {
        height: 1fr;
        scrollbar-size: 1 1;
    }
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.border_title = f"{EMOJI['chat']} Chat History"
        self.history_path = None  # Will be set from config
        
    def compose(self) -> ComposeResult:
        """Compose the chat history panel."""
        # Header with gear button (config) and close button at top right
        with Horizontal(id="history-header"):
            yield Button(EMOJI['gear'], id="history-config-btn", variant="primary")
            yield Button("✕", id="history-close-btn", variant="primary")
        
        # Get history directory from config
        history_config = load_history_config()
        if history_config:
            location = history_config.get('location', str(Path.cwd()))
            self.history_path = Path(location) / "valbot_tui_chat_history"
        else:
            # Fallback to current directory
            self.history_path = Path.cwd() / "valbot_tui_chat_history"
        
        # Create history directory if it doesn't exist
        self.history_path.mkdir(parents=True, exist_ok=True)
        
        # Directory tree showing only log files
        yield ChatHistoryTree(str(self.history_path), id="history-tree")
    
    async def on_mount(self) -> None:
        """Load the history tree on mount."""
        try:
            tree = self.query_one("#history-tree", ChatHistoryTree)
            tree.reload()
        except Exception:
            pass
    
    def refresh_history(self) -> None:
        """Refresh the history tree to show new entries."""
        try:
            tree = self.query_one("#history-tree", ChatHistoryTree)
            tree.reload()
        except Exception:
            pass
    
    @on(Button.Pressed, "#history-close-btn")
    def action_close_history(self):
        """Close the chat history panel."""
        self.remove_class("visible")
    
    @on(Button.Pressed, "#history-config-btn")
    def action_open_history_config(self):
        """Open the history configuration screen."""
        def on_config_complete(result):
            """Callback after config screen is closed."""
            if result:
                # Configuration was saved, refresh the history tree with new location
                self.refresh_history()
        
        # Push the HistoryConfigScreen
        self.app.push_screen(HistoryConfigScreen(), callback=on_config_complete)


class StatusBar(Container):
    """Modern Material Design status bar with clickable files and model buttons."""
    
    DEFAULT_CSS = """
    StatusBar {
        dock: bottom;
        height: 3;
        background: $surface;
        padding: 1 2;
        layout: horizontal;
        align-vertical: bottom;
    }
    
    StatusBar #files-button {
        dock: left;
        min-width: 5;
        background: transparent !important;
        content-align: center middle;
        text-style: none;
        border: none !important;
    }
    
    StatusBar #right-buttons-container {
        dock: right;
        layout: horizontal;
        height: auto;
        width: auto;
    }
    
    StatusBar #model-button {
        min-width: 10;
        background: transparent !important;
        padding: 0 2;
        content-align: center middle;
        text-style: none;
        border: none !important;
    }
    
    StatusBar #submit-button {
        min-width: 5;
        background: transparent !important;
        padding: 0 2;
        content-align: center middle;
        text-style: none;
        border: none !important;
        margin-left: 1;
    }
    """
    
    model_name = reactive("gpt-4o")
    processing = reactive(False)
    
    def _truncate_model_name(self, model_name: str, max_length: int = 20) -> str:
        """Truncate model name if it exceeds max_length."""
        if len(model_name) <= max_length:
            return model_name
        return model_name[:max_length - 3] + "..."
    
    def compose(self) -> ComposeResult:
        """Compose the status bar with files button (left) and model/submit buttons in container (right)."""
        yield Button(EMOJI['link'], id="files-button")
        with Horizontal(id="right-buttons-container"):
            display_name = self._truncate_model_name(self.model_name)
            yield Button(f"{display_name} {EMOJI['gear']}", id="model-button")
            button_label = EMOJI['stop'] if self.processing else EMOJI['send']
            yield Button(button_label, id="submit-button")
    
    def watch_model_name(self, new_model: str) -> None:
        """Update button label when model changes."""
        try:
            button = self.query_one("#model-button", Button)
            display_name = self._truncate_model_name(new_model)
            button.label = f"{display_name} {EMOJI['gear']}"
        except Exception:
            pass
    
    def watch_processing(self, is_processing: bool) -> None:
        """Update submit button label when processing state changes."""
        try:
            button = self.query_one("#submit-button", Button)
            button.label = EMOJI['stop'] if is_processing else EMOJI['send']
        except Exception:
            pass
    
    @on(Button.Pressed, "#files-button")
    def open_files_panel(self):
        """Open files panel when button is clicked."""
        # Get the main screen and toggle files
        main_screen = self.screen
        if hasattr(main_screen, 'action_toggle_files'):
            main_screen.action_toggle_files()
    
    @on(Button.Pressed, "#model-button")
    async def open_model_picker(self):
        """Open model picker when button is clicked."""
        # Get the main screen and call its action method
        main_screen = self.screen
        if hasattr(main_screen, 'action_change_model'):
            await main_screen.action_change_model()
    
    @on(Button.Pressed, "#submit-button")
    def submit_or_stop(self):
        """Submit input when idle, or stop streaming when processing."""
        main_screen = self.screen
        
        if self.processing:
            # Stop streaming (same as Esc key behavior)
            if hasattr(main_screen, 'action_cancel'):
                main_screen.action_cancel()
        else:
            # Submit user input
            try:
                command_input = main_screen.query_one("#command-input", CommandInput)
                command_input.submit()
            except Exception:
                pass

class CommandInput(TextArea):
    """Modern Material Design multiline input widget with elegant styling."""
    
    DEFAULT_CSS = """
    CommandInput {
        background: rgba(30, 41, 59, 0);
        padding: 1 3;
        height: auto;
        min-height: 3;
        max-height: 10;
    }
    
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(
            placeholder=" Type your message or /command... (↓ for new line, Enter to submit)",
            highlight_cursor_line = False,
            tab_behavior = "indent",
            *args,
            **kwargs
        )
        self.show_line_numbers = False
        self._last_text = ""
    
    def on_mount(self) -> None:
        """Handle mount event."""
        super().on_mount()
        # Store reference in screen for autocomplete
        if hasattr(self.screen, '__class__') and 'MainScreen' in self.screen.__class__.__name__:
            self.screen._command_input = self
    
    def on_text_area_changed(self, event: TextArea.Changed) -> None:
        """Watch for text changes to trigger autocomplete."""
        new_text = self.text
        
        # Check if current word at cursor starts with # - trigger file autocomplete
        # This works for both regular input and commands (e.g., /create_database #chatbot)
        # Check this FIRST and let it determine if file autocomplete should be shown
        file_autocomplete_shown = False
        if hasattr(self.screen, 'on_text_changed_check_paths'):
            file_autocomplete_shown = self.screen.on_text_changed_check_paths(new_text, self.cursor_location)
        
        # Only trigger command autocomplete if text starts with / AND we're not showing file autocomplete
        if new_text.startswith("/") and not file_autocomplete_shown:
            if new_text != self._last_text:
                self._last_text = new_text
                # Notify screen about text change
                if hasattr(self.screen, 'on_command_input_changed'):
                    self.screen.on_command_input_changed(new_text)
        else:
            # Hide command autocomplete if / was removed (but not if file autocomplete is shown)
            if self._last_text.startswith("/") and not file_autocomplete_shown:
                if hasattr(self.screen, 'hide_autocomplete'):
                    self.screen.hide_autocomplete()
            
            self._last_text = new_text
    
    def _on_key(self, event: events.Key) -> None:
        """Handle key events for multiline input."""
        # Check if autocomplete is visible
        autocomplete_visible = False
        autocomplete = None
        if hasattr(self.screen, '_autocomplete_overlay'):
            autocomplete = self.screen._autocomplete_overlay
            if autocomplete and autocomplete.styles.display != "none":
                autocomplete_visible = True
        
        if event.key == "enter":
            # Check if AI is processing - block Enter/submission
            if hasattr(self.screen, 'processing') and self.screen.processing:
                # Show toast notification
                if hasattr(self.screen, 'app'):
                    self.screen.app.notify(
                        "⚠️ AI is busy.\nPress `Esc` to cancel current response.",
                        severity="warning",
                        timeout=2
                    )
                event.prevent_default()
                event.stop()
                return
            
            if autocomplete_visible and autocomplete:
                # Trigger selection from autocomplete
                option_list = autocomplete.get_option_list()
                if option_list and option_list.highlighted is not None:
                    # Get the selected command and post the selection message
                    selected_idx = option_list.highlighted
                    if selected_idx < len(autocomplete._commands):
                        command = autocomplete._commands[selected_idx][0]
                        autocomplete.post_message(AutocompleteSelected(autocomplete, command))
                event.prevent_default()
                return
            
            # Regular Enter submits the message
            self.submit()
            event.prevent_default()
            
        elif event.key == "down":
            if autocomplete_visible and autocomplete:
                # Focus autocomplete for navigation
                option_list = autocomplete.get_option_list()
                if option_list:
                    option_list.focus()
                event.prevent_default()
                return
            
            # Explicitly handle for new lines
            self.insert("\n")
            return
            
        elif event.key == "up":
            if autocomplete_visible and autocomplete:
                # Focus autocomplete for navigation
                option_list = autocomplete.get_option_list()
                if option_list:
                    option_list.focus()
                    # Move to last item
                    if len(option_list._options) > 0:
                        option_list.highlighted = len(option_list._options) - 1
                event.prevent_default()
                return
                
        elif event.key == "tab":
            # Check if autocomplete is visible (command or file autocomplete)
            if autocomplete_visible and autocomplete:
                # Tab selects the highlighted item (trigger Enter behavior)
                option_list = autocomplete.get_option_list()
                if option_list and option_list.highlighted is not None:
                    # Get the selected command/file and post the selection message
                    selected_idx = option_list.highlighted
                    if selected_idx < len(autocomplete._commands):
                        command = autocomplete._commands[selected_idx][0]
                        autocomplete.post_message(AutocompleteSelected(autocomplete, command))
                event.prevent_default()
                return
            
            # Default: indent behavior (let TextArea handle it)
            # Don't prevent default - let the TextArea's tab_behavior work
                
        elif event.key == "escape":
            if autocomplete_visible:
                self.screen.hide_autocomplete()
                event.prevent_default()
                return
            
            # Propagate Esc to parent screen's action_cancel
            if hasattr(self.screen, 'action_cancel'):
                self.screen.action_cancel()
            event.prevent_default()
            event.stop()
    
    def submit(self):
        """Submit the current text content."""
        content = self.text
        # Always submit if content exists (even if blank) - let handle_input decide
        # Post a custom submission event
        self.post_message(TextAreaSubmitted(self, content))
        self.clear()

class TextAreaSubmitted(Message):
    """Custom message for TextArea submission."""
    def __init__(self, text_area, value):
        super().__init__()
        self.text_area = text_area
        self.value = value


class InlinePickerSelected(Message):
    """Custom message for inline picker selection."""
    def __init__(self, picker, value):
        super().__init__()
        self.picker = picker
        self.value = value


class InlinePicker(OptionList):
    """Inline picker that temporarily replaces the textarea for selections."""
    
    BINDINGS = [
        Binding("ctrl+d", "parent_agent_picker", "Agent", show=True, priority=True),
        Binding("ctrl+f", "parent_toggle_files", "Files", show=True, priority=True),
        Binding("ctrl+g", "parent_toggle_history", "History", show=True, priority=True),
        Binding("ctrl+n", "parent_new_chat", "New Chat", show=True, priority=True),
        Binding("ctrl+o", "parent_change_model", "Model", show=True, priority=True),
        Binding("escape", "cancel_picker", "Cancel", show=True, priority=True),
        Binding("ctrl+q", "parent_quit", "Quit", show=True, priority=True),
    ]
    
    can_focus = True
    
    DEFAULT_CSS = """
    InlinePicker {
        height: auto;
        min-height: 10;
        max-height: 20;
        background: $surface;
        border: solid $primary;
        padding: 1;
        overflow-y: auto;
        scrollbar-size: 1 1;
        width: 100%;
    }
    
    InlinePicker .option-list--option {
        background: $panel;
        padding: 0 2;
        height: auto;
        min-height: 3;
        margin: 0 0 1 0;
        width: 100%;
        align-vertical: middle;
    }
    
    InlinePicker .option-list--option-highlighted {
        background: $primary-darken-3 !important;
        text-style: bold;
    }

    """
    
    def __init__(self, title: str, options: list, *args, **kwargs):
        """
        Initialize inline picker.
        
        Args:
            title: The title to show at the top
            options: List of (label, value) tuples or just strings
        """
        super().__init__(*args, **kwargs)
        self.title = title
        self._options_data = options
        # Create a mapping from label to value for tuple options
        self._value_map = {}
        # Set borders (using primary color from theme)
        self.border_title = f"{title}"
        self.border_subtitle = f"{len(options)} items"
        
    def on_mount(self) -> None:
        """Add options when mounted."""
        # Set border using the theme's primary color
        # Access the app's theme to get the actual color value
        try:
            theme = self.app.theme_variables
            primary_color = theme.get("primary", "#989af5")
            self.styles.border = ("solid", primary_color)
        except Exception:
            # Fallback to the valbot-dark primary color
            self.styles.border = ("solid", "#989af5")
        
        # Clear any existing options first
        self.clear_options()
        
        # Build list of options - use Rich Text objects that respect theme colors
        from rich.text import Text
        
        for idx, item in enumerate(self._options_data):
            if isinstance(item, tuple):
                label, value = item
                # Store the mapping from index to value
                self._value_map[idx] = value
                # Create a Rich Text object without hardcoded color - will use theme $text
                text = Text(str(label))
                self.add_option(text)
            else:
                # For simple strings, value is the same as label
                self._value_map[idx] = item
                # Create a Rich Text object without hardcoded color - will use theme $text
                text = Text(str(item))
                self.add_option(text)
        
        # Log for debugging
        self.log(f"InlinePicker mounted with {len(self._options_data)} options")
        self.log(f"Option count after adding: {len(self._options)}")
        
        # Force a refresh after adding all options
        self.refresh(layout=True)
        
        # Auto-highlight the first option after refresh completes
        if len(self._options_data) > 0:
            def highlight_first():
                self.highlighted = 0
                self.refresh()
            self.call_after_refresh(highlight_first)
    
    def on_option_list_option_selected(self, event: OptionList.OptionSelected) -> None:
        """Handle option selection."""
        # Get the actual value from our mapping using the option index
        selected_idx = event.option_index
        selected_value = self._value_map.get(selected_idx, event.option.id)
        self.post_message(InlinePickerSelected(self, selected_value))
    
    # Delegate parent actions to the main screen
    async def action_parent_agent_picker(self) -> None:
        """Delegate to parent screen's agent picker action."""
        if hasattr(self.screen, 'action_agent_picker'):
            await self.screen.action_agent_picker()
    
    async def action_parent_change_model(self) -> None:
        """Delegate to parent screen's change model action."""
        if hasattr(self.screen, 'action_change_model'):
            await self.screen.action_change_model()
    
    def action_parent_toggle_files(self) -> None:
        """Delegate to parent screen's toggle files action."""
        if hasattr(self.screen, 'action_toggle_files'):
            self.screen.action_toggle_files()
    
    def action_parent_toggle_history(self) -> None:
        """Delegate to parent screen's toggle history action."""
        if hasattr(self.screen, 'action_toggle_history'):
            self.screen.action_toggle_history()
    
    def action_parent_new_chat(self) -> None:
        """Delegate to parent screen's new chat action."""
        if hasattr(self.screen, 'action_new_chat'):
            self.screen.action_new_chat()
    
    def action_parent_quit(self) -> None:
        """Delegate to parent app's quit action."""
        self.app.exit()


class AutocompleteSelected(Message):
    """Custom message for autocomplete selection."""
    def __init__(self, picker, command):
        super().__init__()
        self.picker = picker
        self.command = command


class AutocompleteOverlay(Container):
    """Floating autocomplete overlay for command suggestions."""
    
    DEFAULT_CSS = """
    AutocompleteOverlay {
        layer: overlay;
        width: 1fr;
        height: auto;
        max-height: 20;
        dock: bottom;
        offset-y: -8;
        background: $surface;
        padding: 0;
        display: none;
        overflow-y: auto;
        border: solid $primary;
    }
    
    AutocompleteOverlay > #autocomplete-list {
        width: 100%;
        height: auto;
        max-height: 18;
        background: transparent;
        border: none;
        padding: 0 1;
        scrollbar-size: 1 1;
    }
    
    AutocompleteOverlay #autocomplete-list > .option-list--option {
        background: $panel;
        padding: 0 2;
        height: auto;
        min-height: 1;
        margin: 0 0 1 0;
    }
    
    AutocompleteOverlay #autocomplete-list > .option-list--option-highlighted {
        background: $primary-darken-2 !important;
        text-style: bold;
    }
    """
    
    def __init__(self, *args, **kwargs):
        """Initialize autocomplete overlay."""
        super().__init__(*args, **kwargs)
        self._commands = []
        self._option_list = None
        self.border_title = f"{EMOJI['lightbulb']} Suggestions"
    
    def compose(self) -> ComposeResult:
        """Compose the overlay contents."""
        option_list = OptionList(id="autocomplete-list")
        option_list.can_focus = True
        self._option_list = option_list
        yield option_list
    
    def on_key(self, event: events.Key) -> None:
        """Handle key events for the autocomplete overlay."""
        if event.key == "escape":
            self.styles.display = "none"
            # Clear file autocomplete context
            if hasattr(self.screen, '_file_autocomplete_context'):
                self.screen._file_autocomplete_context = None
            # Return focus to input
            if hasattr(self.screen, '_command_input'):
                self.screen._command_input.focus()
            event.prevent_default()
            event.stop()
        elif event.key == "tab":
            # Tab selects the highlighted item
            if self._option_list and self._option_list.highlighted is not None:
                selected_idx = self._option_list.highlighted
                if 0 <= selected_idx < len(self._commands):
                    command = self._commands[selected_idx][0]
                    self.post_message(AutocompleteSelected(self, command))
            event.prevent_default()
            event.stop()
        elif event.key == "enter":
            # Enter also selects the highlighted item
            if self._option_list and self._option_list.highlighted is not None:
                selected_idx = self._option_list.highlighted
                if 0 <= selected_idx < len(self._commands):
                    command = self._commands[selected_idx][0]
                    self.post_message(AutocompleteSelected(self, command))
            event.prevent_default()
            event.stop()
        elif event.key == "up":
            # If at the top, return to input
            if self._option_list and self._option_list.highlighted == 0:
                self.styles.display = "none"
                # Clear file autocomplete context
                if hasattr(self.screen, '_file_autocomplete_context'):
                    self.screen._file_autocomplete_context = None
                if hasattr(self.screen, '_command_input'):
                    self.screen._command_input.focus()
                event.prevent_default()
                event.stop()

    
    def update_suggestions(self, commands: list) -> None:
        """Update the list of command suggestions."""
        self._commands = commands
        
        if not self._option_list:
            return
        
        self._option_list.clear_options()
        
        if not commands:
            self.styles.display = "none"
            return
        
        from rich.text import Text
        
        for command, description in commands:
            # Format: /command - description
            text = Text()
            text.append(command)
            if description:
                text.append(f" - {description}", style="dim")
            self._option_list.add_option(text)
        
        # Auto-highlight the first option
        if len(commands) > 0:
            self._option_list.highlighted = 0
        
        self.styles.display = "block"
        self.refresh(layout=True)
    
    def on_option_list_option_selected(self, event: OptionList.OptionSelected) -> None:
        """Handle option selection."""
        selected_idx = event.option_index
        if 0 <= selected_idx < len(self._commands):
            command = self._commands[selected_idx][0]
            self.post_message(AutocompleteSelected(self, command))
        event.stop()
    
    def get_option_list(self) -> Optional[OptionList]:
        """Get the option list widget."""
        return self._option_list


class ContextChipRemoved(Message):
    """Custom message for context chip removal."""
    def __init__(self, chip, item_type, item_name):
        super().__init__()
        self.chip = chip
        self.item_type = item_type  # 'file' or 'database'
        self.item_name = item_name


class ContextChip(Container):
    """A chip widget displaying a file or database with a remove button."""
    
    DEFAULT_CSS = """
    ContextChip {
        width: auto;
        height: 1;
        background: $panel;
        padding: 0 1;
        margin: 0 1 1 0;
        layout: horizontal;
        align: center middle;
    }
    
    ContextChip .chip-label {
        width: auto;
        height: 1;
        color: $text;
        content-align: left middle;
    }
    
    ContextChip .chip-remove-btn {
        width: 3;
        height: 1;
        min-width: 3;
        padding: 0;
        margin: 0 0 0 1;
        background: transparent;
        color: $text;
        border: none;
        text-style: bold;
        content-align: center middle;
    }
    
    ContextChip .chip-remove-btn:hover {
        color: $text;
        background: transparent;
        border: none;
    }
    """
    
    def __init__(self, item_type: str, item_name: str, *args, **kwargs):
        """
        Initialize a context chip.
        
        Args:
            item_type: Either 'file' or 'database'
            item_name: The name/path of the file or database
        """
        super().__init__(*args, **kwargs)
        self.item_type = item_type
        self.item_name = item_name
    
    def compose(self) -> ComposeResult:
        """Compose the chip with label and remove button."""
        # Determine icon based on type
        if self.item_type == 'file':
            icon = EMOJI['file']
        elif self.item_type == 'folder':
            icon = EMOJI['folder']
        elif self.item_type == 'database':
            icon = EMOJI['database']
        elif self.item_type == 'agent':
            icon = EMOJI['agent']
        else:
            icon = EMOJI['file']  # Default fallback
        
        # Show basename for files, full name for databases/agents
        if self.item_type == 'file':
            display_name = os.path.basename(self.item_name)
        else:
            display_name = self.item_name
        
        # Truncate long names
        if len(display_name) > 40:
            display_name = "..." + display_name[-37:]
        
        yield Static(f"{icon} {display_name}", classes="chip-label")
        yield Button("✕", classes="chip-remove-btn")
    
    @on(Button.Pressed, ".chip-remove-btn")
    def remove_chip(self):
        """Handle remove button click."""
        self.post_message(ContextChipRemoved(self, self.item_type, self.item_name))


class ContextChipBar(Container):
    """Horizontal bar displaying loaded files and databases as chips."""
    
    DEFAULT_CSS = """
    ContextChipBar {
        width: 100%;
        height: auto;
        max-height: 3;
        background: transparent;
        padding: 0;
        layout: horizontal;
        overflow-x: auto;
        scrollbar-size: 1 0;
        dock: none;
        display: none;
    }
    
    ContextChipBar.visible {
        display: block;
        padding: 0 2;
        margin-bottom: 0;
    }
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.loaded_files = []
        self.loaded_databases = []  # Changed from single to list
        self.active_agent = None  # Track currently active agent
    
    def add_file(self, file_path: str):
        """Add a file chip to the bar."""
        if file_path not in self.loaded_files:
            self.loaded_files.append(file_path)
            chip = ContextChip("file", file_path)
            self.mount(chip)
            self.add_class("visible")
    
    def add_database(self, db_path: str):
        """Add a database chip to the bar."""
        # Add to list if not already present
        if db_path not in self.loaded_databases:
            self.loaded_databases.append(db_path)
            chip = ContextChip("database", db_path)
            self.mount(chip)
            self.add_class("visible")
    
    def add_agent(self, agent_name: str):
        """Add an agent chip to the bar."""
        # Remove any existing agent chip first (only one agent at a time)
        if self.active_agent:
            self.remove_agent(self.active_agent)
        
        self.active_agent = agent_name
        chip = ContextChip("agent", agent_name)
        self.mount(chip)
        self.add_class("visible")
    
    def remove_file(self, file_path: str):
        """Remove a file chip from the bar."""
        if file_path in self.loaded_files:
            self.loaded_files.remove(file_path)
        
        # Find and remove the chip widget
        for chip in self.query(ContextChip):
            if chip.item_type == "file" and chip.item_name == file_path:
                chip.remove()
                break
        
        # Hide the bar if no chips remain
        if not self.loaded_files and not self.loaded_databases and not self.active_agent:
            self.remove_class("visible")
    
    def remove_database(self, db_path: str):
        """Remove a database chip from the bar."""
        if db_path in self.loaded_databases:
            self.loaded_databases.remove(db_path)
        
        # Find and remove the chip widget
        for chip in self.query(ContextChip):
            if chip.item_type == "database" and chip.item_name == db_path:
                chip.remove()
                break
        
        # Hide the bar if no chips remain
        if not self.loaded_files and not self.loaded_databases and not self.active_agent:
            self.remove_class("visible")
    
    def remove_agent(self, agent_name: str = None):
        """Remove the agent chip from the bar."""
        # If no agent_name specified, remove current active agent
        if agent_name is None:
            agent_name = self.active_agent
        
        if agent_name and agent_name == self.active_agent:
            self.active_agent = None
        
        # Find and remove the chip widget
        for chip in self.query(ContextChip):
            if chip.item_type == "agent" and (agent_name is None or chip.item_name == agent_name):
                chip.remove()
                break
        
        # Hide the bar if no chips remain
        if not self.loaded_files and not self.loaded_databases and not self.active_agent:
            self.remove_class("visible")
    
    def clear_all(self):
        """Remove all chips from the bar."""
        # Clear the lists
        self.loaded_files.clear()
        self.loaded_databases.clear()
        self.active_agent = None
        
        # Remove all chip widgets
        for chip in list(self.query(ContextChip)):
            chip.remove()
        
        # Hide the bar
        self.remove_class("visible")


# ============================================================================
# History Configuration Screen
# ============================================================================

class HistoryConfigScreen(Screen):
    """Full-screen overlay for configuring chat history settings on first run."""
    
    BINDINGS = [
        Binding("escape", "cancel", "Cancel", show=False),
    ]
    
    CSS = """
    HistoryConfigScreen {
        align: center middle;
        background: $background 90%;
    }
    
    #config-dialog {
        width: 80;
        height: auto;
        max-width: 100;
        background: $surface;
        border: solid $primary;
        padding: 2 4;
    }
    
    #config-title {
        width: 100%;
        text-align: center;
        text-style: bold;
        color: $primary;
        margin-bottom: 2;
    }
    
    #config-description {
        width: 100%;
        text-align: center;
        color: $text-muted;
        margin-bottom: 3;
    }
    
    .config-label {
        width: 100%;
        color: $text;
        text-style: bold;
        margin-top: 2;
        margin-bottom: 1;
    }
    
    .config-input {
        width: 100%;
        height: 1;
        margin-bottom: 1;
        padding: 0 1;
        background: $panel;
        border: none;
    }
    
    .config-help {
        width: 100%;
        color: $text-muted;
        margin-top: 0;
        margin-bottom: 2;
    }
    
    #button-container {
        width: 100%;
        height: 1;
        layout: horizontal;
        margin-top: 3;
        margin-bottom: 1;
    }
    
    #button-container Button {
        width: 2fr;
        height: 1;
        min-width: 0;
        padding: 0;
        margin: 0 1 0 0;
        border: none;
        background: $panel;
        content-align: center middle;
    }
    
    #button-container Button:hover {
        background: $surface-lighten-1;
    }
    
    #button-container Button:last-child {
        margin: 0;
    }
    
    #button-container Button#save-btn {
        background: $primary;
    }
    
    #button-container Button#save-btn:hover {
        background: $primary-lighten-1;
    }
    """
    
    def __init__(self):
        super().__init__()
        self.default_location = str(Path.cwd())
        self.default_max_size = 5
    
    def compose(self) -> ComposeResult:
        """Compose the configuration dialog."""
        with Container(id="config-dialog"):
            yield Static(f"{EMOJI['gear']} Chat History Configuration", id="config-title")
            yield Static("Configure where to save chat history and maximum storage size", id="config-description")
            
            yield Static("History Location:", classes="config-label")
            yield Input(
                value=self.default_location,
                placeholder="Enter path to save chat history",
                id="location-input",
                classes="config-input"
            )
            yield Static("Path where chat history will be saved", classes="config-help")
            
            yield Static("Maximum Size (GB):", classes="config-label")
            yield Input(
                value=str(self.default_max_size),
                placeholder="Enter maximum size in GB",
                id="maxsize-input",
                classes="config-input"
            )
            yield Static("Maximum storage space for chat history (default: 5 GB)", classes="config-help")
            
            with Container(id="button-container"):
                yield Button("Save", id="save-btn", variant="primary")
                yield Button("Cancel", id="cancel-btn", variant="default")
    
    def on_mount(self) -> None:
        """Focus the location input when mounted."""
        self.query_one("#location-input", Input).focus()
    
    @on(Button.Pressed, "#save-btn")
    async def save_configuration(self) -> None:
        """Save the configuration and dismiss the screen."""
        try:
            location_input = self.query_one("#location-input", Input)
            maxsize_input = self.query_one("#maxsize-input", Input)
            
            location = location_input.value.strip()
            if not location:
                location = self.default_location
            
            # Validate and parse max size
            try:
                max_size = int(maxsize_input.value.strip())
                if max_size < 1:
                    max_size = self.default_max_size
            except ValueError:
                max_size = self.default_max_size
            
            # Save configuration
            save_history_config(location, max_size)
            
            # Dismiss the screen
            self.dismiss(True)
            
        except Exception as e:
            # On error, use defaults
            save_history_config(self.default_location, self.default_max_size)
            self.dismiss(True)
    
    @on(Button.Pressed, "#cancel-btn")
    async def cancel_configuration(self) -> None:
        """Cancel and use default configuration."""
        save_history_config(self.default_location, self.default_max_size)
        self.dismiss(False)
    
    def action_cancel(self) -> None:
        """Handle escape key - use default configuration."""
        save_history_config(self.default_location, self.default_max_size)
        self.dismiss(False)


class SidebarButtonPressed(Message):
    """Message sent when a sidebar button is pressed."""
    def __init__(self, button_id: str):
        super().__init__()
        self.button_id = button_id


class Sidebar(Container):
    """Collapsible sidebar with icon buttons for main actions."""
    
    DEFAULT_CSS = """
    Sidebar {
        width: auto;
        height: 100%;
        background: $background;
        border-right: solid $panel;
        padding: 0;
        margin: 0 !important;
    }
    
    Sidebar.expanded {
        width: auto;
    }
    
    Sidebar Button {
        width: auto;
        height: 3;
        min-height: 3;
        min-width: 0;
        margin: 0;
        border: none !important;
        background: transparent !important;
        content-align: center middle;
        text-align: center;
        padding: 0;
    }
    
    Sidebar.collapsed Button .button-label-text {
        display: none;
    }
    
    Sidebar.expanded Button {
        content-align: left middle;
        text-align: left;
        padding: 0 2;
    }
    """
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.expanded = False
        self.add_class("collapsed")
    
    def compose(self) -> ComposeResult:
        """Compose the sidebar buttons."""
        yield Button(">", id="btn-expand", variant="default")
        yield Button("+", id="btn-new-chat", variant="default")
        yield Button("◈", id="btn-agent", variant="default")
        yield Button("◷", id="btn-history", variant="default")
    
    def toggle_expand(self) -> None:
        """Toggle sidebar expansion."""
        self.expanded = not self.expanded
        
        if self.expanded:
            self.remove_class("collapsed")
            self.add_class("expanded")
            # Update button labels to show text
            try:
                self.query_one("#btn-expand", Button).label = "< Collapse"
                self.query_one("#btn-new-chat", Button).label = f"+ New Chat"
                self.query_one("#btn-agent", Button).label = f"◈ Agent"
                self.query_one("#btn-history", Button).label = f"◷ History"
            except Exception:
                pass
        else:
            self.remove_class("expanded")
            self.add_class("collapsed")
            # Update button labels to show only icons
            try:
                self.query_one("#btn-expand", Button).label = ">"
                self.query_one("#btn-new-chat", Button).label = "+"
                self.query_one("#btn-agent", Button).label = "◈"
                self.query_one("#btn-history", Button).label = "◷"
            except Exception:
                pass
        
        self.refresh(layout=True)
    
    @on(Button.Pressed, "#btn-expand")
    def handle_expand_button(self) -> None:
        """Handle expand/collapse button press."""
        self.toggle_expand()
    
    @on(Button.Pressed, "#btn-new-chat")
    def handle_new_chat_button(self) -> None:
        """Handle new chat button press."""
        self.post_message(SidebarButtonPressed("new-chat"))
    
    @on(Button.Pressed, "#btn-agent")
    def handle_agent_button(self) -> None:
        """Handle agent button press."""
        self.post_message(SidebarButtonPressed("agent"))
    
    @on(Button.Pressed, "#btn-history")
    def handle_history_button(self) -> None:
        """Handle history button press."""
        self.post_message(SidebarButtonPressed("history"))


class MainScreen(Screen):
    """Main application screen with modern Material Design layout."""
    
    BINDINGS = [
        Binding("ctrl+d", "agent_picker", "Agent", priority=True),
        Binding("ctrl+f", "toggle_files", "Files", priority=True),
        Binding("ctrl+g", "toggle_history", "History", priority=True),
        Binding("ctrl+n", "new_chat", "New Chat", priority=True),
        Binding("ctrl+o", "change_model", "Model", priority=True),
        Binding("escape", "cancel", "Cancel", priority=True),
        Binding("ctrl+q", "quit", "Quit", priority=True),
    ]
    
    CSS = """
    MainScreen {
        background: $background;
        layers: base overlay;
        layout: horizontal;
    }
    
    #sidebar {
        height: 100%;
    }
    
    #main-content-area {
        layout: vertical;
        width: 1fr;
        height: 100%;
    }
    
    #main-container {
        layout: vertical;
        height: 1fr;
        width: 100%;
        background: $background;
    }
    
    #chat-container {
        height: 1fr;
        width: 100%;
    }
    
    #file-panel {
        dock: right;
        layer: overlay;
    }
    
    #history-panel {
        dock: left;
        layer: overlay;
    }
    
    #input-container {
        height: auto;
        width: 100%;
        padding: 1 2;
        background: $surface;
    }
    
    #command-input {
        width: 100%;
        border: none;
    }
    
    /* Markdown Styling - Let Textual handle natural styling */
    Markdown {
        background: transparent;
        padding: 0;
        margin: 0;
    }
    
    /* Scrollbar Styling */
    ScrollableContainer > .scrollbar {
        background: $surface;
        color: $primary;
    }
    
    ScrollableContainer > .scrollbar:hover {
        background: $panel;
        color: $accent;
    }
    
    /* Remove focus borders */
    *:focus {
        border: none !important;
    }
    
    Input:focus {
        border: none !important;
    }
    
    Static:focus {
        border: none !important;
    }
    
    Container:focus {
        border: none !important;
    }
    
    """
    
    def __init__(self, config_manager: ConfigManager, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.config_manager = config_manager
        self.chatbot = None
        self.processing = False
        self.show_files = False
        self.show_history = False
        self._active_picker_type = None  # Track which picker is currently open ('agent', 'model', etc.)
        self._cancel_streaming = False  # Flag to cancel ongoing streaming
        self._cancel_agent = False  # Flag to cancel ongoing agent execution
        self._file_autocomplete_context = None  # Track file path autocomplete context
        
        # Tool execution task tracking for cancellation
        self._current_tool_task = None  # Current async task when using tools
        self._current_tool_loop = None  # Current event loop for tool execution
        
        # Chat history tracking
        self._current_session_file = None  # Path to current session log file
        self._session_start_time = None  # Timestamp when session started
        self._session_initialized = False  # Flag to track if session file has been created
        self._last_autosave_content = ""  # Track last saved content to avoid duplicate saves
        self._first_user_input = None  # Track first user input for filename generation
        
    def compose(self) -> ComposeResult:
        """Compose the modern Material Design layout."""
        yield Header(show_clock=True, icon="☰")
        
        # Sidebar (docked to left)
        yield Sidebar(id="sidebar")
        
        # File explorer overlay (appears on top when toggled)
        yield FileExplorerPanel(id="file-panel")
        
        # Chat history overlay (appears on top when toggled)
        yield ChatHistoryPanel(id="history-panel")
        
        # Main content area (everything to the right of sidebar)
        with Container(id="main-content-area"):
            with Container(id="main-container"):
                # Main chat area (scrollable)
                with Vertical(id="chat-container"):
                    yield ChatPanel(id="chat-panel")
            
            # Input area at bottom (fixed)
            with Vertical(id="input-container"):
                # Context chip bar (shows loaded files/databases)
                yield ContextChipBar(id="context-chip-bar")
                yield CommandInput(id="command-input")
                yield StatusBar(id="status-bar")
        
        # Autocomplete overlay (on its own layer)
        yield AutocompleteOverlay(id="autocomplete-overlay")
    
    def _initialize_session_file(self, first_message: str = "") -> None:
        """Initialize a new session log file with timestamp-based naming.
        
        Args:
            first_message: The first user message, used to generate a descriptive filename
        """
        if self._session_initialized:
            return
        
        try:
            # Get history directory
            history_dir = ensure_history_directory()
            
            # Check and cleanup old history files if exceeding max_size
            history_config = load_history_config()
            if history_config:
                max_size_gb = history_config.get("max_size", 5)  # Default to 5 GB
                cleanup_old_history(history_dir, max_size_gb)
            
            # Create filename based on first message (first 20 chars) and timestamp
            self._session_start_time = datetime.now()
            timestamp = self._session_start_time.strftime("%Y%m%d_%H%M%S")
            
            # Generate descriptive prefix from first message
            if first_message:
                # Take first 30 characters, replace spaces with underscores, remove special chars
                prefix = first_message[:30].replace(' ', '_')
                # Keep only alphanumeric and underscores
                prefix = ''.join(c for c in prefix if c.isalnum() or c == '_')
                # Remove consecutive underscores
                while '__' in prefix:
                    prefix = prefix.replace('__', '_')
                # Remove leading/trailing underscores
                prefix = prefix.strip('_')
                # Ensure prefix is not empty
                if prefix:
                    filename = f"{prefix}_{timestamp}.log"
                else:
                    filename = f"chat_{timestamp}.log"
            else:
                filename = f"chat_{timestamp}.log"
            
            self._current_session_file = history_dir / filename
            
            # Create the file with a header
            with open(self._current_session_file, 'w', encoding='utf-8') as f:
                f.write(f"ValBot TUI Chat Log\n")
                f.write(f"Session Started: {self._session_start_time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("=" * 80 + "\n\n")
            
            self._session_initialized = True
            
        except Exception as e:
            # Silently fail if we can't create the session file
            pass
    
    async def _autosave_message(self, role: str, content: str, agent_name: Optional[str] = None) -> None:
        """
        Autosave a message to the current session file.
        This runs in the background and doesn't block the UI.
        """
        try:
            # Initialize session file if this is the first message
            if not self._session_initialized:
                # Store first user input for filename generation
                if role == "user" and not self._first_user_input:
                    self._first_user_input = content
                self._initialize_session_file(self._first_user_input or "")
            
            if not self._current_session_file:
                return
            
            # Prepare the message entry
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            role_label = role.upper()
            
            if role == "assistant" and agent_name:
                role_label += f" (Agent: {agent_name})"
            
            # Append to the file asynchronously
            await asyncio.to_thread(self._write_to_session_file, timestamp, role_label, content)
            
        except Exception as e:
            # Silently fail to not disrupt the user experience
            pass
    
    def _write_to_session_file(self, timestamp: str, role_label: str, content: str) -> None:
        """Synchronous write operation to be run in a thread."""
        try:
            with open(self._current_session_file, 'a', encoding='utf-8') as f:
                f.write(f"[{timestamp}] {role_label}:\n")
                f.write(f"{content}\n")
                f.write("-" * 80 + "\n\n")
        except Exception:
            pass
        
    async def on_mount(self) -> None:
        """Initialize the application after mounting."""
        await self.initialize_chatbot()
        
        # Store reference to autocomplete overlay
        self._autocomplete_overlay = self.query_one("#autocomplete-overlay", AutocompleteOverlay)
        
        # Store reference to context chip bar
        self._context_chip_bar = self.query_one("#context-chip-bar", ContextChipBar)
        
        self.query_one("#command-input", CommandInput).focus()
        
        # Initialize CommandManager with TUI console wrapper
        if self.chatbot:
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            tui_console = TUIConsoleWrapper(chat_panel, self.app)
            self.chatbot.console = tui_console
            self.chatbot.command_manager = CommandManager(
                self.chatbot, 
                self.chatbot.plugin_manager,
                tui_console
            )
        
        # Display welcome message
        self.show_welcome_message()
    
    def on_context_chip_removed(self, message: ContextChipRemoved) -> None:
        """Handle removal of a context chip."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        if message.item_type == "file":
            # Remove file from context manager
            if self.chatbot and self.chatbot.context_manager:
                self.chatbot.context_manager.remove_file_from_context(message.item_name)
            
            # Remove from chip bar
            self._context_chip_bar.remove_file(message.item_name)
            
            # Show notification
            chat_panel.add_message("system", 
                f"{EMOJI['checkmark']} Removed file from context: `{message.item_name}`")
        
        elif message.item_type == "database":
            # Unload the specific database from the active RAG knowledge bases
            if hasattr(self, '_active_rag_kb') and self._active_rag_kb is not None:
                # _active_rag_kb is now a dict, remove this specific database
                if isinstance(self._active_rag_kb, dict) and message.item_name in self._active_rag_kb:
                    del self._active_rag_kb[message.item_name]
                    # If dict is now empty, set to None
                    if not self._active_rag_kb:
                        self._active_rag_kb = None
            
            # Remove from chip bar
            self._context_chip_bar.remove_database(message.item_name)
            
            # Show notification
            chat_panel.add_message("system", 
                f"{EMOJI['checkmark']} Unloaded database: `{message.item_name}`")
        
        elif message.item_type == "agent":
            # Cancel the running agent
            if hasattr(self, '_agent_input_state') and self._agent_input_state:
                self._agent_input_state['cancelled'] = True
                # Wake up any waiting prompts
                if self._agent_input_state.get('waiting_for_input'):
                    self._agent_input_state['ready'].set()
            
            # Remove from chip bar
            self._context_chip_bar.remove_agent(message.item_name)
            
            # Show notification
            chat_panel.add_message("system", 
                f"{EMOJI['cross']} Cancelling agent: **{message.item_name}**")
        
        # Auto-focus the input text area after removing a context chip
        try:
            command_input = self.query_one("#command-input", CommandInput)
            command_input.focus()
        except Exception:
            pass
    
    def on_sidebar_button_pressed(self, message: SidebarButtonPressed) -> None:
        """Handle sidebar button presses."""
        if message.button_id == "new-chat":
            asyncio.create_task(self.action_new_chat())
        elif message.button_id == "agent":
            asyncio.create_task(self.action_agent_picker())
        elif message.button_id == "files":
            self.action_toggle_files()
        elif message.button_id == "history":
            self.action_toggle_history()
    
    def show_welcome_message(self):
        """Display the welcome message."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        welcome = f"""# ░▒▓ Welcome to ValBot TUI ▓▒░

### {EMOJI['clipboard']} Available Commands

Type `/help` for complete list of commands, or try these:

- `/` - Show all commands
- `/agent` - Run an agent workflow
- `/model` - Change AI model
- `/create_database #<file_path>` - create RAG database from documents (.pdf, .docx, .csv, .py, etc.)
- `/terminal <cmd>` - Run shell commands


### {EMOJI['info']} Tips
- Type `#` to reference local files directly in chat
- Press `TAB` to autocomplete file paths
- Use key-bindings for quick access. E.g.: `ctrl + a` for Agent
- Change themes from palette


### {EMOJI['lightbulb']} Getting Started

Type your message and press **Enter** to chat with ValBot!

Try:
> Help me write a Python function to parse JSON

> Summarise *<file_path>*
"""
        # Use 'assistant' role but with plain styling (no border/background)
        msg = ChatMessage("assistant", welcome)
        msg.remove_class("assistant-message")  # Remove default styling
        msg.add_class("welcome-message")  # Add plain styling
        chat_panel.messages.append(msg)
        chat_panel.mount(msg)
        chat_panel.scroll_end(animate=True)
        
    async def initialize_chatbot(self):
        """Initialize the ChatBot instance."""
        try:
            agent_config = self.config_manager.get_setting('agent_model_config')
            agent_model = initialize_agent_model(
                endpoint=agent_config['default_endpoint'], 
                model_name=agent_config['default_model']
            )
            plugin_manager = PluginManager(self.config_manager, model=agent_model)
            
            self.chatbot = ChatBot(
                agent_model=agent_model,
                config_manager=self.config_manager,
                plugin_manager=plugin_manager,
                is_cli=False
            )
            
            # Add system prompt to conversation history (like CLI does)
            system_prompt = self.config_manager.get_setting(
                "chat_model_config.system_prompt", 
                "You are a helpful assistant. Prioritize markdown format and code blocks when applicable."
            )
            self.chatbot.context_manager.conversation_history.append({
                "role": "system", 
                "content": system_prompt
            })
            
            # Update status bar
            status_bar = self.query_one("#status-bar", StatusBar)
            status_bar.model_name = self.config_manager.get_setting("chat_model_config.default_model", "gpt-4o")
            
        except Exception as e:
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.add_message("error", f"""### {EMOJI['cross']} Initialization Error

Failed to initialize chatbot:

```
{str(e)}
```

Please check your configuration and try again.
""")
    
    def get_available_commands(self) -> list:
        """Get list of all available commands with descriptions."""
        commands = []
        
        # Built-in TUI commands
        builtin_commands = [
            ("/help", "Show help information"),
            ("/quit", "Exit the application"),
            ("/new", "Start a new conversation"),
            ("/agent", "Select and run an agent workflow"),
            ("/model", "Change the AI model"),
            ("/file", "Display file contents"),
            ("/terminal", "Execute shell commands"),
            ("/multi", "Multi-line input via system editor"),
            ("/prompts", "Show custom prompts"),
            ("/commands", "Show all available commands"),
            ("/settings", "Show settings information"),
            ("/reload", "Reinitialize chatbot"),
            ("/update", "Check for and install updates from GitHub"),
            ("/add_agent", "Add agent information"),
            ("/add_tool", "Add tool information"),
            ("/create_database", "Create a RAG database from documents"),
            ("/load_database", "Load and query an existing database"),
        ]
        commands.extend(builtin_commands)
        
        # Add commands from CommandManager if available
        if self.chatbot and hasattr(self.chatbot, 'command_manager'):
            cmd_mgr = self.chatbot.command_manager
            
            # Add registered commands
            if hasattr(cmd_mgr, 'command_registry'):
                for cmd_name, cmd_info in cmd_mgr.command_registry.items():
                    if not any(c[0] == cmd_name for c in commands):
                        commands.append((cmd_name, cmd_info.get('description', '')))
            
            # Add custom prompts
            if hasattr(cmd_mgr, 'prompts'):
                for prompt_name, prompt_obj in cmd_mgr.prompts.items():
                    if not any(c[0] == prompt_name for c in commands):
                        commands.append((prompt_name, prompt_obj.description))
        
        return commands
    
    def filter_commands(self, input_text: str) -> list:
        """Filter commands based on input text."""
        if not input_text.startswith("/"):
            return []
        
        # Get the command part (without /)
        query = input_text[1:].lower()
        
        # Get all available commands
        all_commands = self.get_available_commands()
        
        # If query is empty, show all commands
        if not query:
            return all_commands
        
        # Filter commands that start with or contain the query
        filtered = []
        for cmd, desc in all_commands:
            cmd_name = cmd[1:] if cmd.startswith("/") else cmd  # Remove leading /
            if cmd_name.lower().startswith(query):
                filtered.append((cmd, desc))
        
        # If no exact prefix matches, try contains matching
        if not filtered:
            for cmd, desc in all_commands:
                cmd_name = cmd[1:] if cmd.startswith("/") else cmd
                if query in cmd_name.lower():
                    filtered.append((cmd, desc))
        
        return filtered
    
    def on_command_input_changed(self, text: str) -> None:
        """Handle command input text changes to show/update autocomplete."""
        if not text.startswith("/"):
            self.hide_autocomplete()
            return
        
        # Don't show command autocomplete if we're currently in file autocomplete mode
        if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
            return
        
        # Filter commands based on input
        filtered_commands = self.filter_commands(text)
        
        if not filtered_commands:
            self.hide_autocomplete()
            return
        
        # Show or update autocomplete overlay
        self.show_autocomplete(filtered_commands)
    
    def on_text_changed_check_paths(self, text: str, cursor_pos: tuple) -> bool:
        """Check if current word is a partial path and show autocomplete.
        
        Returns:
            True if file autocomplete is shown, False otherwise
        """
        if not text.strip():
            self.hide_autocomplete()
            return False
        
        # Get the word at cursor (returns empty string if not starting with #)
        word, start_col, end_col = self.get_word_at_cursor(text, cursor_pos)
        
        # Check if the character at start_col is # (meaning we're in a file reference)
        row, col = cursor_pos
        lines = text.split('\n')
        if row < len(lines):
            line = lines[row]
            # Check if we're at a # position
            if start_col < len(line) and line[start_col] == '#':
                # We're at a file reference position
                # word contains the text after #, which could be empty (just #) or have content (#src)
                pass  # Continue to show autocomplete
            else:
                # Not at a # position, hide autocomplete only if we had file autocomplete before
                if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
                    self.hide_autocomplete()
                return False
        else:
            # Invalid row position
            if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
                self.hide_autocomplete()
            return False
        
        # Get matching file paths
        matches = self.get_file_path_matches(word)
        
        if not matches:
            # Hide autocomplete if no matches
            if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
                self.hide_autocomplete()
            return False
        
        # Format matches for display
        formatted_matches = []
        for display_name, completion, is_dir in matches:
            desc = f"{EMOJI['folder']} Directory" if is_dir else f"{EMOJI['file']} File"
            # Use display_name for showing in list (has / for dirs), completion for actual path
            formatted_matches.append((display_name, desc))
        
        # Store the current word info for completion
        self._file_autocomplete_context = {
            'word': word,
            'start_col': start_col,
            'end_col': end_col,
            'cursor_row': cursor_pos[0],
            'matches': matches  # Store full match info including is_dir flag
        }
        
        # Show autocomplete with file paths
        self.show_autocomplete(formatted_matches)
        return True
    
    def show_autocomplete(self, commands: list) -> None:
        """Show autocomplete overlay with filtered commands."""
        if not hasattr(self, '_autocomplete_overlay') or self._autocomplete_overlay is None:
            return
        
        # Update suggestions
        self._autocomplete_overlay.update_suggestions(commands)
    
    def hide_autocomplete(self) -> None:
        """Hide autocomplete overlay."""
        if hasattr(self, '_autocomplete_overlay') and self._autocomplete_overlay:
            self._autocomplete_overlay.styles.display = "none"
            # Clear file autocomplete context when hiding
            self._file_autocomplete_context = None
            # Return focus to input
            command_input = self.query_one("#command-input", CommandInput)
            command_input.focus()
    
    def get_word_at_cursor(self, text: str, cursor_pos: tuple) -> tuple:
        """
        Get the word/path at cursor position.
        For file path autocomplete, only returns words that start with #.
        
        Returns:
            (word, start_col, end_col) tuple
        """
        row, col = cursor_pos
        lines = text.split('\n')
        
        if row >= len(lines):
            return ("", col, col)
        
        line = lines[row]
        if col > len(line):
            col = len(line)
        
        # Find start of word (go backwards until whitespace or start)
        start = col
        while start > 0 and not line[start - 1].isspace():
            start -= 1
        
        # Find end of word (go forwards until whitespace or end)
        end = col
        while end < len(line) and not line[end].isspace():
            end += 1
        
        word = line[start:end]
        
        # Only return the word if it starts with # (for file path autocomplete)
        # Strip the # from the returned word, but remember the position includes it
        if word.startswith('#'):
            # Return word without the # prefix, but keep start position at the #
            return (word[1:], start, end)
        
        # If word doesn't start with #, return empty (no autocomplete)
        return ("", start, end)
    
    def get_file_path_matches(self, partial_path: str, max_results: int = 20) -> list:
        """
        Get file/folder paths that match the partial path.
        If partial_path is empty, returns all files/folders in current directory.
        
        Returns:
            List of (display_name, full_path, is_dir) tuples
        """
        import os
        from pathlib import Path
        
        # Expand user home directory if provided
        if partial_path:
            partial_path = os.path.expanduser(partial_path)
        
        # Handle both absolute and relative paths
        if partial_path and os.path.isabs(partial_path):
            base_path = os.path.dirname(partial_path)
            prefix = os.path.basename(partial_path)
        else:
            # Relative path or empty - search from current working directory
            base_path = os.getcwd()
            if partial_path and os.path.dirname(partial_path):
                base_path = os.path.join(base_path, os.path.dirname(partial_path))
            prefix = os.path.basename(partial_path) if partial_path else ""
        
        matches = []
        
        try:
            # If base path doesn't exist, try parent
            if not os.path.exists(base_path):
                return []
            
            # List directory contents
            entries = os.listdir(base_path)
            
            for entry in entries:
                # Filter by prefix
                if prefix and not entry.lower().startswith(prefix.lower()):
                    continue
                
                full_path = os.path.join(base_path, entry)
                is_dir = os.path.isdir(full_path)
                
                # Create display name (add OS-appropriate separator for directories)
                display_name = entry + (os.sep if is_dir else "")
                
                # Create the completion path
                if os.path.dirname(partial_path):
                    completion = os.path.join(os.path.dirname(partial_path), entry)
                else:
                    completion = entry
                
                matches.append((display_name, completion, is_dir))
                
                if len(matches) >= max_results:
                    break
            
            # Sort: directories first, then files, both alphabetically
            matches.sort(key=lambda x: (not x[2], x[0].lower()))
            
        except (PermissionError, OSError):
            # Can't read directory, return empty
            pass
        
        return matches
    
    def try_file_path_autocomplete(self, text: str, cursor_pos: tuple) -> bool:
        """
        Try to show file path autocomplete or trigger if already showing.
        Only works if the word at cursor starts with #.
        
        Returns:
            True if autocomplete was shown/triggered, False otherwise
        """
        # Check if autocomplete is already showing file paths
        if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
            if hasattr(self, '_autocomplete_overlay') and self._autocomplete_overlay.styles.display != "none":
                # Autocomplete is showing - focus it for navigation
                option_list = self._autocomplete_overlay.get_option_list()
                if option_list:
                    option_list.focus()
                return True
        
        # Get the word at cursor (only returns non-empty if starts with #)
        word, start_col, end_col = self.get_word_at_cursor(text, cursor_pos)
        
        # If word is empty, it means it doesn't start with # - no autocomplete
        if not word:
            return False
        
        # Get matching file paths
        matches = self.get_file_path_matches(word)
        
        if not matches:
            return False
        
        # If only one exact match, complete it directly
        if len(matches) == 1:
            command_input = self.query_one("#command-input", CommandInput)
            lines = text.split('\n')
            row, col = cursor_pos
            
            if row < len(lines):
                line = lines[row]
                # Get the match info
                display_name, completion, is_dir = matches[0]
                
                # Add appropriate suffix: separator for directories, space for files
                if is_dir:
                    completion = completion + os.sep
                else:
                    completion = completion + " "
                
                # Replace the word with the completion, preserving the # prefix
                # start_col points to the # character, so we keep it
                new_line = line[:start_col] + "#" + completion + line[end_col:]
                lines[row] = new_line
                command_input.text = '\n'.join(lines)
                
                # Move cursor to end of completed word (+ 1 for the # we added)
                new_cursor_col = start_col + len(completion) + 1
                command_input.move_cursor((row, new_cursor_col))
            
            return True
        
        # Multiple matches - trigger the real-time check to show autocomplete
        self.on_text_changed_check_paths(text, cursor_pos)
        
        return True
    
    @on(AutocompleteSelected)
    def handle_autocomplete_selected(self, event: AutocompleteSelected) -> None:
        """Handle autocomplete selection."""
        import os
        selection = event.command
        command_input = self.query_one("#command-input", CommandInput)
        
        # Check if this is a file path completion
        if hasattr(self, '_file_autocomplete_context') and self._file_autocomplete_context:
            context = self._file_autocomplete_context
            text = command_input.text
            lines = text.split('\n')
            row = context['cursor_row']
            
            # Find the selected match to check if it's a directory
            # selection now contains display_name (with / for dirs), find the actual completion path
            is_directory = False
            actual_completion = selection
            if 'matches' in context:
                for display_name, completion, is_dir in context['matches']:
                    if display_name == selection:
                        is_directory = is_dir
                        actual_completion = completion
                        break
            
            # Add appropriate suffix for directories vs files
            if is_directory:
                # Use OS-appropriate separator for directories
                # Don't add if already ends with a separator
                if not actual_completion.endswith(os.sep):
                    actual_completion = actual_completion + os.sep
            else:
                # Add space after file selection
                actual_completion = actual_completion + " "
            
            if row < len(lines):
                line = lines[row]
                # Replace the word with the selected path, preserving the # prefix
                # start_col points to the # character, so we keep it
                new_line = line[:context['start_col']] + "#" + actual_completion + line[context['end_col']:]
                lines[row] = new_line
                command_input.text = '\n'.join(lines)
                
                # Move cursor to end of completed path (+ 1 for the # we added)
                new_cursor_col = context['start_col'] + len(actual_completion) + 1
                command_input.move_cursor((row, new_cursor_col))
            
            # Always close the autocomplete and clear context after selection
            # User must press Tab again to trigger autocomplete
            self.hide_autocomplete()
            self._file_autocomplete_context = None
            
        else:
            # Command completion - replace entire input with command + space
            command_input.text = selection + " "
            command_input.move_cursor_relative(columns=len(selection) + 1)
            
            # Hide autocomplete
            self.hide_autocomplete()
        
        # Return focus to input
        command_input.focus()
        command_input.focus()
    
    def _update_footer_bindings(self, picker_open: bool):
        """Placeholder for footer binding updates (footer removed but bindings still active).
        
        Args:
            picker_open: True if picker is open, False otherwise
        """
        # Footer widget removed, but BINDINGS are still active
        # Key bindings work automatically without the footer widget
        pass
    
    async def show_inline_picker(self, title: str, options: list, picker_type: str = None):
        """Show an inline picker that replaces the textarea.
        
        Args:
            title: The title for the picker
            options: List of options to display
            picker_type: Type identifier for the picker ('agent', 'model', etc.)
        """
        # Check if same picker type is already open - if so, close it instead
        if picker_type and self._active_picker_type == picker_type:
            self.hide_inline_picker()
            return
        
        # Debug: Log what we're about to show
        self.log(f"show_inline_picker called with title='{title}', options={options[:3] if len(options) > 3 else options}...")
        
        # First, hide any existing inline picker (but don't clear callback yet)
        try:
            existing_picker = self.query_one("#inline-picker", InlinePicker)
            await existing_picker.remove()
        except Exception:
            pass
        
        input_container = self.query_one("#input-container", Vertical)
        command_input = self.query_one("#command-input", CommandInput)
        
        # Hide the command input
        command_input.display = False
        
        # Create and mount the picker
        picker = InlinePicker(title, options, id="inline-picker")
        await input_container.mount(picker, before=0)  # Mount at the top of the container
        
        # Track which picker is now open
        self._active_picker_type = picker_type
        
        # Update footer bindings to show the desired order when picker is open
        self._update_footer_bindings(picker_open=True)
        
        # Debug: Check the picker state after mounting
        self.log(f"Picker mounted. Option count: {len(picker._options)}")
        for idx, opt in enumerate(picker._options[:3]):
            self.log(f"  Option {idx}: prompt='{opt.prompt}'")
        
        picker.focus()
    
    def hide_inline_picker(self):
        """Hide the inline picker and restore the textarea."""
        try:
            picker = self.query_one("#inline-picker", InlinePicker)
            picker.remove()
        except Exception:
            pass
        
        # Clear the active picker type
        self._active_picker_type = None
        
        # Restore original footer bindings
        self._update_footer_bindings(picker_open=False)
        
        # Show the command input again
        command_input = self.query_one("#command-input", CommandInput)
        command_input.display = True
        command_input.focus()
    
    @on(InlinePickerSelected)
    async def handle_picker_selection(self, event: InlinePickerSelected) -> None:
        """Handle selection from inline picker."""
        selected_value = event.value
        
        # Check if AI is processing - block selection EXCEPT for command confirmation
        if self.processing and selected_value is not None and self._active_picker_type != 'command_confirm':
            self.app.notify(
                "⚠️ AI is busy.\nPress `Esc` to cancel current response.",
                severity="warning",
                timeout=2
            )
            # Don't hide the picker - let user try again or cancel
            return
        
        # Hide the picker first
        self.hide_inline_picker()
        
        # Check what type of picker this was
        if hasattr(self, '_picker_callback'):
            callback = self._picker_callback
            self._picker_callback = None
            
            # Execute the callback only if not cancelled
            if callback and selected_value is not None:
                # Check if callback is async or sync
                import inspect
                if inspect.iscoroutinefunction(callback):
                    await callback(selected_value)
                else:
                    # Call sync function directly
                    callback(selected_value)
    
    def action_cancel(self):
        """Cancel current operation (e.g., close inline picker, file panel, cancel command/agent prompts)."""
        # Priority 1: Check if inline picker or file explorer panel is visible - close both if open
        picker_closed = False
        file_panel_closed = False
        autocomplete_closed = False
        
        # Priority 1: Close agent & file panels if open
        try:
            picker = self.query_one("#inline-picker", InlinePicker)
            self.hide_inline_picker()
            # Clear any pending callback
            if hasattr(self, '_picker_callback'):
                self._picker_callback = None
            picker_closed = True
        except Exception:
            pass
        
        if self.show_files:
            file_panel = self.query_one("#file-panel", FileExplorerPanel)
            file_panel.remove_class("visible")
            self.show_files = False
            file_panel_closed = True
        
        # Check if autocomplete/suggestions list is open
        if hasattr(self, '_autocomplete_overlay') and self._autocomplete_overlay:
            if self._autocomplete_overlay.styles.display != "none":
                self.hide_autocomplete()
                autocomplete_closed = True
        
        # If we closed any UI elements, return focus to command input and don't cancel AI
        if picker_closed or file_panel_closed or autocomplete_closed:
            command_input = self.query_one("#command-input", CommandInput)
            command_input.focus()
            return
        
        # Priority 2: Cancel any ongoing streaming (only if no UI elements are open)
        if self._cancel_streaming == False:  # Only set flag and notify if not already cancelled
            self._cancel_streaming = True
            # Check if we're actually streaming
            if self.processing:
                self.app.notify(f"{EMOJI['cross']} Response cancelled.", severity="warning", timeout=2)
                
                # If using tool agent, cancel the async task
                if hasattr(self, '_current_tool_task') and self._current_tool_task and not self._current_tool_task.done():
                    # Cancel the task (this will raise CancelledError in the task)
                    self._current_tool_task.cancel()
        
        # Priority 4: Cancel any ongoing agent execution
        if hasattr(self, '_agent_input_state') and self._agent_input_state:
            # Check if agent is actively running (not just waiting for input)
            if self._agent_input_state.get('waiting_for_input', False):
                # Agent is waiting for input - abort it
                self._agent_input_state['result'] = None
                self._agent_input_state['waiting_for_input'] = False
                self._agent_input_state['cancelled'] = True  # Mark as cancelled
                self._agent_input_state['ready'].set()
                # Clear current agent indicator
                chat_panel = self.query_one("#chat-panel", ChatPanel)
                chat_panel.current_agent = None
                # Don't show message here - the agent thread will show it
                return
        
        # Priority 5: Check if we're waiting for command input
        if hasattr(self, '_command_input_state') and self._command_input_state.get('waiting_for_input', False):
            self._command_input_state['result'] = None
            self._command_input_state['waiting_for_input'] = False
            self._command_input_state['ready'].set()
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.add_message("system", f"{EMOJI['cross']} Command cancelled.")
            return
    
    @on(TextAreaSubmitted)
    async def handle_input(self, event: TextAreaSubmitted) -> None:
        """Handle user input submission."""
        message = event.value
        
        # Check if we're waiting for command input (like /add_agent prompts)
        if hasattr(self, '_command_input_state') and self._command_input_state.get('waiting_for_input', False):
            # Display user's response
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            display_msg = message if message else ""
            chat_panel.add_message("user", display_msg)
            
            # Provide the response to the command
            self._command_input_state['result'] = message
            self._command_input_state['waiting_for_input'] = False
            self._command_input_state['ready'].set()
            return
        
        # Check if we're waiting for agent input (allow blank entries for agents)
        if hasattr(self, '_agent_input_state') and self._agent_input_state.get('waiting_for_input', False):
            # Display user's response (show empty message as well)
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            display_msg = message if message else ""
            chat_panel.add_message("user", display_msg)
            
            # Provide the response to the agent (can be empty string)
            self._agent_input_state['result'] = message
            self._agent_input_state['waiting_for_input'] = False  # Reset the flag
            self._agent_input_state['ready'].set()
            return
        
        # For normal chat/commands, require non-empty input
        if not message.strip():
            return
        
        # Process command or message
        if message.strip().startswith("/"):
            command_handled = await self.handle_command(message.strip())
            # If command was not recognized, treat it as normal chat input
            if not command_handled:
                # Remove the leading '/' since it's not actually a command
                message_without_slash = message.lstrip('/').strip()
                
                # Strip # from file references for the actual message to send
                message_to_send = self._strip_file_markers(message_without_slash)
                
                # Format file references for display (#file -> `file`)
                message_for_display = self._format_file_references(message_without_slash)
                
                # For multiline messages, we need to format them properly for Markdown
                # Convert single newlines to double newlines for proper Markdown line breaks
                formatted_message = message_for_display.replace('\n', '\n\n')
                
                # Display user message immediately (without the /)
                chat_panel = self.query_one("#chat-panel", ChatPanel)
                chat_panel.add_message("user", formatted_message, raw_content=message_to_send)
                
                # Note: autosave is now handled in ChatPanel.add_message()
                
                # Send the message without slash to the chatbot (without double newlines or # markers) in background thread
                import threading
                thread = threading.Thread(
                    target=self._send_chat_message_in_thread,
                    args=(message_to_send,),
                    daemon=True
                )
                thread.start()
        else:
            # Strip # from file references for the actual message to send
            message_to_send = self._strip_file_markers(message)
            
            # Format file references for display (#file -> `file`)
            message_for_display = self._format_file_references(message)
            
            # For multiline messages, we need to format them properly for Markdown
            # Convert single newlines to double newlines for proper Markdown line breaks
            formatted_message = message_for_display.replace('\n', '\n\n')
            
            # Display user message immediately
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.add_message("user", formatted_message, raw_content=message_to_send)
            
            # Note: autosave is now handled in ChatPanel.add_message()
            
            # Send the original message to the chatbot (without double newlines or # markers) in background thread
            import threading
            thread = threading.Thread(
                target=self._send_chat_message_in_thread,
                args=(message_to_send,),
                daemon=True
            )
            thread.start()
    
    def _strip_file_markers(self, command: str) -> str:
        """Strip # markers from file paths in commands.
        
        The # prefix is used in the TUI for file autocomplete suggestions,
        but needs to be removed before processing actual file paths.
        
        Args:
            command: The command string that may contain #file paths
            
        Returns:
            Command string with # markers removed from file paths
        """
        import re
        # Replace #filepath patterns with filepath (preserving spaces)
        # Match # followed by non-whitespace characters (file path)
        return re.sub(r'#([^\s]+)', r'./\1' if IS_LINUX else r'.\\\1', command)
    
    def _format_file_references(self, message: str) -> str:
        """Format file references in chat messages.
        
        Strips # markers and adds backtick highlighting for file references.
        Example: "Check #config.py and #src/main.py" -> "Check `config.py` and `src/main.py`"
        
        Args:
            message: The chat message that may contain #file references
            
        Returns:
            Formatted message with file references highlighted
        """
        import re
        # Replace #filepath patterns with `filepath` (backtick highlighting)
        # Match # followed by non-whitespace characters (file path)
        return re.sub(r'#([^\s]+)', r'`\1`', message)
    
    async def handle_command(self, command: str) -> bool:
        """Handle slash commands using CommandManager.
        
        Returns:
            bool: True if command was handled, False if command is not recognized
        """
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        # Strip # markers from file paths before processing
        command = self._strip_file_markers(command)
        
        parts = command.split(maxsplit=1)
        cmd = parts[0].lower()
        args = parts[1] if len(parts) > 1 else ""
        
        # Check for /add_agent and /add_tool commands - handle with threading directly
        if cmd == "/add_agent":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Start in a separate thread using Python's threading module directly
            import threading
            thread = threading.Thread(
                target=self._run_add_agent_in_thread,
                args=(args,),
                daemon=True
            )
            thread.start()
            return True
        
        elif cmd == "/add_tool":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Start in a separate thread using Python's threading module directly
            import threading
            thread = threading.Thread(
                target=self._run_add_tool_in_thread,
                args=(args,),
                daemon=True
            )
            thread.start()
            return True
        
        # Handle TUI-specific commands first (these need special UI handling)
        if cmd == "/quit":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            self.app.exit()
            return True
        
        elif cmd == "/reload":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Handle /reload BEFORE CommandManager to avoid blocking Confirm.ask()
            chat_panel.add_message("system", """### ↻ Reload Configuration

Restarting TUI to reload configuration...
""")
            # Set restart flag and schedule exit
            self.app.should_restart = True
            
            def do_restart():
                self.app.exit()
            
            self.set_timer(0.5, do_restart)
            return True  # Return immediately after scheduling
        
        elif cmd == "/update":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Run update in TUI using ValbotUpdater in a background thread
            import threading
            
            chat_panel.add_message("system", f"{EMOJI['gear']} Checking for updates...")
            
            # Run update in background thread
            thread = threading.Thread(
                target=self._run_update_in_thread,
                daemon=True
            )
            thread.start()
            return True
        
        elif cmd == "/agent":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Override CLI's agent command with TUI-specific inline picker
            if self.chatbot and hasattr(self.chatbot, 'plugin_manager'):
                agents = self.chatbot.plugin_manager.plugin_info
                if agents:
                    async def on_agent_selected(selected_agent_name):
                        if selected_agent_name:
                            # Find the agent details
                            for agent_name, agent_desc in agents:
                                if agent_name == selected_agent_name:
                                    # Run agent in background thread using threading directly
                                    import threading
                                    thread = threading.Thread(
                                        target=self._run_agent_workflow_in_thread,
                                        args=(agent_name, agent_desc),
                                        daemon=True
                                    )
                                    thread.start()
                                    break
                    
                    # Format agents for display: show name and description
                    agent_options = [(f"{name}: {desc}", name) for name, desc in agents]
                    
                    # Show the inline picker first
                    await self.show_inline_picker(f"{EMOJI['agent']} Select Agent", agent_options)
                    
                    # Store callback for when selection is made (after picker is shown)
                    self._picker_callback = on_agent_selected
                else:
                    chat_panel.add_message("system", "No agents available. Use `/add_agent` to add one.")
            else:
                chat_panel.add_message("error", "Plugin manager not initialized.")
            return True
        
        elif cmd == "/model":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Override CLI's model command with TUI-specific picker
            await self.action_change_model()
            return True
        
        elif cmd == "/new":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Use the proper action_new_chat method which handles all cleanup
            await self.action_new_chat()
            chat_panel.add_message("system", f"{EMOJI['checkmark']} Started new chat.")
            return True
            
        elif cmd == "/help":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            chat_panel.add_message("system", """### ⌨️ ValBot Help

**Available Commands:**
- `/agent` - Select and run an agent workflow
- `/model` - Change the AI model
- `/new` - Start a new conversation
- `/prompts` - Show available custom prompts
- `/commands` - Show all available commands
- `/settings` - Show settings (CLI mode only)
- `/create_database <folder_path> #<file1> #<file1>` - Create a RAG database (PDF, DOCX, TXT, MD, etc.)
- `/load_database <folder_path>` - Load RAG database(s)
- `/quit` - Exit the application

**Usage Tips:**
- Type messages normally for chat
- Use `/` prefix for commands
- Agent workflows provide specialized functionality
- Context files are remembered throughout conversation
- **RAG databases** enable Q&A with your documents
- Load multiple databases with space-separated paths
- Click ✕ on database chips to unload specific databases
- Use `/new` to unload all databases and start fresh

For more detailed help, try the CLI mode with `--help` flag.
""")
            return True
            
        elif cmd == "/settings":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            chat_panel.add_message("system", """### {EMOJI['gear']} Settings

Settings management is available in CLI mode only.
To modify settings, exit TUI and run:

```bash
python app.py
/settings
```

**Current Model:** """ + str(self.config_manager.get_setting("chat_model_config.default_model", "gpt-4o")) + """

Use `/model` to change the current model.
""")
            return True
        
        # Try to use CommandManager for standard commands
        if hasattr(self.chatbot, 'command_manager'):
            # Check if command exists in CommandManager
            if cmd in self.chatbot.command_manager.command_registry or cmd in self.chatbot.command_manager.prompts:
                # Display the user's command input in the chat with backtick formatting
                chat_panel.add_message("user", f"`{command}`")
                try:
                    # Let CommandManager handle the command
                    self.chatbot.command_manager.handle_command(command)
                    return True
                except Exception as e:
                    chat_panel.add_message("error", f"""### {EMOJI['cross']} Command Error

Error executing command `{cmd}`:

```
{str(e)}
```
""")
                    return True
        
        # Handle additional TUI-specific commands not in CommandManager
        if cmd == "/help":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            help_text = """# 📚 ValBot TUI Help

### Available Commands

### Chat Management
- `/new` - Start a new chat session
- `/quit` - Exit the application
- `/prompts` - Show available custom prompts
- `/commands` - Show all available commands

### Model & AI Configuration
- `/model` - Change AI model (interactive picker with ↑/↓ keys)
  - Available models: gpt-4o, gpt-5, gpt-4.1, gpt-oss:20b
  
### Context & File Management
- `/file <path>` - Display file content with syntax highlighting
  - Example: `/file ./config.py`

### RAG database (Document Q&A)
- `/create_database <folder> <files>` - Create a searchable database
  - Example: `/create_database ./kb doc1.pdf doc2.txt notes.md data.xlsx`
  - Supports: PDF, DOCX, XLSX, XLS, TXT, MD, PDL, PY, C, CPP, H files
  - Also supports .gz compressed files (.pdf.gz, .txt.gz, .log.gz, etc.)
  - Creates vector embeddings for semantic search
- `/load_database <folder_path>` - Load one or more databases
  - Example: `/load_database ./kb` - Load single database
  - Example: `/load_database ./kb1 ./kb2 ./kb3` - Load multiple databases
  - After loading, just ask questions normally in chat!
  - System automatically searches all loaded databases and provides answers
  - Click the ✕ button on a database chip to unload that specific database
  - Use `/new` to unload all databases and return to normal chat

### Agent System (Advanced Workflows)
- `/agent` - Run an agent flow (interactive selection with ↑/↓ keys)
  - Agents can perform complex multi-step tasks
  - Select from available agents like File Edit, Terminal, Spec Expert, etc.
- `/add_agent <url_or_path>` - Add a new agent from git repo or local path
- `/add_tool <url_or_path>` - Add a new tool from git repo or local path

### System Integration
- `/terminal <command>` - Execute a shell command
  - Example: `/terminal ls -la`
  - Example: `/terminal python --version`
  - Example: `/terminal git status`
- `/multi` - Multi-line input using system editor
  - Opens your default editor (set via $EDITOR env variable)
  - Default: notepad (Windows) or vim (Linux/Mac)

### Configuration & Updates
- `/settings` - View settings information
- `/reload` - Reinitialize chatbot with current configuration
- `/update` - Check for and install updates from GitHub
  - Automatically fetches latest changes
  - Shows commit log of updates
  - Handles git pull and requirements updates
  - Prompts before applying changes

### {EMOJI['keyboard']} Keyboard Shortcuts

| Shortcut | Action | Description |
|----------|--------|-------------|
| **ctrl+d** | Agent | Select and run an agent workflow |
| **ctrl+f** | Files | Toggle file explorer panel |
| **ctrl+g** | History | Toggle chat history panel |
| **ctrl+n** | New Chat | Clear conversation and start fresh |
| **ctrl+o** | Model | Open model picker dialog |
| **escape** | Cancel | Cancel current operation/close dialogs |
| **ctrl+q** | Quit | Exit the application |

### 💬 Chat Features

- **Markdown Support**: Full markdown rendering with headers, lists, tables, quotes
- **Code Highlighting**: Syntax-highlighted code blocks with copy buttons
  - Supports: Python, JavaScript, Java, C++, Go, Rust, and more
- **Streaming**: Real-time response streaming as AI generates text
- **Autosave**: Chat history is automatically saved to disk
- **Reasoning Display**: GPT-5 shows its thinking process (if enabled)
  - Configure in `user_config.json`: `"display_reasoning": true`
  - Set effort level: `"reasoning_effort": "low|medium|high"`
- **Context Awareness**: Maintains full conversation history
- **RAG database**: Q&A with your documents using semantic search
  - Create once, query anytime
  - Supports multiple file formats
  - Powered by vector embeddings and ChromaDB
- **Agents**: Run complex agentic workflows for tasks like:
  - File editing
  - Terminal operations
  - Spec analysis
  - Project scaffolding
  - And more!

### 🎨 UI Features

- **Material Design**: Modern, beautiful dark theme interface
- **Gradient Accents**: Colorful borders and highlights
- **Message Types**: Color-coded by role (user/assistant/system/error)
  - 🔵 Blue border = Assistant response
  - 🟡 Yellow border = System message
  - 🔴 Red border = Error message
- **Responsive Layout**: Adapts to your terminal size
- **Smooth Scrolling**: Elegant animations and transitions
- **Interactive Dialogs**: Beautiful modal dialogs for selection tasks

### 📚 Custom Prompts

ValBot supports custom prompts defined in your config file. These are shortcuts for common tasks:

- Use `/prompts` to see all available custom prompts
- Custom prompts can include placeholders and execute shell commands
- Example: `/my-prompt "argument1" "argument2"`

**Define custom prompts in `user_config.json`:**
```json
{
  "custom_prompts": [
    {
      "prompt_cmd": "review",
      "description": "Review code in a file",
      "prompt": "Please review this code: {file}",
      "args": ["file"]
    }
  ]
}
```

### {EMOJI['robot']} Agent Plugins

Agents are powerful workflows that can perform complex tasks:

**Built-in Agents:**
- **Spec Expert** - Q&A with specification documents
- **File Edit** - Edit files with AI assistance
- **File Creator** - Create new files with AI assistance
- **Terminal** - Run terminal commands intelligently
- **Repo Index** - Summarize and index repository context
- **Project** - Make project-wide changes and add features

**Using Agents:**
1. Type `/agent` or press ctrl+o
2. Use ↑/↓ arrow keys to select
3. Press Enter to run
4. Follow the agent's prompts

### 🔧 Configuration Tips

Edit `user_config.json` to customize your experience:

```json
{
  "chat_model_config": {
    "default_model": "gpt-5",
    "You are a helpful assistant. Prioritize markdown format and code blocks when applicable.",
    "display_reasoning": false,
    "reasoning_effort": "low"
  },
  "agent_model_config": {
    "default_model": "gpt-4.1"
  },
  "general": {
    "display_commands_on_startup": true
  }
}
```

### 🆘 Troubleshooting

**Q: Commands not working?**
- Make sure command starts with `/`
- Check spelling carefully
- Use `/commands` to see all available commands

**Q: Can't see GPT-5 reasoning?**
- Set `"display_reasoning": true` in config
- Only works with gpt-5 models
- Restart TUI after config changes

**Q: Agent not found?**
- Use `/agent` to see available agents
- Some agents may need to be installed
- Check config for agent definitions

**Q: How do I exit?**
- Press **Ctrl+Q** or type `/quit`

### 📖 More Resources

- **Full Documentation**: See `README_TUI.md`
- **Feature List**: See `TUI_FEATURES_IMPLEMENTED.md`
- **Quick Reference**: See `TUI_QUICK_REFERENCE_CARD.md`

Need more help? Just ask ValBot directly!

---
**ValBot TUI** - Your AI-Powered Assistant
"""
            chat_panel.add_message("system", help_text)
            return True
                
        elif cmd == "/terminal":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            if args:
                await chat_panel.add_terminal_output(args)
            else:
                chat_panel.add_message("system", """### {EMOJI['keyboard']} Terminal Command

**Usage**: `/terminal <command>`

**Examples**:
- `/terminal ls -la`
- `/terminal python --version`
- `/terminal git status`
""")
            return True
                
        elif cmd == "/file":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            if args:
                try:
                    file_path = Path(args).expanduser()
                    if file_path.exists() and file_path.is_file():
                        # Check file size before reading
                        file_size = file_path.stat().st_size
                        max_size_bytes = 10 * 1024 * 1024  # 10 MB limit for display
                        
                        if file_size > max_size_bytes:
                            chat_panel.add_message("error", 
                                f"{EMOJI['cross']} File too large to display: {file_size / (1024*1024):.2f} MB (limit: 10 MB)\n"
                                f"Reference the file directly in your message instead.")
                        else:
                            # Read in chunks for better memory handling
                            content_chunks = []
                            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                                while True:
                                    chunk = f.read(8192)  # Read in 8KB chunks
                                    if not chunk:
                                        break
                                    content_chunks.append(chunk)
                            content = ''.join(content_chunks)
                            
                            # Determine file type for syntax highlighting
                            file_ext = file_path.suffix.lower()
                            lang_map = {
                                '.py': 'python', '.js': 'javascript', '.ts': 'typescript',
                                '.java': 'java', '.cpp': 'cpp', '.c': 'c', '.cs': 'csharp',
                                '.go': 'go', '.rs': 'rust', '.rb': 'ruby', '.php': 'php',
                                '.html': 'html', '.css': 'css', '.json': 'json', '.xml': 'xml',
                                '.yaml': 'yaml', '.yml': 'yaml', '.sh': 'bash', '.md': 'markdown'
                            }
                            lang = lang_map.get(file_ext, 'text')
                            
                            formatted_content = f"""### 📄 File: `{file_path}`

```{lang}
{content}
```

**Lines**: {len(content.splitlines())} | **Size**: {file_size:,} bytes ({file_size / 1024:.2f} KB)
"""
                            chat_panel.add_message("system", formatted_content)
                    else:
                        chat_panel.add_message("error", f"{EMOJI['cross']} File not found: `{args}`")
                except MemoryError:
                    chat_panel.add_message("error", f"{EMOJI['cross']} Memory error: File too large to read")
                except Exception as e:
                    chat_panel.add_message("error", f"{EMOJI['cross']} Error reading file: {str(e)}")
            else:
                chat_panel.add_message("system", "**Usage**: `/file <path>`\n\nExample: `/file ./config.py`")
            return True
        
        elif cmd == "/multi":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Multi-line input using system editor
            chat_panel.add_message("system", """### 📝 Multi-line Input

Opening your default editor for multi-line input...

**Note**: Set your `EDITOR` environment variable to choose your preferred editor.
Default: vim (Linux/Mac) or notepad (Windows)
""")
            try:
                editor = os.environ.get('EDITOR', 'notepad' if sys.platform == 'win32' else 'vim')
                with tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.md') as temp_file:
                    temp_file_name = temp_file.name
                
                # Open editor
                subprocess.run([editor, temp_file_name], check=True)
                
                # Read content
                with open(temp_file_name, 'r') as f:
                    content = f.read().strip()
                
                os.remove(temp_file_name)
                
                if content:
                    chat_panel.add_message("user", content)
                    # Send in background thread
                    import threading
                    thread = threading.Thread(
                        target=self._send_chat_message_in_thread,
                        args=(content,),
                        daemon=True
                    )
                    thread.start()
                else:
                    chat_panel.add_message("system", "No content entered.")
                    
            except Exception as e:
                chat_panel.add_message("error", f"{EMOJI['cross']} Error with multi-line input: {str(e)}")
            return True
        
        elif cmd == "/create_database":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Create a new RAG database
            if not args:
                chat_panel.add_message("system", """### 📚 Create RAG Database

**Usage**: `/create_database <output_folder> <file1> [file2] [file3] ...`

**Example**:
- `/create_database ./my_kb document1.pdf document2.docx notes.txt`
- `/create_database C:/kb *.pdf *.txt`

**Supported file types**: PDF, DOCX, TXT, MD, PDL, PY, C, CPP, H
**Compressed files**: All formats also support .gz compression (e.g., .pdf.gz, .txt.gz, .log.gz)

This will process your documents and create a searchable database with vector embeddings.
The cache and database files will be stored in the output folder.
**Note**: Source files are NOT copied - only the database cache is created.
""")
            else:
                # Run in background thread
                import threading
                thread = threading.Thread(
                    target=self._run_create_database_in_thread,
                    args=(args,),
                    daemon=True
                )
                thread.start()
            return True
        
        elif cmd == "/load_database":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            # Load and query an existing RAG database
            if not args:
                chat_panel.add_message("system", """### 📚 Load RAG database

**Usage**: `/load_database <kb_directory>`

**Example**:
- `/load_database ./my_kb`
- `/load_database C:/Users/myuser/Documents/my_kb`

This will load an existing database and allow you to ask questions about the documents.
After loading, just type your questions normally in chat!
""")
            else:
                # Run in background thread
                import threading
                thread = threading.Thread(
                    target=self._run_load_database_in_thread,
                    args=(args.strip(),),
                    daemon=True
                )
                thread.start()
            return True
        
        elif cmd == "/settings":
            # Display the user's command input in the chat with backtick formatting
            chat_panel.add_message("user", f"`{command}`")
            chat_panel.add_message("system", """### {EMOJI['gear']} Settings

The settings TUI is not yet available in this version.

You can manually edit your configuration file at:
`user_config.json`

**Available Settings**:
- `chat_model_config.default_model` - Default AI model
- `chat_model_config.system_prompt` - System instructions
- `chat_model_config.display_reasoning` - Show reasoning (GPT-5)
- `chat_model_config.reasoning_effort` - Reasoning effort level
- `agent_model_config.default_model` - Agent model
- `general.ascii_banner_size` - Banner size
- `general.display_commands_on_startup` - Show commands on start
""")
            return True
                
        else:
            # Command not recognized - return False so it can be treated as normal input
            return False
    
    
    def _run_add_agent_in_thread(self, args: str):
        """Run /add_agent in a separate thread with monkey patches (not @work decorator)."""
        import asyncio
        import sys
        import threading
        import traceback
        import re
        import importlib
        from rich.prompt import Prompt, Confirm
        from rich.console import Console
        from prompt_toolkit import PromptSession
        import rich.prompt
        import rich.console
        from extension_manager import agent_adder
        
        # Set up event loop for this thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        if sys.version_info >= (3, 10):
            try:
                asyncio.get_event_loop_policy()
            except RuntimeError:
                asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
        
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        console_wrapper = self.chatbot.command_manager.console
        
        # Save original classes/functions
        original_prompt_ask = Prompt.ask
        original_confirm_ask = Confirm.ask
        original_prompt_session_prompt = PromptSession.prompt
        original_console_class = rich.console.Console
        
        # Create shared state for input handling
        command_input_state = {
            'waiting_for_input': False,
            'prompt_text': None,
            'result': None,
            'ready': threading.Event()
        }
        
        self._command_input_state = command_input_state
        
        # Monkey-patch Prompt.ask
        def tui_prompt_ask(prompt="", **kwargs):
            """Replacement for Prompt.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            
            choices = kwargs.get('choices', None)
            default = kwargs.get('default', None)
            
            if choices:
                choices_text = ", ".join(f"`{c}`" for c in choices)
                full_prompt = f"**{clean_prompt}**\n\nChoices: {choices_text}"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            else:
                full_prompt = f"**{clean_prompt}**"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default:
                return default
            return result if result else ""
        
        # Monkey-patch Confirm.ask
        def tui_confirm_ask(prompt="", **kwargs):
            """Replacement for Confirm.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            default = kwargs.get('default', True)
            
            full_prompt = f"**{clean_prompt}**\n\nChoices: `y`, `n`"
            if default is not None:
                full_prompt += f"\n\n_Default: `{'y' if default else 'n'}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default is not None:
                return default
            
            return result.lower() in ['y', 'yes', 'true', '1']
        
        # Create TUI Console wrapper
        _captured_console_wrapper = console_wrapper
        
        class TUIConsole:
            """Console replacement for TUI mode."""
            def __init__(self, *args, **kwargs):
                self._wrapper = _captured_console_wrapper
            
            def print(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            def rule(self, *args, **kwargs):
                return self._wrapper.rule(*args, **kwargs)
            
            def status(self, *args, **kwargs):
                return self._wrapper.status(*args, **kwargs)
            
            def input(self, prompt="", **kwargs):
                return self._wrapper.input(prompt, **kwargs)
            
            def log(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            def path_prompt(self, prompt="", **kwargs):
                """Handle path_prompt used by AgentAdder."""
                return tui_prompt_ask(prompt, **kwargs)
            
            @property
            def width(self):
                return self._wrapper.width
            
            @property
            def height(self):
                return self._wrapper.height
            
            def is_terminal(self):
                return self._wrapper.is_terminal()
        
        # Apply monkey patches
        rich.prompt.Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        rich.prompt.Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        rich.console.Console = TUIConsole
        
        # Reload extension_manager modules
        importlib.reload(agent_adder)
        
        # Add starting message
        self.app.call_from_thread(
            chat_panel.add_message,
            "system",
            f"{EMOJI['refresh']} Starting `/add_agent` command..."
        )
        
        try:
            # Use the reloaded module
            AgentAdder = agent_adder.AgentAdder
            tui_console_instance = TUIConsole()
            adder = AgentAdder(tui_console_instance, self.chatbot.config_manager)
            adder.run(args if args else None)
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "system",
                f"{EMOJI['checkmark']} Agent added successfully! Run `/reload` to load the new agent."
            )
            
        except Exception as e:
            error_details = f"{EMOJI['cross']} Error adding agent: {str(e)}\n\n```\n{traceback.format_exc()}\n```"
            self.app.call_from_thread(
                chat_panel.add_message,
                "error",
                error_details
            )
        
        finally:
            # Clean up
            if hasattr(self, '_command_input_state'):
                self._command_input_state['waiting_for_input'] = False
                delattr(self, '_command_input_state')
            
            # Restore originals
            rich.prompt.Prompt.ask = original_prompt_ask
            rich.prompt.Confirm.ask = original_confirm_ask
            Prompt.ask = original_prompt_ask
            Confirm.ask = original_confirm_ask
            rich.console.Console = original_console_class
            
            try:
                if not loop.is_closed():
                    loop.close()
            except Exception:
                pass
    
    def _run_add_tool_in_thread(self, args: str):
        """Run /add_tool in a separate thread with monkey patches (not @work decorator)."""
        import asyncio
        import sys
        import threading
        import traceback
        import re
        import importlib
        from rich.prompt import Prompt, Confirm
        from rich.console import Console
        from prompt_toolkit import PromptSession
        import rich.prompt
        import rich.console
        from extension_manager import tool_adder
        
        # Set up event loop for this thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        if sys.version_info >= (3, 10):
            try:
                asyncio.get_event_loop_policy()
            except RuntimeError:
                asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
        
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        console_wrapper = self.chatbot.command_manager.console
        
        # Save original classes/functions
        original_prompt_ask = Prompt.ask
        original_confirm_ask = Confirm.ask
        original_prompt_session_prompt = PromptSession.prompt
        original_console_class = rich.console.Console
        
        # Create shared state for input handling
        command_input_state = {
            'waiting_for_input': False,
            'prompt_text': None,
            'result': None,
            'ready': threading.Event()
        }
        
        self._command_input_state = command_input_state
        
        # Monkey-patch Prompt.ask
        def tui_prompt_ask(prompt="", **kwargs):
            """Replacement for Prompt.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            
            choices = kwargs.get('choices', None)
            default = kwargs.get('default', None)
            
            if choices:
                choices_text = ", ".join(f"`{c}`" for c in choices)
                full_prompt = f"**{clean_prompt}**\n\nChoices: {choices_text}"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            else:
                full_prompt = f"**{clean_prompt}**"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default:
                return default
            return result if result else ""
        
        # Monkey-patch Confirm.ask
        def tui_confirm_ask(prompt="", **kwargs):
            """Replacement for Confirm.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            default = kwargs.get('default', True)
            
            full_prompt = f"**{clean_prompt}**\n\nChoices: `y`, `n`"
            if default is not None:
                full_prompt += f"\n\n_Default: `{'y' if default else 'n'}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default is not None:
                return default
            
            return result.lower() in ['y', 'yes', 'true', '1']
        
        # Create TUI Console wrapper
        _captured_console_wrapper = console_wrapper
        
        class TUIConsole:
            """Console replacement for TUI mode."""
            def __init__(self, *args, **kwargs):
                self._wrapper = _captured_console_wrapper
            
            def print(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            def rule(self, *args, **kwargs):
                return self._wrapper.rule(*args, **kwargs)
            
            def status(self, *args, **kwargs):
                return self._wrapper.status(*args, **kwargs)
            
            def input(self, prompt="", **kwargs):
                return self._wrapper.input(prompt, **kwargs)
            
            def log(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            def path_prompt(self, prompt="", **kwargs):
                """Handle path_prompt used by ToolAdder."""
                return tui_prompt_ask(prompt, **kwargs)
            
            @property
            def width(self):
                return self._wrapper.width
            
            @property
            def height(self):
                return self._wrapper.height
            
            def is_terminal(self):
                return self._wrapper.is_terminal()
        
        # Apply monkey patches
        rich.prompt.Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        rich.prompt.Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        rich.console.Console = TUIConsole
        
        # Reload extension_manager modules
        importlib.reload(tool_adder)
        
        # Add starting message
        self.app.call_from_thread(
            chat_panel.add_message,
            "system",
            f"{EMOJI['refresh']} Starting `/add_tool` command..."
        )
        
        try:
            # Use the reloaded module
            ToolAdder = tool_adder.ToolAdder
            tui_console_instance = TUIConsole()
            adder = ToolAdder(tui_console_instance, self.chatbot.config_manager)
            adder.run(args if args else None)
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "system",
                f"{EMOJI['checkmark']} Tool added successfully! Run `/reload` to load the new tool."
            )
            
        except Exception as e:
            error_details = f"{EMOJI['cross']} Error adding tool: {str(e)}\n\n```\n{traceback.format_exc()}\n```"
            self.app.call_from_thread(
                chat_panel.add_message,
                "error",
                error_details
            )
        
        finally:
            # Clean up
            if hasattr(self, '_command_input_state'):
                self._command_input_state['waiting_for_input'] = False
                delattr(self, '_command_input_state')
            
            # Restore originals
            rich.prompt.Prompt.ask = original_prompt_ask
            rich.prompt.Confirm.ask = original_confirm_ask
            Prompt.ask = original_prompt_ask
            Confirm.ask = original_confirm_ask
            rich.console.Console = original_console_class
            
            try:
                if not loop.is_closed():
                    loop.close()
            except Exception:
                pass
    
    def _run_update_in_thread(self):
        """Run /update using ValbotUpdater - non-blocking and TUI-friendly."""
        import threading
        import rich
        import rich.prompt
        import rich.console
        from rich.prompt import Prompt, Confirm
        import re
        
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        console_wrapper = self.chatbot.command_manager.console
        
        # Save original classes/functions
        original_prompt_ask = Prompt.ask
        original_confirm_ask = Confirm.ask
        original_console_class = rich.console.Console
        
        # Create shared state for input handling
        command_input_state = {
            'waiting_for_input': False,
            'prompt_text': None,
            'result': None,
            'ready': threading.Event()
        }
        
        self._command_input_state = command_input_state
        
        # Monkey-patch Prompt.ask
        def tui_prompt_ask(prompt="", **kwargs):
            """Replacement for Prompt.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            
            choices = kwargs.get('choices', None)
            default = kwargs.get('default', None)
            
            if choices:
                choices_text = ", ".join(f"`{c}`" for c in choices)
                full_prompt = f"**{clean_prompt}**\n\nChoices: {choices_text}"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            else:
                full_prompt = f"**{clean_prompt}**"
                if default:
                    full_prompt += f"\n\n_Default: `{default}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default:
                return default
            return result if result else ""
        
        # Monkey-patch Confirm.ask
        def tui_confirm_ask(prompt="", **kwargs):
            """Replacement for Confirm.ask that uses chat-based input."""
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            default = kwargs.get('default', True)
            
            full_prompt = f"**{clean_prompt}**\n\nChoices: `y`, `n`"
            if default is not None:
                full_prompt += f"\n\n_Default: `{'y' if default else 'n'}`_"
            
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            command_input_state['waiting_for_input'] = True

            
            command_input_state['prompt_text'] = clean_prompt

            
            command_input_state['result'] = None

            
            command_input_state['ready'].clear()

            
            

            
            # Wait with timeout to prevent hanging

            
            if not command_input_state['ready'].wait(timeout=300):  # 5 minute timeout

            
                command_input_state['waiting_for_input'] = False

            
                self.app.call_from_thread(

            
                    chat_panel.add_message,

            
                    "error",

            
                    "⏱️ Input timeout - command cancelled"

            
                )

            
                return default if default else ""

            
            

            
            result = command_input_state['result']
            command_input_state['waiting_for_input'] = False
            
            if not result and default is not None:
                return default
            
            return result.lower() in ['y', 'yes', 'true', '1']
        
        # Create TUI Console wrapper
        _captured_console_wrapper = console_wrapper
        
        class TUIConsole:
            """Console replacement for TUI mode."""
            def __init__(self, *args, **kwargs):
                self._wrapper = _captured_console_wrapper
            
            def print(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            def rule(self, *args, **kwargs):
                return self._wrapper.rule(*args, **kwargs)
            
            def status(self, *args, **kwargs):
                return self._wrapper.status(*args, **kwargs)
            
            def input(self, prompt="", **kwargs):
                return self._wrapper.input(prompt, **kwargs)
            
            def log(self, *args, **kwargs):
                return self._wrapper.print(*args, **kwargs)
            
            @property
            def width(self):
                return 80
            
            @property
            def height(self):
                return 24
            
            def is_terminal(self):
                return self._wrapper.is_terminal()
        
        # Apply monkey patches
        rich.prompt.Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        rich.prompt.Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        Prompt.ask = classmethod(lambda cls, *args, **kwargs: tui_prompt_ask(*args, **kwargs))
        Confirm.ask = classmethod(lambda cls, *args, **kwargs: tui_confirm_ask(*args, **kwargs))
        rich.console.Console = TUIConsole
        
        try:
            # Create TUI console wrapper
            tui_console = TUIConsole()
            
            # Define callbacks for reload and exit
            def reload_callback():
                """Called when update completes and reload is needed."""
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "system",
                    f"{EMOJI['checkmark']} **Update complete!**\n\n"
                    f"Please type `/reload` or manually restart the TUI to apply changes."
                )
            
            def exit_callback():
                """Called when running from exe and rebuild is needed."""
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "system",
                    f"{EMOJI['info']} Update complete. Please rebuild the executable as instructed.\n\n"
                    f"Press `Ctrl+Q` to exit."
                )
            
            # Create and run the updater
            updater = ValbotUpdater(
                console=tui_console,
                reload_callback=reload_callback,
                exit_callback=exit_callback
            )
            
            # Run the update
            updater.update()
            
        except Exception as e:
            import traceback
            error_details = f"{EMOJI['cross']} Unexpected error during update:\n\n```\n{str(e)}\n\n{traceback.format_exc()}\n```"
            self.app.call_from_thread(
                chat_panel.add_message,
                "error",
                error_details
            )
        
        finally:
            # Clean up
            if hasattr(self, '_command_input_state'):
                self._command_input_state['waiting_for_input'] = False
                delattr(self, '_command_input_state')
            
            # Restore originals
            rich.prompt.Prompt.ask = original_prompt_ask
            rich.prompt.Confirm.ask = original_confirm_ask
            Prompt.ask = original_prompt_ask
            Confirm.ask = original_confirm_ask
            rich.console.Console = original_console_class

    
    def _run_agent_workflow_in_thread(self, agent_name: str, agent_desc: str, agent_init_args: dict = None):
        """Run an agent workflow in a separate thread to avoid asyncio conflicts."""
        import asyncio
        import sys
        import threading
        from rich.prompt import Prompt
        from rich.console import Console
        from prompt_toolkit import PromptSession
        
        # For worker threads, we need to ensure asyncio works properly
        # Set a new event loop for this thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        # Store original policy if needed
        if sys.version_info >= (3, 10):
            # In Python 3.10+, asyncio.run() can have issues in threads
            # We need to ensure the event loop policy is set correctly
            try:
                asyncio.get_event_loop_policy()
            except RuntimeError:
                asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
        
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        console_wrapper = self.chatbot.command_manager.console
        
        # Save original classes/functions for restoration
        original_prompt_ask = Prompt.ask
        original_prompt_session_prompt = PromptSession.prompt
        import rich.console
        original_console_class = rich.console.Console
        
        # Create a shared state for agent input handling
        agent_input_state = {
            'waiting_for_input': False,
            'prompt_text': None,
            'result': None,
            'ready': threading.Event(),
            'last_menu_choices': None,  # Track menu choices from printed tables
            'cancelled': False  # Flag to cancel agent execution
        }
        
        # Store reference in MainScreen for access in handle_input
        self._agent_input_state = agent_input_state
        
        # Monkey-patch Prompt.ask to use chat-based input
        
        def tui_prompt_ask(prompt="", **kwargs):
            """Replacement for Prompt.ask that uses chat-based input."""
            import re
            
            # Check if agent was cancelled before prompting
            if agent_input_state.get('cancelled', False):
                raise KeyboardInterrupt("Agent cancelled by user")
            
            # Clean up rich markup from prompt
            clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
            
            # Check if there are choices
            choices = kwargs.get('choices', None)
            if choices:
                choices_text = ", ".join(f"`{c}`" for c in choices)
                full_prompt = f"**{clean_prompt}**\n\nChoices: {choices_text}"
            else:
                full_prompt = f"**{clean_prompt}**"
            
            # Show prompt in chat
            self.app.call_from_thread(
                chat_panel.add_message,
                "assistant",
                full_prompt
            )
            
            # Set state to waiting
            agent_input_state['waiting_for_input'] = True
            agent_input_state['prompt_text'] = clean_prompt
            agent_input_state['result'] = None
            agent_input_state['ready'].clear()
            
            # Wait for user input
            agent_input_state['ready'].wait()
            
            # Check if cancelled while waiting
            if agent_input_state.get('cancelled', False):
                agent_input_state['waiting_for_input'] = False
                raise KeyboardInterrupt("Agent cancelled by user")
            
            # Get result and reset state
            result = agent_input_state['result']
            agent_input_state['waiting_for_input'] = False
            
            return result if result else ""
        
        def tui_prompt_session_prompt(self, *args, **kwargs):
            """Replacement for PromptSession.prompt that uses chat-based input."""
            prompt_text = args[0] if args else "Enter valuesss"
            return tui_prompt_ask(prompt_text, **kwargs)
        
        # Create a TUI Console wrapper class that captures the console_wrapper in closure
        _captured_console_wrapper = console_wrapper
        main_app = self.app  # Capture the ValbotTUI app instance (self is MainScreen, self.app is ValbotTUI)
        _original_console_class = original_console_class  # Capture original Console before patching
        main_screen = self  # Capture MainScreen instance to access _agent_input_state dynamically
        
        class TUIConsole:
            """A Console replacement that forwards to our TUI console wrapper."""
            def __init__(self, *args, **kwargs):
                # Capture the console wrapper for this instance
                self._wrapper = _captured_console_wrapper
                self._main_screen = main_screen  # Store MainScreen to access current _agent_input_state
                self._chat_panel = chat_panel  # Capture chat panel
                self._main_app = main_app  # Capture main ValbotTUI instance (outer 'self')
                self._original_console = _original_console_class  # Store original Console class
            
            @property
            def _agent_input_state(self):
                """Get the current agent_input_state from MainScreen dynamically."""
                return self._main_screen._agent_input_state if hasattr(self._main_screen, '_agent_input_state') else None
            
            def print(self, *args, **kwargs):
                # Check if printing a Rich Table with menu choices
                from rich.table import Table
                if args and isinstance(args[0], Table):
                    table = args[0]
                    
                    # Try to extract table data by rendering it
                    choices = []
                    try:
                        # Use the ORIGINAL Rich Console (not our patched one) to render the table
                        from io import StringIO
                        
                        # Create a temporary console using the original Console class
                        string_buffer = StringIO()
                        temp_console = self._original_console(file=string_buffer, force_terminal=False, width=80, legacy_windows=False)
                        temp_console.print(table)
                        rendered_output = string_buffer.getvalue()
                        
                        # Parse the rendered table to extract rows
                        # Look for lines that match the pattern: "│ key │ description │"
                        import re
                        lines = rendered_output.split('\n')
                        for line in lines:
                            # Match table rows (lines with │ separators)
                            if '│' in line and not any(char in line for char in ['─', '┌', '┐', '└', '┘', '├', '┤', '┬', '┴', '┼']):
                                # Split by │ and extract cells
                                cells = [cell.strip() for cell in line.split('│') if cell.strip()]
                                if len(cells) >= 2:
                                    # First cell is key, second is description
                                    key = cells[0].strip()
                                    description = cells[1].strip()
                                    if key and description:
                                        choices.append((description, key))
                        
                        if choices and self._agent_input_state:
                            self._agent_input_state['last_menu_choices'] = choices
                        
                        # Don't print the rendered table - it will be shown as InlinePicker
                        return
                    except Exception as e:
                        # Fallback - just pass through
                        pass
                
                # For non-table content, pass through to wrapper
                return self._wrapper.print(*args, **kwargs)
            
            def rule(self, *args, **kwargs):
                return self._wrapper.rule(*args, **kwargs)
            
            def status(self, *args, **kwargs):
                return self._wrapper.status(*args, **kwargs)
            
            def input(self, prompt="", **kwargs):
                """Handle input requests from agents - use InlinePicker for menus, text input otherwise."""
                import re
                import threading
                # Clean up rich markup from prompt
                clean_prompt = re.sub(r'\[/?[^\]]+\]', '', str(prompt))
                
                # Check if we have menu choices from a recently printed table
                menu_choices = self._agent_input_state.get('last_menu_choices') if self._agent_input_state else None
                
                if menu_choices:
                    # Show InlinePicker for menu selection
                    if self._agent_input_state:
                        self._agent_input_state['last_menu_choices'] = None  # Clear after using
                    
                    # Use a container to store the result
                    result_container = {'value': None, 'ready': threading.Event()}
                    
                    # Define the callback that will be triggered when user selects
                    def on_picker_selection(selected_value):
                        result_container['value'] = selected_value
                        result_container['ready'].set()
                    
                    # Call the show_inline_picker method from the main thread
                    # This uses the proper async flow that MainScreen already has
                    import asyncio
                    
                    async def show_picker_wrapper():
                        main_screen = self._main_app.screen
                        # Set up the callback before showing picker
                        main_screen._picker_callback = on_picker_selection
                        # Show the picker
                        await main_screen.show_inline_picker(
                            title=clean_prompt,
                            options=menu_choices,
                            picker_type="agent_menu"
                        )
                    
                    # Schedule the async function on the main event loop
                    self._main_app.call_from_thread(
                        lambda: asyncio.create_task(show_picker_wrapper())
                    )
                    
                    # Wait for user to make a selection
                    result_container['ready'].wait()
                    return result_container['value'] if result_container['value'] else ""
                else:
                    # No menu - use text input
                    if not self._agent_input_state:
                        return ""  # Safety check - no state available
                    
                    # Show prompt in chat
                    self._main_app.call_from_thread(
                        self._chat_panel.add_message,
                        "assistant",
                        f"**{clean_prompt}**"
                    )
                    
                    # Set state to waiting
                    self._agent_input_state['waiting_for_input'] = True
                    self._agent_input_state['prompt_text'] = clean_prompt
                    self._agent_input_state['result'] = None
                    self._agent_input_state['ready'].clear()
                    
                    # Wait for user input
                    self._agent_input_state['ready'].wait()
                    
                    # Get result and reset state
                    result = self._agent_input_state['result']
                    self._agent_input_state['waiting_for_input'] = False
                    return result if result is not None else ""
            
            def log(self, *args, **kwargs):
                # Treat log same as print
                return self._wrapper.print(*args, **kwargs)
            
            @property
            def width(self):
                return self._wrapper.width
            
            @property
            def height(self):
                return self._wrapper.height
            
            def is_terminal(self):
                return self._wrapper.is_terminal()
            
            def set_live(self, live_instance):
                """Stub for Rich Live compatibility - we'll intercept Live context manager instead."""
                return True
            
            def clear_live(self):
                """Stub for Rich Live compatibility."""
                pass
            
            def show_cursor(self, show=True):
                """Stub for Rich cursor control - not needed in TUI."""
                pass
            
            def push_render_hook(self, hook):
                """Stub for Rich render hooks - not needed in TUI."""
                pass
            
            def pop_render_hook(self):
                """Stub for Rich render hooks - not needed in TUI."""
                pass
            
            def __enter__(self):
                """Context manager support for Rich Live."""
                return self
            
            def __exit__(self, exc_type, exc_val, exc_tb):
                """Context manager support for Rich Live."""
                return False
        
        # Monkey patch Rich's Live class to intercept interactive menus
        from rich.live import Live as OriginalLive
        
        class TUILive:
            """Intercepts Rich Live displays and converts them to TUI InlinePicker."""
            def __init__(self, renderable=None, *, console=None, **kwargs):
                self.renderable = renderable
                self.console = console
                self.kwargs = kwargs
                self._entered = False
            
            def __enter__(self):
                self._entered = True
                # Don't actually start Live display in TUI
                return self
            
            def __exit__(self, exc_type, exc_val, exc_tb):
                self._entered = False
                return False
            
            def start(self, refresh=True):
                pass
            
            def stop(self):
                pass
            
            def update(self, renderable):
                pass
        
        # Apply monkey patches
        Prompt.ask = staticmethod(tui_prompt_ask)
        PromptSession.prompt = tui_prompt_session_prompt
        
        # Replace Console class entirely with our wrapper
        import rich.console
        import rich.live
        original_console_class = rich.console.Console
        rich.console.Console = TUIConsole
        
        # Replace Live class with our interceptor
        rich.live.Live = TUILive
        
        # Monkey patch readchar to force fallback to console.input()
        try:
            import readchar
            original_readkey = readchar.readkey
            def tui_readkey():
                """Raise exception to force agent to use console.input() fallback."""
                raise RuntimeError("readchar not available in TUI - using console.input() fallback")
            readchar.readkey = tui_readkey
        except ImportError:
            pass  # readchar not installed, that's fine
        
        # Patch AgentRunResult to add .data property as alias for .output (for backward compatibility)
        try:
            from pydantic_ai import AgentRunResult
            if not hasattr(AgentRunResult, 'data'):
                # Add data property that returns output
                AgentRunResult.data = property(lambda self: self.output)
        except (ImportError, AttributeError):
            pass  # If AgentRunResult doesn't exist or is different, skip
        
        # Also patch Console and Live in any modules that may have already imported them
        import sys
        for module_name, module in list(sys.modules.items()):
            if module:
                # Patch Console
                if hasattr(module, 'Console'):
                    try:
                        if module.Console is original_console_class:
                            module.Console = TUIConsole
                    except (AttributeError, TypeError):
                        pass
                # Patch Live
                if hasattr(module, 'Live'):
                    try:
                        module.Live = TUILive
                    except (AttributeError, TypeError):
                        pass
        
        # Add starting message
        self.app.call_from_thread(
            chat_panel.add_message, 
            "system", 
            f"{EMOJI['robot']} Running agent: **{agent_name}**\n\n[dim]{agent_desc}[/dim]"
        )
        
        # Set the current agent in the chat panel
        self.app.call_from_thread(setattr, chat_panel, 'current_agent', agent_name)
        
        # Add agent chip to context pill bar
        try:
            context_chip_bar = self.query_one("#context-chip-bar", ContextChipBar)
            self.app.call_from_thread(context_chip_bar.add_agent, agent_name)
        except Exception:
            pass
        
        # CRITICAL: Monkey-patch the SPF2PDL agent's get_initializer_args method
        # to completely bypass its interactive menu system
        original_spf2pdl_get_init = None
        try:
            from agent_plugins.agents import pdl_agent
            original_spf2pdl_get_init = pdl_agent.MyCustomPlugin.get_initializer_args
            
            def tui_get_initializer_args_replacement(agent_self):
                """TUI-compatible version that returns mode from initializer_args without prompting."""
                # Simply return what was already set via kwargs in __init__
                mode = agent_self.initializer_args.get('mode', '2')
                result = {'mode': mode}
                
                # If mode is 1 and input_path not provided, we need it
                if mode == '1' and not agent_self.initializer_args.get('input_path'):
                    result['input_path'] = agent_self.initializer_args.get('input_path', '')
                
                return result
            
            # Apply the monkey patch
            pdl_agent.MyCustomPlugin.get_initializer_args = tui_get_initializer_args_replacement
            self._original_get_initializer_args = original_spf2pdl_get_init
        except (ImportError, AttributeError) as e:
            pass
        
        try:
            # Run the agent with proper context (pass init args if provided)
            # All console.print() calls from the agent will now appear as assistant messages
            kwargs = agent_init_args or {}
            
            # Monkey-patch PluginManager.run_plugin to replace agent console after creation
            original_run_plugin = self.chatbot.plugin_manager.run_plugin
            tui_console_instance = TUIConsole()  # Create our TUI console
            
            def patched_run_plugin(plugin_name, context, **run_kwargs):
                """Patched run_plugin that replaces agent's console with TUIConsole."""
                def normalize(name):
                    return name.replace('_', '').replace(' ', '').lower()

                normalized_input = normalize(plugin_name)
                matched_plugin_class = None
                for key in self.chatbot.plugin_manager.plugins:
                    if normalize(key) == normalized_input:
                        matched_plugin_class = self.chatbot.plugin_manager.plugins[key]
                        break
                if matched_plugin_class:
                    plugin_instance = matched_plugin_class(self.chatbot.plugin_manager.model, **run_kwargs)
                    
                    # CRITICAL: Replace the agent's console BEFORE running it
                    if hasattr(plugin_instance, 'console'):
                        plugin_instance.console = tui_console_instance
                    
                    plugin_instance.run_agent_flow_with_init_args(context, **run_kwargs)
                else:
                    print(f"Plugin {plugin_name} not found.")
            
            # Apply the patch
            self.chatbot.plugin_manager.run_plugin = patched_run_plugin
            
            # Now run the plugin
            self.chatbot.plugin_manager.run_plugin(
                agent_name, 
                self.chatbot.context_manager.conversation_history,
                **kwargs
            )
            
            # Restore original
            self.chatbot.plugin_manager.run_plugin = original_run_plugin
            
            # Check if cancelled after completion
            if agent_input_state.get('cancelled', False):
                # Add cancellation message
                self.app.call_from_thread(
                    chat_panel.add_message, 
                    "system", 
                    f"{EMOJI['cross']} Agent **{agent_name}** was cancelled."
                )
            else:
                # Add success message
                self.app.call_from_thread(
                    chat_panel.add_message, 
                    "assistant", 
                    f"{EMOJI['checkmark']} Agent **{agent_name}** completed successfully!"
                )
        except KeyboardInterrupt:
            # Agent was cancelled by user
            self.app.call_from_thread(
                chat_panel.add_message, 
                "system", 
                f"{EMOJI['cross']} Agent: **{agent_name}** was cancelled by user."
            )
        except Exception as e:
            import traceback
            
            # Add error message with traceback for debugging
            error_details = f"{EMOJI['cross']} Agent error: {str(e)}\n\n```\n{traceback.format_exc()}\n```"
            self.app.call_from_thread(
                chat_panel.add_message, 
                "error", 
                error_details
            )
        finally:
            # Restore SPF2PDL agent's get_initializer_args if it was patched
            try:
                from agent_plugins.agents import pdl_agent
                if hasattr(self, '_original_get_initializer_args'):
                    pdl_agent.MyCustomPlugin.get_initializer_args = self._original_get_initializer_args
                    delattr(self, '_original_get_initializer_args')
            except (ImportError, AttributeError):
                pass
            
            # Clear the current agent in the chat panel
            self.app.call_from_thread(setattr, chat_panel, 'current_agent', None)
            
            # Remove agent chip from context pill bar
            try:
                context_chip_bar = self.query_one("#context-chip-bar", ContextChipBar)
                self.app.call_from_thread(context_chip_bar.remove_agent)
            except Exception:
                pass
            
            # Clean up agent input state (but don't delete it - just reset)
            if hasattr(self, '_agent_input_state'):
                self._agent_input_state['waiting_for_input'] = False
                self._agent_input_state['last_menu_choices'] = None
                self._agent_input_state['cancelled'] = False  # Reset cancellation flag
                # Keep _agent_input_state around for next agent run
            
            # Restore original prompt functions and Console class
            Prompt.ask = original_prompt_ask
            PromptSession.prompt = original_prompt_session_prompt
            import rich.console
            rich.console.Console = original_console_class
            
            # Restore Console in modules that were patched
            import sys
            for module_name, module in list(sys.modules.items()):
                if module and hasattr(module, 'Console'):
                    try:
                        if module.Console is TUIConsole:
                            module.Console = original_console_class
                    except (AttributeError, TypeError):
                        pass
            
            # Clean up the event loop
            try:
                if not loop.is_closed():
                    loop.close()
            except Exception:
                pass
    
    def _run_create_database_in_thread(self, args: str):
        """Create a new RAG database (runs in background thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        try:
            # Strip # markers from file paths (used for autocomplete)
            args = self._strip_file_markers(args)
            
            # Parse arguments: first is output folder, rest are file paths
            parts = args.split()
            if len(parts) < 2:
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "error",
                    f"{EMOJI['cross']} Invalid arguments. Usage: `/create_database <output_folder> <file1> [file2] ...`"
                )
                return
            
            output_folder = parts[0]
            file_paths_input = parts[1:]
            
            # Show initial message
            total_files = len(file_paths_input)
            self.app.call_from_thread(
                chat_panel.add_message,
                "system",
                f"""### {EMOJI['lightbulb']} Creating Database

**Output folder**: `{output_folder}`
**Input files**: {total_files} file(s)"""
            )
            
            # Create a custom progress widget
            # Calculate total steps: each file has ~5 sub-steps (read, parse, chunk, embed, store)
            # This gives granular progress instead of jumping per file
            STEPS_PER_FILE = 100  # Treat each file as 100 steps for smooth progress
            total_steps = total_files * STEPS_PER_FILE
            progress_widget = None
            
            def create_progress_ui():
                nonlocal progress_widget
                
                # Create and mount the progress widget with total steps
                progress_widget = DatabaseProgressWidget(total=total_steps)
                chat_panel.mount(progress_widget)
                chat_panel.scroll_end(animate=False)
            
            self.app.call_from_thread(create_progress_ui)
            
            # Check dependencies
            missing = []
            if PyPDF2 is None:
                missing.append("PyPDF2")
            if SentenceTransformer is None:
                missing.append("sentence-transformers")
            if chromadb is None:
                missing.append("chromadb")
            
            if missing:
                error_msg = f"""### {EMOJI['cross']} Missing Dependencies

The following packages are required but not installed:
{chr(10).join(f'- `{pkg}`' for pkg in missing)}

**Install with**:
```bash
pip install {' '.join(missing)}
```"""
                self.app.call_from_thread(chat_panel.add_message, "error", error_msg)
                return
            
            # Prepare output folder path (don't create yet - wait until we have valid files)
            output_path = Path(output_folder)
            folder_created = False
            rag_kb = None
            
            # Process each file
            processed_count = 0
            failed_count = 0
            current_step = 0  # Track overall progress in steps
            
            for idx, file_path_str in enumerate(file_paths_input, 1):
                file_path = Path(file_path_str.strip())
                
                # If path is not absolute, resolve it relative to current working directory
                if not file_path.is_absolute():
                    file_path = file_path.resolve()
                
                if not file_path.exists():
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "error",
                        f"{EMOJI['cross']} File not found: `{file_path_str}`\n\nSearched at: `{file_path}`\n\nMake sure the file exists or provide the full absolute path."
                    )
                    failed_count += 1
                    # Update progress widget - skip this file's steps
                    current_step += STEPS_PER_FILE
                    if progress_widget:
                        def update_failed():
                            progress_widget.update_progress(current_step, f"✗ Failed: `{file_path.name}` ({idx}/{total_files}, {failed_count} errors)")
                        self.app.call_from_thread(update_failed)
                    continue
                
                # Update status - starting file processing
                if progress_widget:
                    def update_starting():
                        progress_widget.update_progress(current_step, f"{EMOJI['hourglass']} Starting: `{file_path.name}` ({idx}/{total_files})")
                    self.app.call_from_thread(update_starting)
                
                try:
                    suffix = file_path.suffix.lower()
                    
                    # Handle .gz compressed files - detect the underlying file type
                    actual_suffix = suffix
                    if suffix == '.gz':
                        # Get the extension before .gz (e.g., .pdf from .pdf.gz)
                        file_stem = file_path.stem  # filename without last extension
                        if '.' in file_stem:
                            actual_suffix = Path(file_stem).suffix.lower()
                        else:
                            # Just .gz with no underlying extension, treat as text
                            actual_suffix = '.txt'
                    
                    # Check if this is an invalid file type
                    if actual_suffix in ['.xlsx']:
                        self.app.call_from_thread(
                            chat_panel.add_message,
                            "error",
                            f"{EMOJI['cross']} Unsupported file type: `{actual_suffix}` for file `{file_path.name}`\n\nThis file type cannot be processed by the RAG database."
                        )
                        failed_count += 1
                        current_step += STEPS_PER_FILE
                        if progress_widget:
                            def update_failed():
                                progress_widget.update_progress(current_step, f"✗ Skipped: `{file_path.name}` ({idx}/{total_files}, {failed_count} errors)")
                            self.app.call_from_thread(update_failed)
                        continue
                    
                    # Create folder and initialize RAG database on first valid file
                    if not folder_created:
                        output_path.mkdir(parents=True, exist_ok=True)
                        folder_created = True
                        
                        # Create a console wrapper to capture RAG output to TUI
                        import io
                        from rich.console import Console as RichConsole
                        
                        # Initialize RAG database with a StringIO console to suppress rich output
                        string_buffer = io.StringIO()
                        rag_console = RichConsole(file=string_buffer, force_terminal=True)
                        
                        rag_kb = RAGKnowledgeBase(
                            pdf_dir=output_path,
                            cache_dir=output_path / ".rag_cache",
                            chunk_size=1000,
                            chunk_overlap=200,
                            collection_name="rag_knowledge_base"
                        )
                        # Override the console to suppress rich output in thread
                        rag_kb.console = rag_console
                    
                    # Helper function to update progress with sub-step info
                    def update_substep(substep_num, substep_desc):
                        """Update progress for a sub-step within file processing."""
                        nonlocal current_step
                        # Each sub-step is 20% of STEPS_PER_FILE (5 sub-steps total)
                        current_step = (idx - 1) * STEPS_PER_FILE + (substep_num * 20)
                        if progress_widget:
                            def update_ui():
                                progress_widget.update_progress(
                                    current_step, 
                                    f"{EMOJI['hourglass']} {substep_desc}: `{file_path.name}` ({idx}/{total_files})"
                                )
                            self.app.call_from_thread(update_ui)
                        # Add a small delay to make progress visible
                        import time
                        time.sleep(0.15)
                    
                    # Sub-step 1: Reading file
                    update_substep(1, "Reading")
                    
                    # Process file directly from its location (no copying needed)
                    # Use actual_suffix to determine file type (handles .gz transparently)
                    if actual_suffix == '.pdf':
                        # Sub-step 2: Parsing PDF
                        update_substep(2, "Parsing PDF")
                        
                        # Sub-step 3: Chunking text
                        update_substep(3, "Chunking text")
                        
                        # Sub-step 4: Generating embeddings
                        update_substep(4, "Generating embeddings")
                        
                        rag_kb.load_pdfs([str(file_path)], force_reload=True)
                        
                        # Sub-step 5: Storing in database
                        update_substep(5, "Storing in DB")
                        
                        processed_count += 1
                    elif actual_suffix == '.docx':
                        # Sub-step 2: Parsing DOCX
                        update_substep(2, "Parsing DOCX")
                        
                        # Sub-step 3: Chunking text
                        update_substep(3, "Chunking text")
                        
                        # Sub-step 4: Generating embeddings
                        update_substep(4, "Generating embeddings")
                        
                        rag_kb.load_docx_files([str(file_path)], force_reload=True)
                        
                        # Sub-step 5: Storing in database
                        update_substep(5, "Storing in DB")
                        
                        processed_count += 1
                    elif actual_suffix in ['.txt', '.pdl', '.md', '.py', '.c', '.cpp', '.h']:
                        # Sub-step 2: Reading text
                        update_substep(2, "Reading text")
                        
                        # Sub-step 3: Chunking text
                        update_substep(3, "Chunking text")
                        
                        # Sub-step 4: Generating embeddings
                        update_substep(4, "Generating embeddings")
                        
                        rag_kb.load_text_file(file_path, force_reload=True)
                        
                        # Sub-step 5: Storing in database
                        update_substep(5, "Storing in DB")
                        
                        processed_count += 1
                    else:
                        # Try to process as text file (default fallback for unknown types)
                        # Sub-step 2: Reading text
                        update_substep(2, "Reading text")
                        
                        # Sub-step 3: Chunking text
                        update_substep(3, "Chunking text")
                        
                        # Sub-step 4: Generating embeddings
                        update_substep(4, "Generating embeddings")
                        
                        rag_kb.load_text_file(file_path, force_reload=True)
                        
                        # Sub-step 5: Storing in database
                        update_substep(5, "Storing in DB")
                        
                        processed_count += 1
                    
                    # Move to end of this file's progress
                    current_step = idx * STEPS_PER_FILE
                    
                    # Update progress widget after successful processing
                    if progress_widget:
                        def update_success():
                            progress_widget.update_progress(current_step, f"✓ Completed: `{file_path.name}` ({idx}/{total_files})")
                        self.app.call_from_thread(update_success)
                    
                except Exception as e:
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "error",
                        f"{EMOJI['cross']} Error processing `{file_path.name}`: {str(e)}"
                    )
                    failed_count += 1
                    
                    # Update progress widget after error - skip remaining steps for this file
                    current_step = idx * STEPS_PER_FILE
                    if progress_widget:
                        def update_error():
                            progress_widget.update_progress(current_step, f"✗ Error: `{file_path.name}` ({idx}/{total_files}, {failed_count} errors)")
                        self.app.call_from_thread(update_error)
            
            # Remove progress widget
            if progress_widget:
                self.app.call_from_thread(progress_widget.remove)
            
            # Display summary only if at least one file was successfully processed
            if processed_count > 0:
                stats = rag_kb.get_stats()
                summary = f"""### {EMOJI['checkmark']} Database Created

**Summary**:
- Files processed: {processed_count}
- Files failed: {failed_count}
- Total chunks: {stats['total_chunks']}
- Embedding model: `{stats['embedding_model']}`
- Location: `{output_folder}`

You can now load this database with:
```
/load_database {output_folder}
```"""
                
                self.app.call_from_thread(chat_panel.add_message, "system", summary)
            # elif failed_count > 0:
            #     # All files failed - show appropriate message
            #     self.app.call_from_thread(
            #         chat_panel.add_message,
            #         "error",
            #         f"{EMOJI['cross']} No files were successfully processed. All {failed_count} file(s) failed or were unsupported."
            #     )
            
        except Exception as e:
            # Remove progress widget on error
            if 'progress_widget' in locals() and progress_widget:
                self.app.call_from_thread(progress_widget.remove)
            
            error_msg = f"""### {EMOJI['cross']} Error Creating database

```
{str(e)}
```

{traceback.format_exc()}"""
            self.app.call_from_thread(chat_panel.add_message, "error", error_msg)
    
    def _run_load_database_in_thread(self, kb_directories: str):
        """Load one or more existing RAG databases and enable Q&A mode (runs in background thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        try:
            # Parse space-separated database directories
            db_paths = kb_directories.strip().split()
            
            if not db_paths:
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "error",
                    f"{EMOJI['cross']} No database directories specified. Usage: `/load_database <dir1> [dir2] ...`"
                )
                return
            
            # Show initial message
            self.app.call_from_thread(
                chat_panel.add_message,
                "system",
                f"""### {EMOJI['lightbulb']} Loading Database(s)

Loading {len(db_paths)} database(s)...

Please wait..."""
            )
            
            # Check dependencies
            missing = []
            if PyPDF2 is None:
                missing.append("PyPDF2")
            if SentenceTransformer is None:
                missing.append("sentence-transformers")
            if chromadb is None:
                missing.append("chromadb")
            
            if missing:
                error_msg = f"""### {EMOJI['cross']} Missing Dependencies

The following packages are required but not installed:
{chr(10).join(f'- `{pkg}`' for pkg in missing)}

**Install with**:
```bash
pip install {' '.join(missing)}
```"""
                self.app.call_from_thread(chat_panel.add_message, "error", error_msg)
                return
            
            # Initialize _active_rag_kb as dict if not already
            if not hasattr(self, '_active_rag_kb') or self._active_rag_kb is None:
                self._active_rag_kb = {}
            elif not isinstance(self._active_rag_kb, dict):
                # Convert old single instance to dict
                self._active_rag_kb = {}
            
            # Track success/failure (using built-in RAGKnowledgeBase class)
            loaded_count = 0
            failed_count = 0
            all_stats = []
            
            # Load each database
            for kb_directory in db_paths:
                kb_directory = kb_directory.strip()
                
                try:
                    # Check if directory exists
                    kb_path = Path(kb_directory)
                    if not kb_path.exists():
                        self.app.call_from_thread(
                            chat_panel.add_message,
                            "error",
                            f"{EMOJI['cross']} Directory does not exist: `{kb_directory}`"
                        )
                        failed_count += 1
                        continue
                    
                    cache_dir = kb_path / ".rag_cache"
                    if not cache_dir.exists():
                        self.app.call_from_thread(
                            chat_panel.add_message,
                            "error",
                            f"{EMOJI['cross']} No database cache found in `{kb_directory}`\n\nThis directory may not contain a valid database."
                        )
                        failed_count += 1
                        continue
                    
                    # Initialize RAG database
                    rag_kb = RAGKnowledgeBase(
                        pdf_dir=kb_path,
                        cache_dir=cache_dir,
                        chunk_size=1000,
                        chunk_overlap=200,
                        collection_name="rag_knowledge_base"
                    )
                    
                    # Store in dictionary for query handling
                    self._active_rag_kb[kb_directory] = rag_kb
                    
                    # Add database to chip bar
                    self.app.call_from_thread(self._context_chip_bar.add_database, kb_directory)
                    
                    # Get stats
                    stats = rag_kb.get_stats()
                    stats['location'] = kb_directory
                    all_stats.append(stats)
                    
                    loaded_count += 1
                    
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "system",
                        f"{EMOJI['checkmark']} Loaded: `{kb_directory}` ({stats['total_chunks']} chunks)"
                    )
                    
                except Exception as e:
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "error",
                        f"{EMOJI['cross']} Error loading `{kb_directory}`: {str(e)}"
                    )
                    failed_count += 1
            
            # Display summary
            if loaded_count > 0:
                # Build stats summary
                stats_lines = []
                for stats in all_stats:
                    stats_lines.append(f"- `{stats['location']}`: {stats['total_chunks']} chunks")
                
                success_msg = f"""### {EMOJI['checkmark']} Database(s) Loaded

**Summary**:
- Successfully loaded: {loaded_count}
- Failed: {failed_count}
- Total chunks across all databases: {sum(s['total_chunks'] for s in all_stats)}

**Loaded databases**:
{chr(10).join(stats_lines)}

**You can now ask questions about your documents!**

Just type your questions normally in the chat. The system will automatically search all loaded databases and provide answers based on the document content.

To unload a specific database, click the ✕ on its chip above.
To unload all databases and return to normal chat, use `/new`.
"""
                
                self.app.call_from_thread(chat_panel.add_message, "system", success_msg)
            else:
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "error",
                    f"{EMOJI['cross']} Failed to load any databases. Check the errors above."
                )
            
        except Exception as e:
            error_msg = f"""### {EMOJI['cross']} Error Loading Database

```
{str(e)}
```

{traceback.format_exc()}"""
            self.app.call_from_thread(chat_panel.add_message, "error", error_msg)
    
    def _handle_rag_query(self, question: str, additional_context: str = ""):
        """Handle a query to the RAG database(s) (runs in background thread).
        
        Args:
            question: The user's question
            additional_context: Optional additional context (e.g., from local files)
        """
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        try:
            # _active_rag_kb is now a dict of {db_path: RAGKnowledgeBase}
            if not isinstance(self._active_rag_kb, dict) or not self._active_rag_kb:
                # Fallback - shouldn't happen but just in case
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "error",
                    f"{EMOJI['cross']} No databases loaded."
                )
                return
            
            num_dbs = len(self._active_rag_kb)
            db_label = "database" if num_dbs == 1 else f"{num_dbs} databases"
            
            # Show that we're searching
            self.app.call_from_thread(
                chat_panel.add_message,
                "system",
                f"{EMOJI['lightbulb']} Searching {db_label}..."
            )
            
            # Query all loaded databases and combine results
            all_contexts = []
            total_retrieved_chunks = 0
            
            for db_path, rag_kb in self._active_rag_kb.items(): 
                try:
                    # Check stats first
                    stats = rag_kb.get_stats()
                    if stats['total_chunks'] == 0:
                        self.app.call_from_thread(
                            chat_panel.add_message,
                            "system",
                            f"⚠️ Database `{Path(db_path).name}` is empty (0 chunks)."
                        )
                        continue

                    # Retrieve relevant chunks
                    # We use a generous top_k because we want to capture enough context
                    chunks = rag_kb.retrieve_relevant_context(
                        query=question,
                        top_k=20
                    )
                    
                    if not chunks:
                        self.app.call_from_thread(
                            chat_panel.add_message,
                            "system",
                            f"ℹ️ No relevant chunks found in `{Path(db_path).name}`."
                        )
                        continue
                        
                    # Format context
                    context_parts = []
                    for chunk in chunks:
                        meta = chunk['metadata']
                        source = meta.get('source', 'Unknown')
                        sheet = meta.get('sheet', '')
                        type_ = meta.get('type', '')
                        
                        header = f"Source: {source}"
                        if sheet:
                            header += f" | Sheet: {sheet}"
                        if type_ == 'summary':
                            header += " | [SUMMARY]"
                            
                        context_parts.append(f"[{header}]\n{chunk['text']}")
                    
                    db_context = "\n\n".join(context_parts)
                    
                    # Add database source to context
                    all_contexts.append(f"### From database: {Path(db_path).name}\n\n{db_context}")
                    
                    total_retrieved_chunks += len(chunks)
                    
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "system",
                        f"{EMOJI['database']} Retrieved {len(chunks)} chunks from `{Path(db_path).name}`"
                    )
                    
                except Exception as e:
                    # Log error but continue with other databases
                    self.app.call_from_thread(
                        chat_panel.add_message,
                        "system",
                        f"{EMOJI['cross']} Error searching `{db_path}`: {str(e)}"
                    )
            
            # Combine all contexts
            if not all_contexts:
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "assistant",
                    "I couldn't find any relevant information in the loaded databases for your question. Please try rephrasing or ask a different question."
                )
                return
            
            relevant_context = "\n\n---\n\n".join(all_contexts)
            
            # Add any additional context (e.g., from local files)
            if additional_context:
                relevant_context = f"{additional_context}\n\n---\n\n{relevant_context}"
            
            # Truncate context if it's too large to prevent context window overflow
            # Most models have 128K token limit, but conversation history also counts
            # Reserve space for conversation history (~20K tokens) and response (~4K tokens)
            MAX_CONTEXT_TOKENS = 60000  # ~240K characters (safe limit for 128K token models)
            max_context_chars = MAX_CONTEXT_TOKENS * 4
            
            if len(relevant_context) > max_context_chars:
                self.app.call_from_thread(
                    chat_panel.add_message,
                    "system",
                    f"⚠️ Context too large ({len(relevant_context):,} chars), truncating to {max_context_chars:,} chars to fit model limits..."
                )
                relevant_context = relevant_context[:max_context_chars] + "\n\n[... Context truncated due to size ...]"
            
            # Generate answer using the chatbot with the retrieved context
            qa_prompt = f"""User Question: {question}

Relevant Context from Documents:
{relevant_context}

Please answer the user's question based on the context above.
Use proper Markdown formatting in your response.
Cite sources when possible."""
            
            # Set processing state
            self.app.call_from_thread(self._set_processing_state, True)
            self.app.call_from_thread(self._add_loading_message)
            
            # Send the prompt to the chatbot
            self._send_standard_message_sync(qa_prompt)
            
        except Exception as e:
            error_msg = f"""### {EMOJI['cross']} Error Querying Database

```
{str(e)}
```

{traceback.format_exc()}"""
            self.app.call_from_thread(chat_panel.add_message, "error", error_msg)
        finally:
            self.app.call_from_thread(self._set_processing_state, False)
    
    def _send_chat_message_in_thread(self, message: str):
        """Send a message to the chatbot (runs in background thread).
        
        File Handling:
        - When files are referenced in chat, they are detected and loaded into
          context_manager using context_management.py
        - Files are read with chonkie chunking and proper formatting
        - File content persists in conversation history for follow-up questions
        - No separate tool-based file reading to avoid confusion
        """
        # Reset cancellation flag at the start of a new message
        self._cancel_streaming = False
        
        if not self.chatbot:
            self.app.call_from_thread(self._add_error_message, f"""### {EMOJI['cross']} ChatBot Not Initialized

ChatBot is not ready. Please check your configuration.

**Troubleshooting**:
1. Verify your API keys are set
2. Check your model configuration
3. Ensure network connectivity
""")
            return
        
        # Check if the message references local files first (always check this)
        detected_files = []
        enhanced_message = message  # Will be enhanced if files are detected
        
        if self.chatbot.agent_model:
            detected_files = self.chatbot._detect_file_paths(message)
            
            # Add detected files to context chip bar AND load into context manager
            if detected_files:
                for file_path in detected_files:
                    self.app.call_from_thread(self._context_chip_bar.add_file, file_path)
                
                # Load files into context manager using context_management.py
                # This ensures files are read with chonkie chunking and proper formatting
                try:
                    self.chatbot.context_manager.load_context(detected_files, silent=True)
                    
                    # Enhance the message to clarify that detected words are file paths
                    # This helps the AI understand the context better
                    import os
                    cwd = os.getcwd()
                    file_list = ", ".join([f"`{f}`" for f in detected_files])
                    enhanced_message = f"{message}\n\n[Note: The following file(s) from the current directory ({cwd}) have been loaded into context: {file_list}]"
                    
                except Exception as e:
                    # Log error but don't fail the request
                    error_msg = f"⚠️ Warning: Failed to load file(s) into context: {str(e)}"
                    self.app.call_from_thread(self._add_system_message, error_msg)
        
        # Check if RAG database is active
        has_active_database = hasattr(self, '_active_rag_kb') and self._active_rag_kb
        
        # If local files are detected AND database is active, combine both
        if detected_files and has_active_database:
            # Files are already loaded into context manager above
            # Now handle RAG query which will include the file context automatically
            self._handle_rag_query(enhanced_message)
            return
        
        # Check if RAG database is active (and no local files) - handle with RAG only
        if has_active_database:
            self._handle_rag_query(enhanced_message)
            return
        
        # Set processing state in UI thread
        self.app.call_from_thread(self._set_processing_state, True)
        
        # Add loading indicator in chat
        self.app.call_from_thread(self._add_loading_message)
        
        try:
            # Use tool-enabled agent if available - let the AI decide when to use tools dynamically
            # Otherwise fall back to standard chat
            if self.chatbot.agent_model and self.chatbot.tool_agent:
                self._send_with_tools(enhanced_message)
            else:
                self._send_standard_message_sync(enhanced_message)
                
        except Exception as e:

            # Remove loading indicator on error
            self.app.call_from_thread(self._remove_loading_message)
            
            error_message = f"""### {EMOJI['cross']} Error

An error occurred while processing your message:

```
{str(e)}
```
"""
            self.app.call_from_thread(self._add_error_message, error_message)
        
        finally:
            # Reset processing state
            self.app.call_from_thread(self._set_processing_state, False)
    
    # Helper methods for thread-safe UI updates
    def _set_processing_state(self, processing: bool):
        """Set processing state (called from worker thread)."""
        self.processing = processing
        # Update status bar button
        try:
            status_bar = self.query_one("#status-bar", StatusBar)
            status_bar.processing = processing
        except Exception:
            pass
    
    def _add_loading_message(self):
        """Add loading message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        # Check if header should be shown
        show_header = not chat_panel._assistant_header_shown
        chat_panel._assistant_header_shown = True
        loading_msg = ChatMessage("assistant", "⠋ Thinking...", show_header=show_header)
        chat_panel.mount(loading_msg)
        chat_panel.scroll_end(animate=True)
        # Store reference for removal later
        self._loading_msg = loading_msg
        # Start animation
        self._start_loading_animation()
    
    def _remove_loading_message(self):
        """Remove loading message (called from worker thread)."""
        if hasattr(self, '_loading_msg') and self._loading_msg and self._loading_msg.parent:
            self._loading_msg.remove()
            self._loading_msg = None
            # Reset the header flag so the next message (streaming response) will show the header
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel._assistant_header_shown = False
        # Stop animation
        self._stop_loading_animation()
    
    def _start_loading_animation(self):
        """Start the loading spinner animation."""
        self._loading_animation_active = True
        self._loading_spinner_frames = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]
        self._loading_frame_index = 0
        # Schedule the animation update - store the timer handle so we can cancel it
        self._loading_timer = self.app.set_interval(0.1, self._update_loading_animation)
    
    def _stop_loading_animation(self):
        """Stop the loading spinner animation."""
        self._loading_animation_active = False
        # Cancel the timer if it exists
        if hasattr(self, '_loading_timer') and self._loading_timer:
            self._loading_timer.stop()
            self._loading_timer = None
    
    def _update_loading_animation(self):
        """Update the loading spinner frame."""
        if not self._loading_animation_active or not hasattr(self, '_loading_msg') or not self._loading_msg:
            return
        
        if self._loading_msg and self._loading_msg.parent:
            self._loading_frame_index = (self._loading_frame_index + 1) % len(self._loading_spinner_frames)
            spinner = self._loading_spinner_frames[self._loading_frame_index]
            # Update the content of the message widget (it's a Markdown widget)
            try:
                content_widget = self._loading_msg.query_one(".message-content", Markdown)
                content_widget.update(f"{spinner} Just a sec...")
            except Exception:
                # Fallback: update the message content directly
                self._loading_msg.content = f"{spinner} Just a sec..."
                self._loading_msg.refresh()
    
    def _add_error_message(self, message: str):
        """Add error message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        chat_panel.add_message("error", message)
    
    def _add_system_message(self, message: str):
        """Add system message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        chat_panel.add_message("system", message)
    
    def _add_assistant_message(self, message: str):
        """Add assistant message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        chat_panel.add_message("assistant", message)
        # Schedule copy button detection after UI settles (for non-streaming messages)
        try:
            if chat_panel.messages:
                last_msg = chat_panel.messages[-1]
                # Use ChatPanel's method for copy button detection
                self.app.call_later(chat_panel._add_copy_buttons_to_message, last_msg)
        except Exception:
            pass  # Ignore errors
        
        # Note: autosave is now handled in ChatPanel.add_message()
    
    def _update_loading_message(self, content: str):
        """Update loading message content (called from worker thread)."""
        if hasattr(self, '_loading_msg') and self._loading_msg and self._loading_msg.parent:
            self._loading_msg.content = content
            self._loading_msg.refresh()
    
    def _create_streaming_message(self, msg_ref: dict):
        """Create a new streaming message widget (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        # Check if header should be shown
        show_header = not chat_panel._assistant_header_shown
        chat_panel._assistant_header_shown = True
        
        # Create message widget with initial content
        streaming_msg = ChatMessage("assistant", msg_ref['content'], show_header=show_header)
        chat_panel.mount(streaming_msg)
        chat_panel.scroll_end(animate=False)  # Don't animate during streaming for better performance
        
        # Store reference
        msg_ref['msg'] = streaming_msg
    
    def _update_streaming_message(self, msg_ref: dict):
        """Update streaming message content (called from worker thread)."""
        if msg_ref['msg'] and msg_ref['msg'].parent:
            # Update the content using the proper update_content method
            msg_ref['msg'].update_content(msg_ref['content'])
            
            # Scroll to show new content
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.scroll_end(animate=False)
    
    def _finalize_streaming_message(self, msg_ref: dict):
        """Finalize streaming message after completion (called from worker thread)."""
        if msg_ref['msg'] and msg_ref['msg'].parent:
            # Final refresh to ensure all content is displayed
            msg_ref['msg'].update_content(msg_ref['content'])
            
            # Schedule copy button detection using the ChatPanel method
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            self.app.call_later(chat_panel._add_copy_buttons_to_message, msg_ref['msg'])
            
            # Final scroll
            chat_panel.scroll_end(animate=True)
            
            # Autosave the assistant message to the session log
            try:
                # Get agent name if available
                agent_name = msg_ref.get('agent_name', None)
                # Schedule autosave asynchronously
                asyncio.create_task(self._autosave_message("assistant", msg_ref['content'], agent_name))
            except Exception:
                pass  # Ignore errors
    
    def _send_with_tools(self, message: str):
        """Send a message using the tool-enabled pydantic-ai agent with streaming."""
        import asyncio
        import time
        
        # Check if tool agent is available
        if not hasattr(self.chatbot, 'tool_agent') or self.chatbot.tool_agent is None:
            self.app.call_from_thread(self._add_system_message, "⚠️ Tool agent not available. Using standard chat.")
            self._send_standard_message_sync(message)
            return
        
        try:
            # Keep loading indicator until first chunk arrives
            # (will be removed when first chunk triggers _create_streaming_message)
            
            # Create a streaming message widget with thread-safe accumulation
            streaming_msg_ref = {
                'msg': None, 
                'content': '', 
                'first_chunk': True,
                'last_update': 0,
                'update_interval': 0.05,  # Batch updates every 50ms for smoother rendering
                'cancelled': False  # Track if stream was cancelled
            }
            
            def stream_callback(chunk: str, is_complete: bool):
                """Callback for streaming updates from the agent."""
                # Check for cancellation at the start
                if self._cancel_streaming and not streaming_msg_ref['cancelled']:
                    streaming_msg_ref['cancelled'] = True
                    # Return True to signal cancellation to the chatbot
                    return True
                
                if not is_complete:
                    # Remove loading indicator on first chunk
                    if streaming_msg_ref['first_chunk']:
                        self.app.call_from_thread(self._remove_loading_message)
                        streaming_msg_ref['first_chunk'] = False
                    
                    # Accumulate content (thread-safe since we're in the same thread)
                    streaming_msg_ref['content'] += chunk
                    
                    # Batch updates to avoid overwhelming the UI with rapid chunks
                    current_time = time.time()
                    time_since_last = current_time - streaming_msg_ref['last_update']
                    
                    # Update or create message widget
                    if streaming_msg_ref['msg'] is None:
                        # Always create initial message immediately
                        self.app.call_from_thread(self._create_streaming_message, streaming_msg_ref)
                        streaming_msg_ref['last_update'] = current_time
                    elif time_since_last >= streaming_msg_ref['update_interval']:
                        # Throttle updates to improve performance
                        self.app.call_from_thread(self._update_streaming_message, streaming_msg_ref)
                        streaming_msg_ref['last_update'] = current_time
                    # If throttled, content will still accumulate and show on next update or finalization
                else:
                    # Streaming complete - ensure final update with all accumulated content
                    if streaming_msg_ref['msg'] is not None:
                        # Force final update to ensure all content is shown
                        self.app.call_from_thread(self._update_streaming_message, streaming_msg_ref)
                        # Finalize immediately - the content is already accumulated
                        self.app.call_from_thread(self._finalize_streaming_message, streaming_msg_ref)
                    # Remove loading indicator if it's still there (e.g., if no chunks were received)
                    self.app.call_from_thread(self._remove_loading_message)
                
                return False  # Continue streaming
            
            # Run the chatbot's send_message_with_tools method with streaming
            # Create a new event loop for this thread
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
            # Store loop and task references for cancellation
            self._current_tool_loop = loop
            self._current_tool_task = None
            
            try:
                # Create the task so we can cancel it
                self._current_tool_task = loop.create_task(
                    self.chatbot.send_message_with_tools(message, stream_callback=stream_callback)
                )
                # Run the task
                loop.run_until_complete(self._current_tool_task)
                
                # Check if we were cancelled during execution
                if streaming_msg_ref['cancelled']:
                    # Handle cancellation cleanup
                    self._handle_tool_cancellation(streaming_msg_ref)
                    
            except asyncio.CancelledError:
                # Task was cancelled - handle cleanup
                self._handle_tool_cancellation(streaming_msg_ref)
            finally:
                # Clean up references
                self._current_tool_task = None
                self._current_tool_loop = None
                loop.close()
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            
            # Remove loading indicator
            self.app.call_from_thread(self._remove_loading_message)
            
            error_msg = f"""### ⚠️ Tool Error

Error using tools: {str(e)}

Falling back to standard chat...

<details>
<summary>Error Details</summary>

```
{error_details}
```
</details>
"""
            self.app.call_from_thread(self._add_system_message, error_msg)
            
            # Remove the user message if it was added
            if self.chatbot.context_manager.conversation_history and \
               self.chatbot.context_manager.conversation_history[-1]["role"] == "user" and \
               self.chatbot.context_manager.conversation_history[-1]["content"] == message:
                self.chatbot.context_manager.conversation_history.pop()
            
            # Fall back to standard message
            self._send_standard_message_sync(message)
    
    def _handle_tool_cancellation(self, streaming_msg_ref: dict):
        """Handle cleanup when tool execution is cancelled."""
        # Remove loading indicator if still present
        if streaming_msg_ref['first_chunk']:
            self.app.call_from_thread(self._remove_loading_message)
        
        # Finalize any partial content that was streamed
        if streaming_msg_ref['msg'] is not None and streaming_msg_ref['content']:
            self.app.call_from_thread(self._update_streaming_message, streaming_msg_ref)
            self.app.call_from_thread(self._finalize_streaming_message, streaming_msg_ref)
            # Add partial response to conversation history
            self.chatbot.context_manager.conversation_history.append({
                "role": "assistant",
                "content": streaming_msg_ref['content']
            })
        
        # Show cancellation message
        self.app.call_from_thread(self._add_system_message, f"{EMOJI['cross']} *Response cancelled by user.*")
    
    def _send_standard_message_sync(self, message: str):
        """Send a message using standard OpenAI streaming (synchronous for worker thread)."""
        # Add to conversation history (if not already added)
        if not self.chatbot.context_manager.conversation_history or \
           self.chatbot.context_manager.conversation_history[-1]["content"] != message:
            self.chatbot.context_manager.conversation_history.append({
                "role": "user",
                "content": message
            })
        
        # Get response from AI with reasoning support
        model_name = self.chatbot.modelname
        messages = self.chatbot.context_manager.conversation_history
        effort_level = self.config_manager.get_setting("chat_model_config.reasoning_effort", "low")
        
        # Validate effort_level
        valid_effort_levels = ["low", "medium", "high"]
        if effort_level not in valid_effort_levels:
            effort_level = "low"
        
        # Prepare kwargs for API call
        kwargs = {
            "model": model_name,
            "input": messages,
            "stream": True,
        }
        
        # Add reasoning for gpt-5 models
        if model_name.startswith("gpt-5"):
            kwargs["reasoning"] = {"summary": "auto", "effort": effort_level}
        
        try:
            stream = self.chatbot.client.responses.create(**kwargs)
            
            # Collect streaming response with reasoning support
            response_text = ""
            reasoning_text = ""
            display_reasoning = self.config_manager.get_setting("chat_model_config.display_reasoning", False)
            
            # Track state
            reasoning_msg_id = None
            first_response_received = False
            
            for response in stream:
                # Check for cancellation
                if self._cancel_streaming:
                    # Break out of streaming loop
                    break
                    
                if isinstance(response, ResponseAudioDeltaEvent):
                    # This is reasoning content
                    if response.delta:
                        reasoning_text += response.delta
                        if display_reasoning:
                            # Update or create reasoning message
                            reasoning_content = f"**🧠 Thinking... (effort: {effort_level})**\n\n{reasoning_text}"
                            if reasoning_msg_id is None:
                                # Create new reasoning message
                                self.app.call_from_thread(self._add_reasoning_message, reasoning_content)
                                reasoning_msg_id = True
                            else:
                                # Update existing reasoning message
                                self.app.call_from_thread(self._update_reasoning_message, reasoning_content)
                        
                elif isinstance(response, ResponseTextDeltaEvent):
                    # This is the actual response text
                    if response.delta:
                        # Remove loading indicator on first response
                        if not first_response_received:
                            self.app.call_from_thread(self._remove_loading_message)
                            first_response_received = True
                        
                        # Remove reasoning message if it exists
                        if reasoning_msg_id is not None:
                            self.app.call_from_thread(self._remove_reasoning_message)
                            reasoning_msg_id = None
                        
                        response_text += response.delta
                        # Stream the response text as it comes in
                        self.app.call_from_thread(self._update_streaming_response, response_text)
            
            # Handle cancellation cleanup
            if self._cancel_streaming:
                # Remove any loading or reasoning messages
                if not first_response_received:
                    self.app.call_from_thread(self._remove_loading_message)
                if reasoning_msg_id is not None:
                    self.app.call_from_thread(self._remove_reasoning_message)
                # Keep the partial response but finalize it (don't remove it)
                if response_text:
                    self.app.call_from_thread(self._finalize_streaming_response)
                    # Still add partial response to conversation history
                    self.chatbot.context_manager.conversation_history.append({
                        "role": "assistant",
                        "content": response_text
                    })
                # Always show cancellation message (even if no response yet)
                self.app.call_from_thread(self._add_system_message, f"{EMOJI['cross']} *Response cancelled by user.*")
                return
                        
        except Exception as e:
            # Remove loading indicator on error
            self.app.call_from_thread(self._remove_loading_message)
            
            response_text = f"""### {EMOJI['cross']} API Communication Error

```
{str(e)}
```

**Possible causes**:
- Invalid API key
- Network connectivity issues
- Model not available
- Rate limit exceeded
"""
            self.app.call_from_thread(self._add_error_message, response_text)
            return
        
        # Finalize the streamed response
        if response_text:
            self.app.call_from_thread(self._finalize_streaming_response)
            self.chatbot.context_manager.conversation_history.append({
                "role": "assistant",
                "content": response_text
            })
    
    def _add_reasoning_message(self, content: str):
        """Add reasoning message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        reasoning_msg = ChatMessage("system", content)
        chat_panel.mount(reasoning_msg)
        chat_panel.scroll_end(animate=False)
        self._reasoning_msg = reasoning_msg
    
    def _update_reasoning_message(self, content: str):
        """Update reasoning message (called from worker thread)."""
        if hasattr(self, '_reasoning_msg') and self._reasoning_msg and self._reasoning_msg.parent:
            self._reasoning_msg.content = content
            self._reasoning_msg.refresh()
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.scroll_end(animate=False)
    
    def _remove_reasoning_message(self):
        """Remove reasoning message (called from worker thread)."""
        if hasattr(self, '_reasoning_msg') and self._reasoning_msg and self._reasoning_msg.parent:
            self._reasoning_msg.remove()
            self._reasoning_msg = None
    
    def _remove_streaming_message(self):
        """Remove streaming message when cancelled (called from worker thread)."""
        if hasattr(self, '_streaming_msg') and self._streaming_msg and self._streaming_msg.parent:
            self._streaming_msg.remove()
            self._streaming_msg = None
            # Reset the header flag
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel._assistant_header_shown = False
    
    def _update_streaming_response(self, content: str):
        """Update streaming response message (called from worker thread)."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        # Create or update the streaming message
        if not hasattr(self, '_streaming_msg') or self._streaming_msg is None:
            # Check if header should be shown
            show_header = not chat_panel._assistant_header_shown
            chat_panel._assistant_header_shown = True
            self._streaming_msg = ChatMessage("assistant", content, show_header=show_header)
            chat_panel.mount(self._streaming_msg)
        else:
            # Update both content and original_content attributes
            self._streaming_msg.content = content
            self._streaming_msg.original_content = content  # Ensure original_content is updated
            # Find and update the Markdown widget(s) inside the message
            try:
                markdown_widgets = self._streaming_msg.query(".message-content")
                for md_widget in markdown_widgets:
                    if isinstance(md_widget, Markdown):
                        md_widget.update(content)
            except Exception:
                # Fallback: remove and recreate the message
                show_header = not chat_panel._assistant_header_shown
                chat_panel._assistant_header_shown = True
                self._streaming_msg.remove()
                self._streaming_msg = ChatMessage("assistant", content, show_header=show_header)
                chat_panel.mount(self._streaming_msg)
        
        chat_panel.scroll_end(animate=False)
    
    def _finalize_streaming_response(self):
        """Finalize the streaming response (called from worker thread)."""
        # Schedule copy button detection after streaming completes and UI settles
        if hasattr(self, '_streaming_msg') and self._streaming_msg:
            # Store reference to the message for later processing
            final_msg = self._streaming_msg
            final_content = final_msg.original_content if hasattr(final_msg, 'original_content') else final_msg.content
            final_agent_name = final_msg.agent_name if hasattr(final_msg, 'agent_name') else None
            
            # Use the simplified single-phase approach for copy buttons
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            self.app.call_later(chat_panel._add_copy_buttons_to_message, final_msg)
            
            # Autosave the final complete assistant response
            # We're already on the main thread (called via call_from_thread), so use create_task directly
            asyncio.create_task(self._autosave_message("assistant", final_content, final_agent_name))
        
        # Clear the streaming message reference
        if hasattr(self, '_streaming_msg'):
            self._streaming_msg = None
        
        # Reset autosave tracking for next response
        self._last_autosave_content = ""


    
    async def action_new_chat(self):
        """Start a new chat - completely restart everything from scratch."""
        # Cancel any ongoing streaming
        self._cancel_streaming = True
        
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        # Cancel any active agent
        agent_was_active = False
        if hasattr(self, '_agent_input_state') and self._agent_input_state:
            # Check if agent is running (either waiting for input or actively executing)
            if self._agent_input_state.get('waiting_for_input', False) or hasattr(chat_panel, 'current_agent') and chat_panel.current_agent:
                agent_was_active = True
                # Signal the agent to stop by setting the cancelled flag
                self._agent_input_state['cancelled'] = True
                self._agent_input_state['result'] = None
                self._agent_input_state['waiting_for_input'] = False
                self._agent_input_state['ready'].set()  # Unblock the waiting thread if it's waiting
                
                # Give the agent thread a moment to process the cancellation
                await asyncio.sleep(0.2)
                
                # Clean up the agent input state
                if hasattr(self, '_agent_input_state'):
                    delattr(self, '_agent_input_state')
        
        # Clear current agent in chat panel
        if hasattr(chat_panel, 'current_agent') and chat_panel.current_agent:
            agent_was_active = True
            chat_panel.current_agent = None
        
        # Clear messages (this should now catch any lingering agent messages)
        chat_panel.clear_messages()
        
        # Reset session file tracking for new conversation
        self._current_session_file = None
        self._session_start_time = None
        self._session_initialized = False
        self._last_autosave_content = ""
        self._first_user_input = None  # Reset first user input for new session
        
        # Unload any active RAG databases
        if hasattr(self, '_active_rag_kb') and self._active_rag_kb is not None:
            self._active_rag_kb = None
        
        # Clear context chip bar
        try:
            context_chip_bar = self.query_one("#context-chip-bar", ContextChipBar)
            context_chip_bar.clear_all()
        except Exception:
            pass
        
        # Clear any streaming message reference
        if hasattr(self, '_streaming_msg'):
            self._streaming_msg = None
        
        # Completely reinitialize the chatbot (like quitting and relaunching)
        if self.chatbot:
            # Clear conversation history
            self.chatbot.clear_conversation()
            
            # Reinitialize system prompt (like during startup)
            system_prompt = self.config_manager.get_setting(
                "chat_model_config.system_prompt", 
                "You are a helpful assistant. Prioritize markdown format and code blocks when applicable."
            )
            self.chatbot.context_manager.conversation_history.append({
                "role": "system", 
                "content": system_prompt
            })
            
            # Reset conversation ID to start fresh backend session
            self.chatbot._conversation_id = None
        
        # Show welcome message again
        self.show_welcome_message()
        
        # Auto-focus the text input area
        try:
            command_input = self.query_one("#command-input", CommandInput)
            command_input.focus()
        except Exception:
            pass
        
        # Notify user if an agent was cancelled via toast
        if agent_was_active:
            self.app.notify(f"{EMOJI['cross']} Active agent cancelled.", severity="warning", timeout=3)
    
    async def action_agent_picker(self):
        """Open the inline agent picker."""
        chat_panel = self.query_one("#chat-panel", ChatPanel)
        
        if self.chatbot and hasattr(self.chatbot, 'plugin_manager'):
            agents = self.chatbot.plugin_manager.plugin_info
            if agents:
                async def on_agent_selected(selected_agent_name):
                    if selected_agent_name:
                        # Find the agent details
                        for agent_name, agent_desc in agents:
                            if agent_name == selected_agent_name:
                                # Run agent in background thread using threading directly
                                import threading
                                thread = threading.Thread(
                                    target=self._run_agent_workflow_in_thread,
                                    args=(agent_name, agent_desc),
                                    daemon=True
                                )
                                thread.start()
                                break
                # Format agents for display: show name and description
                agent_options = [(f"{name}: {desc}", name) for name, desc in agents]
                # Show the inline picker first (with 'agent' type for tracking)
                await self.show_inline_picker(f"{EMOJI['agent']} Select Agent", agent_options, picker_type="agent")
                # Store callback for when selection is made (after picker is shown)
                self._picker_callback = on_agent_selected
            else:
                chat_panel.add_message("system", "No agents available. Use `/add_agent` to add one.")
        else:
            chat_panel.add_message("error", "Plugin manager not initialized.")
    
    def action_toggle_files(self):
        """Toggle the file explorer panel visibility."""
        file_panel = self.query_one("#file-panel", FileExplorerPanel)
        if self.show_files:
            file_panel.remove_class("visible")
            self.show_files = False
        else:
            file_panel.add_class("visible")
            self.show_files = True
            # Focus the file tree when opening the panel
            file_tree = self.query_one("#file-tree", DirectoryTree)
            file_tree.focus()
    
    def action_toggle_history(self):
        """Toggle the chat history panel visibility."""
        history_panel = self.query_one("#history-panel", ChatHistoryPanel)
        if self.show_history:
            history_panel.remove_class("visible")
            self.show_history = False
        else:
            history_panel.add_class("visible")
            self.show_history = True
            # Refresh the tree to show new entries
            history_panel.refresh_history()
            # Focus the history tree when opening the panel
            history_tree = self.query_one("#history-tree", ChatHistoryTree)
            history_tree.focus()
    
    async def action_change_model(self):
        """Prompt to change the AI model with inline picker."""
        async def handle_model_selection(selected_model):
            if selected_model:
                chat_panel = self.query_one("#chat-panel", ChatPanel)
                chat_panel.add_message("system", f"{EMOJI['checkmark']} Model changed to: **{selected_model}**\n\nUpdating chatbot...")
                
                # Update status bar and chatbot's model
                status_bar = self.query_one("#status-bar", StatusBar)
                status_bar.model_name = selected_model
                
                # Update the chatbot's model directly
                if self.chatbot:
                    self.chatbot.modelname = selected_model
                    chat_panel.add_message("system", f"{EMOJI['checkmark']} Now using model: **{selected_model}**")
                else:
                    chat_panel.add_message("error", f"{EMOJI['cross']} ChatBot not initialized.")
        
        # Show inline picker for model selection
        status_bar = self.query_one("#status-bar", StatusBar)
        current_model = status_bar.model_name
        model_names = ["gpt-4o", "gpt-5", "gpt-4.1", "gpt-oss:20b"]
        
        # Format model names EXACTLY like agent picker does - with newline and description
        model_options = [(f"{name}", name) for name in model_names]
        
        # Show the inline picker first (with 'model' type for tracking)
        await self.show_inline_picker(f"{EMOJI['robot']} Select AI Model", model_options, picker_type="model")
        
        # Store callback for when selection is made (after picker is shown)
        self._picker_callback = handle_model_selection
    
    def on_directory_tree_file_selected(self, event: DirectoryTree.FileSelected) -> None:
        """Handle click on a file in the directory tree.
        
        Single-click behavior:
        - Click on file in file explorer: Adds path to chat and closes file explorer
        - Click on file in history panel: Loads that chat history
        """
        # Only process if it's actually a file (not a directory)
        if not event.path.is_file():
            return
        
        # Check if this is from the history tree by checking the control/node that sent the event
        if isinstance(event.control, ChatHistoryTree) or isinstance(event.node.tree, ChatHistoryTree):
            # Load chat history
            self._load_chat_history(event.path)
            return
        
        # Otherwise, it's from the file explorer - add file to chat with # prefix
        file_path = event.path
        try:
            # Convert to relative path from current directory
            import os
            relative_path = os.path.relpath(file_path, os.getcwd())
            
            # Get the command input widget
            command_input = self.query_one("#command-input", CommandInput)
            
            # Insert the path with # prefix at the cursor position or append if empty
            current_text = command_input.text
            if current_text:
                # Add space before the path only if current text does not end with whitespace
                if not current_text[-1].isspace():
                    command_input.insert(f" #{relative_path} ")
                else:
                    command_input.insert(f"#{relative_path} ")
            else:
                command_input.insert(f"#{relative_path} ")
            
            # Close the file explorer
            self.action_toggle_files()
            
            # Focus the command input
            command_input.focus()
            
        except Exception as e:
            # If there's an error, just use the absolute path with # prefix
            command_input = self.query_one("#command-input", CommandInput)
            command_input.insert(f"#{str(file_path)} ")
            self.action_toggle_files()
            command_input.focus()
    
    def _load_chat_history(self, history_file: Path) -> None:
        """Load a chat history file and restore the conversation.
        
        This method handles loading history whether it's:
        1. The first history being loaded (from fresh state)
        2. Switching from one loaded history to another
        
        In both cases, it completely resets the session state and loads the new history.
        """
        try:
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            
            # STEP 1: Reset session state completely (important when switching between histories)
            # This ensures we don't append to the wrong file or carry over old state
            self._current_session_file = None
            self._session_initialized = False
            self._session_start_time = None
            self._first_user_input = None
            self._last_autosave_content = ""
            
            # STEP 2: Read and parse the log file
            with open(history_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse the log file format
            # Format: [timestamp] ROLE:\ncontent\n----\n\n
            import re
            
            # First, remove the header section (everything before the first message)
            # Header format: ValBot TUI Chat Log\nSession Started: ...\n====\n\n
            header_end = content.find('=' * 80)
            if header_end != -1:
                content = content[header_end + 80:]  # Skip past the header separator
            
            # Split by the separator line (80 dashes)
            sections = content.split('-' * 80)
            
            messages = []
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Match pattern: [timestamp] ROLE:\ncontent
                # Use a more flexible pattern that allows for optional newlines
                match = re.match(r'\[([^\]]+)\]\s+([^:]+):\s*(.*)', section, re.DOTALL)
                if match:
                    timestamp = match.group(1)
                    role_label = match.group(2).strip()
                    message_content = match.group(3).strip()
                    
                    # Parse role and agent name if present
                    agent_name = None
                    if '(Agent:' in role_label:
                        # Extract agent name from "ASSISTANT (Agent: name)"
                        role_match = re.match(r'(\w+)\s*\(Agent:\s*([^)]+)\)', role_label)
                        if role_match:
                            role = role_match.group(1).lower()
                            agent_name = role_match.group(2).strip()
                        else:
                            role = role_label.split('(')[0].strip().lower()
                    else:
                        role = role_label.lower()
                    
                    messages.append({
                        'role': role,
                        'content': message_content,
                        'agent_name': agent_name
                    })
            
            # STEP 3: Clear the chat panel and UI state
            chat_panel.clear_messages()
            
            # Cancel any ongoing streaming
            self._cancel_streaming = True
            
            # Clear any streaming message reference
            if hasattr(self, '_streaming_msg'):
                self._streaming_msg = None
            
            # Unload any active RAG databases
            if hasattr(self, '_active_rag_kb') and self._active_rag_kb is not None:
                self._active_rag_kb = None
            
            # Clear context chip bar
            try:
                context_chip_bar = self.query_one("#context-chip-bar", ContextChipBar)
                context_chip_bar.clear_all()
            except Exception:
                pass
            
            # STEP 4: Completely reinitialize the chatbot conversation history
            if self.chatbot:
                # Clear conversation history
                self.chatbot.clear_conversation()
                
                # Reinitialize system prompt (like during startup)
                system_prompt = self.config_manager.get_setting(
                    "chat_model_config.system_prompt", 
                    "You are a helpful assistant. Prioritize markdown format and code blocks when applicable."
                )
                self.chatbot.context_manager.conversation_history.append({
                    "role": "system", 
                    "content": system_prompt
                })
                
                # Add all loaded messages to conversation history
                for msg in messages:
                    self.chatbot.context_manager.conversation_history.append({
                        'role': msg['role'],
                        'content': msg['content']
                    })
                
                # Reset conversation ID to start fresh backend session
                self.chatbot._conversation_id = None
            
            # STEP 5: Add all messages to chat panel for display
            # IMPORTANT: Skip autosave to prevent creating duplicate history entries
            for msg in messages:
                chat_panel.add_message(
                    role=msg['role'],
                    content=msg['content'],
                    agent_name=msg['agent_name'],
                    skip_autosave=True  # Don't save historical messages again
                )
            
            # STEP 6: CRITICAL - Set the current session file to the loaded history file
            # This ensures new messages are appended to THIS specific history file
            # Not to the previous one if switching histories, and not to a new file
            self._current_session_file = history_file
            self._session_initialized = True
            self._session_start_time = datetime.now()
            self._first_user_input = None  # Already has messages, don't need to track first input
            self._last_autosave_content = ""  # Reset tracking
            
            # Close the history panel
            self.action_toggle_history()
            
            # Show success message
            self.app.notify(f"{EMOJI['checkmark']} Loaded chat history: {history_file.stem}", severity="information", timeout=3)
            
            # Auto-focus the text input area
            try:
                command_input = self.query_one("#command-input", CommandInput)
                command_input.focus()
            except Exception:
                pass
            
        except Exception as e:
            # Show error message
            chat_panel = self.query_one("#chat-panel", ChatPanel)
            chat_panel.add_message("error", f"❌ Failed to load chat history: {str(e)}\n\n```\n{traceback.format_exc()}\n```")
            self.action_toggle_history()


class ValbotTUI(App):
    """ValBot Terminal User Interface with Material Design."""
    
    TITLE = f"{EMOJI['robot']} ValBot TUI - AI Assistant"
    
    # Material Design inspired CSS
    CSS = """
    App {
        background: $background;
    }
    
    Screen {
        background: $background;
    }
    
    /* Enhanced Header */
    Header {
        dock: top;
        height: 1;
        text-style: bold;
        text-align: center;
        align: center middle;
    }
    
    Header .header--title {
        text-style: bold;
        text-align: center;
    }
    
    Header .header--clock {
        color: $primary-lighten-2;
    }
    
    /* Enhanced Footer */
    Footer {
        dock: bottom;
        height: 1;
        background: $surface;
    }
    
    Footer > .footer--highlight {
        background: $accent;
    }
    
    Footer > .footer--key {
        background: $primary-darken-1;
    }
    
    Footer > .footer--description {
        color: $text-muted;
    }
    
    /* Smooth animations */
    * {
        transition: background 200ms, color 200ms;
    }
    
    /* Remove all focus borders */
    *:focus {
        border: none !important;
    }
    
    /* Loading spinner */
    LoadingIndicator {
        background: $surface;
        color: $accent;
    }
    """
    
    def __init__(self, config_path: Optional[str] = None):
        super().__init__()
        self.config_manager = ConfigManager(config_path)
        utilities.set_config_manager(self.config_manager)
        self.should_restart = False  # Flag to trigger restart
        self.should_update = False  # Flag to trigger updater
        # Register and use the custom valbot-dark theme
        self.register_theme(VALBOT_DARK_THEME)
        
        # Load saved theme from config file, or use default
        saved_theme = load_theme_config()
        if saved_theme and saved_theme in self.available_themes:
            self.theme = saved_theme
        else:
            self.theme = "valbot-dark (default)"
    
    def watch_theme(self, new_theme: str) -> None:
        """Watch for theme changes and save to config file."""
        save_theme_config(new_theme)
    
    def on_mount(self) -> None:
        """Set up the application after mounting."""
        # Check if history configuration exists
        history_config = load_history_config()
        
        if history_config is None:
            # Show configuration screen with callback to continue
            def on_config_complete(result):
                # Result doesn't matter - config was saved either way
                # Now show the main screen
                self.push_screen(MainScreen(self.config_manager))
            
            self.push_screen(HistoryConfigScreen(), callback=on_config_complete)
        else:
            # Config exists, go directly to main screen
            self.push_screen(MainScreen(self.config_manager))


def main():
    """Main entry point for the TUI application."""
    import argparse
    import locale
    
    # Force UTF-8 encoding for UNIX systems to properly display emojis
    try:
        # Set locale to UTF-8 if available
        locale.setlocale(locale.LC_ALL, '')
        
        # For Python 3.7+, ensure UTF-8 mode is enabled
        if sys.version_info >= (3, 7):
            # Set environment variables before any I/O operations
            import os
            os.environ['PYTHONIOENCODING'] = 'utf-8'
            
            # Reconfigure stdout/stderr for UTF-8
            if hasattr(sys.stdout, 'reconfigure'):
                sys.stdout.reconfigure(encoding='utf-8', errors='replace')
            if hasattr(sys.stderr, 'reconfigure'):
                sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except (locale.Error, Exception) as e:
        # Fallback if locale setting fails
        import os
        os.environ['LANG'] = 'en_US.UTF-8'
        os.environ['LC_ALL'] = 'en_US.UTF-8'
        pass
    
    parser = argparse.ArgumentParser(description="ValBot TUI - Terminal User Interface")
    parser.add_argument("--config", type=str, help="Path to custom configuration file")
    args = parser.parse_args()
    
    while True:
        try:
            app = ValbotTUI(config_path=args.config)
            app.run()
            
            # Check if we should restart
            if app.should_restart:
                print(f"\n{EMOJI['refresh']} Restarting TUI...\n")
                continue
            elif app.should_update:
                print(f"\n{EMOJI['refresh']} Running updater in CLI mode...\n")
                # Run the CLI app.py with update command
                import subprocess
                subprocess.call([sys.executable, "app.py"])
                break
            else:
                break
        except KeyboardInterrupt:
            print("\nExiting...")
            sys.exit(0)


if __name__ == "__main__":
    main()
