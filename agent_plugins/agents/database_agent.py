from agent_plugins.common_agent_imports import *
from agent_plugins.agent_plugin import AgentPlugin
from paibox import Agent
import asyncio
import json
import traceback
from pathlib import Path
from typing import List, Dict, Union, Optional, Any, Tuple
import hashlib
import re

# Optional dependencies - will check at runtime
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    chromadb = None

try:
    from docx import Document
except ImportError:
    Document = None

from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn
from rich.prompt import Confirm
from rich import box
from rich.table import Table
from rich.panel import Panel
from rich.markdown import Markdown

from agent_plugins.common_tools.terminal_tools import list_files, find_files, grep_in_files, show_file_tree, run_shell_command
from agent_plugins.common_tools.file_tools import read_file, FileContent


class MyCustomPlugin(AgentPlugin):
    """
    RAG Knowledge Base Agent
    
    This agent allows you to create and interact with RAG (Retrieval-Augmented Generation)
    knowledge bases. You can:
    
    1. CREATE KNOWLEDGE BASE:
       - Provide PDF, DOCX, or text files
       - Specify output directory for vector database cache
       - Agent processes documents and creates reusable embeddings
    
    2. LOAD KNOWLEDGE BASE:
       - Point to existing knowledge base directory
       - Query the knowledge base interactively
       - Get AI-powered answers based on document content
    
    The agent uses vector embeddings and semantic search to retrieve relevant
    information from your documents.
    """
    
    REQUIRED_ARGS = {}
    
    def get_initializer_args(self):
        """Prompt user to select mode: load or create knowledge base."""
        from rich.console import Console
        
        console = Console()
        
        # Menu choices
        choices = [
            ("1", "Load existing knowledge base"),
            ("2", "Create new knowledge base")
        ]
        
        # Display header
        console.print("\n[bold cyan]RAG Knowledge Base Agent[/bold cyan]\n")
        console.print("[dim]Select an option:[/dim]\n")
        
        table = Table(show_header=False, box=box.ROUNDED, padding=(0, 2))
        table.add_column("Key", style="cyan", width=5)
        table.add_column("Mode", style="white")
        for key, description in choices:
            table.add_row(key, description)
        
        console.print(table)
        console.print()
        
        mode_value = None
        while mode_value not in ['1', '2']:
            mode_value = console.input("[bold yellow]Enter your choice (1 or 2):[/bold yellow] ").strip()
            if mode_value not in ['1', '2']:
                console.print("[red]Invalid choice. Please enter 1 or 2.[/red]")
        
        console.print(f"\n[green]✓ Selected:[/green] {choices[int(mode_value) - 1][1]}\n")
        
        if mode_value == '1':
            # Load existing knowledge base
            kb_directory = console.input("[bold yellow]Enter the path to knowledge base directory:[/bold yellow] ").strip()
            return {
                'mode': mode_value,
                'kb_directory': kb_directory
            }
        else:
            # Create new knowledge base
            console.print("[dim]Enter file paths separated by commas (supports PDF, DOCX, TXT)[/dim]")
            file_paths_input = console.input("[bold yellow]File paths:[/bold yellow] ").strip()
            output_folder = console.input("[bold yellow]Output folder for knowledge base:[/bold yellow] ").strip()
            
            return {
                'mode': mode_value,
                'file_paths': file_paths_input,
                'output_folder': output_folder
            }

    def __init__(self, model, **kwargs):
        super().__init__(model, **kwargs)
        # Initialize Rich console for beautiful formatting
        self.console = Console()
        
        # RAG knowledge base will be initialized based on mode
        self.rag_kb = None
        
        # Create AI agent for answering questions
        self.qa_agent = self.create_qa_agent(model)
        
        # Add tools for autonomous operation
        self.add_agentic_tools()

    def create_qa_agent(self, model):
        """Create an agent specialized in answering questions based on retrieved context."""
        system_prompt = """
        You are an expert AI assistant that answers questions based on provided document context.
        
        You will receive:
        1. A user's question
        2. Relevant excerpts from documents retrieved via semantic search
        
        Your task:
        - Answer the question using ONLY information from the provided context
        - Be clear, accurate, and concise
        - Cite which documents or sections you're referencing when possible
        - If the context doesn't contain enough information to answer, say so
        - Use proper Markdown formatting in your responses
        
        CRITICAL: Format ALL responses using proper Markdown:
        - Use ### for section headers
        - Use ```language for code blocks
        - Use `code` for inline code or technical terms
        - Use **bold** for emphasis
        - Use bullet points (- or *) for lists
        - Use > for important notes or warnings
        - Use tables when presenting structured information
        
        Always ground your answers in the provided context. Do not make up information.
        """
        
        return Agent(
            model=model,
            system_prompt=system_prompt,
            result_type=str,
            retries=2
        )
    
    def add_agentic_tools(self):
        """Add autonomous tools for the agent."""
        self.qa_agent.tool(list_files)
        self.qa_agent.tool(find_files)
        self.qa_agent.tool(grep_in_files)
        self.qa_agent.tool(show_file_tree)
        self.qa_agent.tool(run_shell_command)
        self.qa_agent.tool(read_file)

    def create_knowledge_base(self, file_paths: List[str], output_folder: str):
        """
        Create a new knowledge base from provided files.
        
        Args:
            file_paths: List of file paths to process
            output_folder: Directory to store the knowledge base
        """
        self.console.rule("[bold blue]Creating Knowledge Base[/bold blue]", style="blue")
        
        # Create output folder if it doesn't exist
        output_path = Path(output_folder)
        output_path.mkdir(parents=True, exist_ok=True)
        self.console.print(f"[green]✓ Output folder created/verified:[/green] {output_folder}\n")
        
        # Initialize RAG Knowledge Base
        try:
            self.rag_kb = RAGKnowledgeBase(
                pdf_dir=output_path,
                cache_dir=output_path / ".rag_cache",
                chunk_size=1000,
                chunk_overlap=200,
                collection_name="rag_knowledge_base"
            )
            self.console.print(f"[green]✓ RAG system initialized[/green]\n")
        except Exception as e:
            self.console.print(f"[red]Error initializing RAG system: {e}[/red]")
            return
        
        # Process each file
        self.console.print(f"[bold cyan]Processing {len(file_paths)} file(s)...[/bold cyan]\n")
        
        processed_count = 0
        failed_count = 0
        
        for file_path_str in file_paths:
            file_path = Path(file_path_str.strip())
            
            if not file_path.exists():
                self.console.print(f"[red]✗ File not found:[/red] {file_path}")
                failed_count += 1
                continue
            
            self.console.print(f"[cyan]Processing:[/cyan] {file_path.name}")
            
            try:
                suffix = file_path.suffix.lower()
                
                if suffix == '.pdf':
                    # Copy PDF to output directory if not already there
                    if file_path.parent != output_path:
                        import shutil
                        shutil.copy(file_path, output_path / file_path.name)
                    self.rag_kb.load_pdfs([file_path.name], force_reload=True)
                    
                elif suffix == '.docx':
                    # Copy DOCX to output directory if not already there
                    if file_path.parent != output_path:
                        import shutil
                        shutil.copy(file_path, output_path / file_path.name)
                    self.rag_kb.load_docx_files([file_path.name], force_reload=True)
                    
                elif suffix in ['.txt', '.pdl', '.md', '.py', '.c', '.cpp', '.h']:
                    # Copy text file to output directory if not already there
                    if file_path.parent != output_path:
                        import shutil
                        shutil.copy(file_path, output_path / file_path.name)
                    self.rag_kb.load_text_file(output_path / file_path.name, force_reload=True)
                    
                else:
                    self.console.print(f"[yellow]⚠ Unsupported file type:[/yellow] {suffix}")
                    failed_count += 1
                    continue
                
                processed_count += 1
                self.console.print(f"[green]✓ Processed successfully[/green]\n")
                
            except Exception as e:
                self.console.print(f"[red]✗ Error processing file: {e}[/red]\n")
                failed_count += 1
        
        # Display summary
        stats = self.rag_kb.get_stats()
        self.console.rule("[bold blue]Knowledge Base Created[/bold blue]", style="blue")
        self.console.print(f"""
[bold green]Summary:[/bold green]
  • Files processed: {processed_count}
  • Files failed: {failed_count}
  • Total chunks: {stats['total_chunks']}
  • Knowledge base location: {output_folder}
        """)

    def load_knowledge_base(self, kb_directory: str):
        """
        Load an existing knowledge base.
        
        Args:
            kb_directory: Path to knowledge base directory
        """
        self.console.rule("[bold blue]Loading Knowledge Base[/bold blue]", style="blue")
        
        kb_path = Path(kb_directory)
        
        if not kb_path.exists():
            self.console.print(f"[red]Error: Directory does not exist: {kb_directory}[/red]")
            return False
        
        cache_dir = kb_path / ".rag_cache"
        if not cache_dir.exists():
            self.console.print(f"[red]Error: No knowledge base cache found in {kb_directory}[/red]")
            self.console.print("[yellow]This directory may not contain a valid knowledge base.[/yellow]")
            return False
        
        try:
            self.rag_kb = RAGKnowledgeBase(
                pdf_dir=kb_path,
                cache_dir=cache_dir,
                chunk_size=1000,
                chunk_overlap=200,
                collection_name="rag_knowledge_base"
            )
            
            stats = self.rag_kb.get_stats()
            self.console.print(f"""
[green]✓ Knowledge base loaded successfully[/green]

[bold cyan]Statistics:[/bold cyan]
  • Total chunks: {stats['total_chunks']}
  • Embedding model: {stats['embedding_model']}
  • Chunk size: {stats['chunk_size']}
  • Location: {kb_directory}
            """)
            return True
            
        except Exception as e:
            self.console.print(f"[red]Error loading knowledge base: {e}[/red]")
            traceback.print_exc()
            return False

    def interactive_qa_mode(self):
        """Interactive Q&A mode with the loaded knowledge base."""
        self.console.rule("[bold blue]Interactive Q&A Mode[/bold blue]", style="blue")
        self.console.print("\n[bold cyan]Ask questions about your documents![/bold cyan]")
        self.console.print("[dim]Type 'exit', 'quit', or 'done' to end the session[/dim]\n")
        
        if not self.rag_kb:
            self.console.print("[red]No knowledge base loaded. Please create or load one first.[/red]")
            return
        
        question_count = 0
        total_tokens = 0
        session_history = []
        
        while True:
            try:
                # Get user question
                self.console.print()
                question = self.console.input("[bold yellow]Your question:[/bold yellow] ").strip()
                
                if question.lower() in ['exit', 'quit', 'done', '']:
                    break
                
                question_count += 1
                
                # Retrieve relevant context
                self.console.print(f"\n[dim]Searching knowledge base...[/dim]")
                relevant_context = self.rag_kb.get_context_for_query(
                    query=question,
                    top_k=8,
                    include_metadata=True
                )
                
                if not relevant_context or relevant_context == "No relevant information found.":
                    self.console.print("[yellow]No relevant information found in the knowledge base.[/yellow]")
                    continue
                
                # Generate answer using AI agent
                self.console.print(f"[dim]Generating answer...[/dim]\n")
                
                qa_prompt = f"""
                User Question: {question}
                
                Relevant Context from Documents:
                {relevant_context}
                
                Please answer the user's question based on the context above.
                Use proper Markdown formatting in your response.
                Cite sources when possible.
                """
                
                result = asyncio.run(self.qa_agent.run(qa_prompt))
                answer = result.data
                tokens = result.usage().total_tokens
                total_tokens += tokens
                
                # Display answer with Markdown rendering
                self.console.print(Panel(
                    Markdown(answer),
                    title="[bold green]Answer[/bold green]",
                    border_style="green",
                    padding=(1, 2)
                ))
                
                # Save to session history
                session_history.append({
                    'question': question,
                    'answer': answer,
                    'tokens': tokens
                })
                
            except KeyboardInterrupt:
                self.console.print("\n[yellow]Session interrupted by user[/yellow]")
                break
            except Exception as e:
                self.console.print(f"\n[red]Error: {e}[/red]")
                traceback.print_exc()
        
        # Display session summary
        if session_history:
            self.console.rule("[bold blue]Session Summary[/bold blue]", style="blue")
            self.console.print(f"""
[bold green]Summary:[/bold green]
  • Questions asked: {question_count}
  • Total tokens used: {total_tokens}
            """)
            
            # Save session history
            try:
                history_file = Path.cwd() / "rag_qa_session_history.json"
                with open(history_file, 'w', encoding='utf-8') as f:
                    json.dump(session_history, f, indent=2, ensure_ascii=False)
                self.console.print(f"\n[green]Session history saved to:[/green] {history_file}")
            except Exception as e:
                self.console.print(f"[yellow]Could not save session history: {e}[/yellow]")

    def run_agent_flow(self, context, **kwargs):
        """Main entry point - routes to appropriate mode based on user selection."""
        # Get mode from initializer_args
        mode = self.initializer_args.get('mode')
        
        if not mode:
            # Call get_initializer_args to get user input
            init_args = self.get_initializer_args()
            self.initializer_args.update(init_args)
            mode = self.initializer_args.get('mode', '1')
        
        if mode == '1':
            # Load existing knowledge base
            kb_directory = self.initializer_args.get('kb_directory')
            if not kb_directory:
                self.console.print("[red]Error: kb_directory is required for load mode[/red]")
                return
            
            success = self.load_knowledge_base(kb_directory)
            if success:
                # Enter interactive Q&A mode
                self.interactive_qa_mode()
        
        elif mode == '2':
            # Create new knowledge base
            file_paths_input = self.initializer_args.get('file_paths')
            output_folder = self.initializer_args.get('output_folder')
            
            if not file_paths_input or not output_folder:
                self.console.print("[red]Error: file_paths and output_folder are required for create mode[/red]")
                return
            
            # Parse file paths
            file_paths = [path.strip() for path in file_paths_input.split(',') if path.strip()]
            
            if not file_paths:
                self.console.print("[red]Error: No valid file paths provided[/red]")
                return
            
            # Create knowledge base
            self.create_knowledge_base(file_paths, output_folder)
            
            # Ask if user wants to query it now
            self.console.print()
            query_now = Confirm.ask("[bold cyan]Would you like to query the knowledge base now?[/bold cyan]")
            
            if query_now and self.rag_kb:
                self.interactive_qa_mode()


class RAGKnowledgeBase:
    """
    RAG-based knowledge base that uses vector embeddings and semantic search
    to retrieve relevant information from documents.
    """
    
    def __init__(
        self,
        pdf_dir: Path,
        cache_dir: Optional[Path] = None,
        embedding_model: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        collection_name: str = "rag_knowledge_base"
    ):
        """
        Initialize RAG Knowledge Base.
        
        Args:
            pdf_dir: Directory containing documents
            cache_dir: Directory for caching embeddings
            embedding_model: HuggingFace model name for embeddings
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks for context preservation
            collection_name: Name of the ChromaDB collection
        """
        self.console = Console()
        self.pdf_dir = Path(pdf_dir)
        self.cache_dir = cache_dir or (self.pdf_dir / ".rag_cache")
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        self.embedding_model_name = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.collection_name = collection_name
        
        # Initialize components
        self.embedding_model = None
        self.chroma_client = None
        self.collection = None
        
        # Check dependencies
        self._check_dependencies()
        
        # Initialize models and database
        self._initialize_components()
    
    def _check_dependencies(self):
        """Check if required dependencies are installed."""
        missing = []
        
        if PyPDF2 is None:
            missing.append("PyPDF2")
        if SentenceTransformer is None:
            missing.append("sentence-transformers")
        if chromadb is None:
            missing.append("chromadb")
        
        if missing:
            error_msg = f"Missing required packages: {', '.join(missing)}"
            self.console.print(f"[bold red]Error: {error_msg}[/bold red]")
            self.console.print("\n[yellow]Install with:[/yellow]")
            self.console.print(f"  pip install {' '.join(missing)}")
            raise ImportError(error_msg)
    
    def _initialize_components(self):
        """Initialize embedding model and vector database."""
        try:
            # Initialize ChromaDB first to check if we have existing data
            chroma_path = str(self.cache_dir / "chroma_db")
            self.chroma_client = chromadb.PersistentClient(path=chroma_path)
            
            # Get or create collection
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            # Check if we have existing embeddings
            existing_count = self.collection.count()
            is_new = existing_count == 0
            
            if is_new:
                self.console.print("[dim]Initializing new knowledge base...[/dim]")
            else:
                self.console.print(f"[dim]Found existing knowledge base with {existing_count} chunks[/dim]")
            
            # Load embedding model
            if is_new:
                self.console.print("[dim]Loading embedding model (this may take a moment)...[/dim]")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            if is_new:
                self.console.print("[green]✓ Embedding model loaded[/green]")
            
        except Exception as e:
            self.console.print(f"[red]Error initializing components: {e}[/red]")
            raise
    
    def _get_pdf_hash(self, file_path: Path) -> str:
        """Generate hash of file for caching."""
        hash_obj = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_obj.update(chunk)
        return hash_obj.hexdigest()
    
    def extract_text_from_pdf(self, pdf_path: Path) -> str:
        """Extract all text from a PDF file."""
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {pdf_path.name}: {e}[/red]")
            return ""
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """Extract all text from a .docx file."""
        if Document is None:
            self.console.print(f"[yellow]⚠ python-docx not installed. Install with: pip install python-docx[/yellow]")
            return ""
        
        try:
            doc = Document(str(docx_path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            return text
        except Exception as e:
            self.console.print(f"[red]Error extracting text from {docx_path.name}: {e}[/red]")
            return ""
    
    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for better context preservation."""
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            # Extract chunk
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk_text.rfind('.')
                last_newline = chunk_text.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > self.chunk_size * 0.5:
                    end = start + break_point + 1
                    chunk_text = text[start:end]
            
            # Create chunk with metadata
            chunk = {
                'text': chunk_text.strip(),
                'metadata': {
                    **metadata,
                    'chunk_id': chunk_id,
                    'start_char': start,
                    'end_char': end
                }
            }
            
            if chunk['text']:
                chunks.append(chunk)
                chunk_id += 1
            
            # Move to next chunk with overlap
            start = end - self.chunk_overlap
        
        return chunks
    
    def load_text_file(self, file_path: Path, source_name: str = None, force_reload: bool = False):
        """Load a text file into the vector database."""
        if not file_path.exists():
            self.console.print(f"[yellow]⚠[/yellow] File not found: {file_path}")
            return
        
        source_name = source_name or file_path.name
        
        # Check if already processed
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
            
            # Check if file is already loaded
            file_hash = self._get_pdf_hash(file_path)
            if source_name in processed_docs and processed_docs[source_name] == file_hash:
                self.console.print(f"[dim]Skipping {source_name} (already processed)[/dim]")
                return
        
        # Read text file
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except Exception as e:
            self.console.print(f"[red]Error reading {file_path.name}: {e}[/red]")
            return
        
        if not text:
            return
        
        # Create chunks
        file_hash = self._get_pdf_hash(file_path)
        metadata = {
            'source': source_name,
            'doc_hash': file_hash,
            'file_type': 'text'
        }
        chunks = self.chunk_text(text, metadata)
        
        if chunks:
            # Generate embeddings and store
            texts = [chunk['text'] for chunk in chunks]
            embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
            
            # Prepare for ChromaDB
            ids = [f"{source_name}_{chunk['metadata']['chunk_id']}" for chunk in chunks]
            metadatas = [chunk['metadata'] for chunk in chunks]
            
            # Add to collection
            self.collection.add(
                ids=ids,
                embeddings=embeddings.tolist(),
                documents=texts,
                metadatas=metadatas
            )
            
            # Update processed docs cache
            processed_docs[source_name] = file_hash
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
    
    def load_pdfs(self, pdf_files: Optional[List[str]] = None, force_reload: bool = False):
        """Load and process PDF files into the vector database."""
        # Check if knowledge base is already fully loaded
        if not force_reload:
            existing_count = self.collection.count()
            cache_file = self.cache_dir / "processed_docs.json"
            
            if existing_count > 0 and cache_file.exists():
                self.console.print(f"[dim]Knowledge base already loaded ({existing_count} chunks)[/dim]")
        
        # Find PDF files
        if pdf_files is None:
            pdf_paths = list(self.pdf_dir.glob("*.pdf"))
        else:
            pdf_paths = [self.pdf_dir / pdf for pdf in pdf_files]
        
        if not pdf_paths:
            self.console.print("[yellow]No PDF files found[/yellow]")
            return
        
        self.console.print(f"\n[bold cyan]Loading {len(pdf_paths)} PDF(s) into RAG Knowledge Base[/bold cyan]")
        
        # Track processed documents
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
        
        all_chunks = []
        total_chunks = 0
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=self.console
        ) as progress:
            
            task = progress.add_task("[cyan]Processing PDFs...", total=len(pdf_paths))
            
            for pdf_path in pdf_paths:
                if not pdf_path.exists():
                    progress.update(task, advance=1)
                    continue
                
                # Check if already processed
                pdf_hash = self._get_pdf_hash(pdf_path)
                if not force_reload and pdf_path.name in processed_docs:
                    if processed_docs[pdf_path.name] == pdf_hash:
                        progress.update(task, advance=1, description=f"[dim]Skipping {pdf_path.name} (cached)[/dim]")
                        continue
                
                progress.update(task, description=f"[cyan]Processing {pdf_path.name}...[/cyan]")
                
                # Extract text
                text = self.extract_text_from_pdf(pdf_path)
                
                if text:
                    # Create chunks
                    metadata = {
                        'source': pdf_path.name,
                        'doc_hash': pdf_hash,
                        'file_type': 'pdf'
                    }
                    chunks = self.chunk_text(text, metadata)
                    all_chunks.extend(chunks)
                    total_chunks += len(chunks)
                    processed_docs[pdf_path.name] = pdf_hash
                
                progress.update(task, advance=1)
        
        # Generate embeddings and store in ChromaDB
        if all_chunks:
            self.console.print(f"\n[yellow]Generating embeddings for {total_chunks} chunks...[/yellow]")
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=self.console
            ) as progress:
                
                task = progress.add_task("[cyan]Creating embeddings...", total=total_chunks)
                
                # Process in batches
                batch_size = 32
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i + batch_size]
                    texts = [chunk['text'] for chunk in batch]
                    embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                    
                    # Prepare for ChromaDB
                    ids = [f"{chunk['metadata']['source']}_{chunk['metadata']['chunk_id']}" for chunk in batch]
                    metadatas = [chunk['metadata'] for chunk in batch]
                    
                    # Add to collection
                    self.collection.add(
                        ids=ids,
                        embeddings=embeddings.tolist(),
                        documents=texts,
                        metadatas=metadatas
                    )
                    
                    progress.update(task, advance=len(batch))
            
            # Save processed documents cache
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
            
            self.console.print(f"[green]✓ Successfully loaded {total_chunks} chunks into vector database[/green]")
        else:
            self.console.print("[yellow]No chunks to process[/yellow]")
    
    def load_docx_files(self, docx_files: Optional[List[str]] = None, force_reload: bool = False):
        """Load and process .docx files into the vector database."""
        if Document is None:
            self.console.print(f"[yellow]⚠ python-docx not installed. Skipping .docx files. Install with: pip install python-docx[/yellow]")
            return
        
        # Find .docx files
        if docx_files is None:
            docx_paths = list(self.pdf_dir.glob("*.docx"))
        else:
            docx_paths = [self.pdf_dir / docx for docx in docx_files]
        
        if not docx_paths:
            self.console.print("[yellow]No .docx files found[/yellow]")
            return
        
        self.console.print(f"\n[bold cyan]Loading {len(docx_paths)} .docx file(s) into RAG Knowledge Base[/bold cyan]")
        
        # Track processed documents
        cache_file = self.cache_dir / "processed_docs.json"
        processed_docs = {}
        
        if cache_file.exists() and not force_reload:
            with open(cache_file, 'r') as f:
                processed_docs = json.load(f)
        
        all_chunks = []
        total_chunks = 0
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            console=self.console
        ) as progress:
            
            task = progress.add_task("[cyan]Processing .docx files...", total=len(docx_paths))
            
            for docx_path in docx_paths:
                if not docx_path.exists():
                    progress.update(task, advance=1)
                    continue
                
                # Check if already processed
                docx_hash = self._get_pdf_hash(docx_path)
                if not force_reload and docx_path.name in processed_docs:
                    if processed_docs[docx_path.name] == docx_hash:
                        progress.update(task, advance=1, description=f"[dim]Skipping {docx_path.name} (cached)[/dim]")
                        continue
                
                progress.update(task, description=f"[cyan]Processing {docx_path.name}...[/cyan]")
                
                # Extract text
                text = self.extract_text_from_docx(docx_path)
                
                if text:
                    # Create chunks
                    metadata = {
                        'source': docx_path.name,
                        'doc_hash': docx_hash,
                        'file_type': 'docx'
                    }
                    chunks = self.chunk_text(text, metadata)
                    all_chunks.extend(chunks)
                    total_chunks += len(chunks)
                    processed_docs[docx_path.name] = docx_hash
                
                progress.update(task, advance=1)
        
        # Generate embeddings and store in ChromaDB
        if all_chunks:
            self.console.print(f"\n[yellow]Generating embeddings for {total_chunks} chunks...[/yellow]")
            
            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                BarColumn(),
                TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
                console=self.console
            ) as progress:
                
                task = progress.add_task("[cyan]Creating embeddings...", total=total_chunks)
                
                # Process in batches
                batch_size = 32
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i + batch_size]
                    texts = [chunk['text'] for chunk in batch]
                    embeddings = self.embedding_model.encode(texts, show_progress_bar=False)
                    
                    # Prepare for ChromaDB
                    ids = [f"{chunk['metadata']['source']}_{chunk['metadata']['chunk_id']}" for chunk in batch]
                    metadatas = [chunk['metadata'] for chunk in batch]
                    
                    # Add to collection
                    self.collection.add(
                        ids=ids,
                        embeddings=embeddings.tolist(),
                        documents=texts,
                        metadatas=metadatas
                    )
                    
                    progress.update(task, advance=len(batch))
            
            # Save processed documents cache
            with open(cache_file, 'w') as f:
                json.dump(processed_docs, f, indent=2)
            
            self.console.print(f"[green]✓ Successfully loaded {total_chunks} chunks into vector database[/green]")
        else:
            self.console.print("[yellow]No chunks to process[/yellow]")
    
    def retrieve_relevant_context(
        self, 
        query: str, 
        top_k: int = 5,
        min_similarity: float = 0.0
    ) -> List[Dict[str, Any]]:
        """Retrieve the most relevant text chunks for a given query."""
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding.tolist()],
                n_results=top_k
            )
            
            # Format results
            chunks = []
            if results['documents'] and results['documents'][0]:
                for i, doc in enumerate(results['documents'][0]):
                    metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                    distance = results['distances'][0][i] if results['distances'] else 0
                    similarity = 1 - distance  # Convert distance to similarity
                    
                    if similarity >= min_similarity:
                        chunks.append({
                            'text': doc,
                            'metadata': metadata,
                            'similarity': similarity
                        })
            
            return chunks
        
        except Exception as e:
            self.console.print(f"[red]Error retrieving context: {e}[/red]")
            return []
    
    def get_context_for_query(
        self, 
        query: str, 
        top_k: int = 5,
        include_metadata: bool = False
    ) -> str:
        """Get formatted context string for a query."""
        chunks = self.retrieve_relevant_context(query, top_k=top_k)
        
        if not chunks:
            return "No relevant information found."
        
        context_parts = []
        for chunk in chunks:
            if include_metadata:
                source = chunk['metadata'].get('source', 'Unknown')
                similarity = chunk['metadata'].get('similarity', 0)
                context_parts.append(f"[Source: {source} | Relevance: {similarity:.2%}]\n{chunk['text']}")
            else:
                context_parts.append(chunk['text'])
        
        return "\n\n---\n\n".join(context_parts)
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the knowledge base."""
        count = self.collection.count()
        
        return {
            "total_chunks": count,
            "embedding_model": self.embedding_model_name,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "cache_dir": str(self.cache_dir)
        }
